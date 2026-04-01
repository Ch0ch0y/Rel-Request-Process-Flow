from fastapi import FastAPI, APIRouter, HTTPException, Depends, status, UploadFile, File, Request
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.staticfiles import StaticFiles
from fastapi.responses import StreamingResponse, FileResponse, HTMLResponse
from fastapi.templating import Jinja2Templates
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
import aiosqlite
import os
import io
import json
import logging
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict, EmailStr
from typing import List, Optional, Dict, Any
import uuid
from datetime import datetime, timezone, timedelta
import bcrypt
import jwt
from enum import Enum
import shutil
import zipfile
import xml.etree.ElementTree as _ET
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
try:
    import xlrd as _xlrd
    _XLRD_AVAILABLE = True
except ImportError:
    _XLRD_AVAILABLE = False

import sys
import re as _re
import asyncio
import threading

# Report generator for PowerPoint exports — parent dir added to path for generate_reports.py
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
try:
    from generate_reports import build_powerpoint
except ImportError:
    build_powerpoint = None

ROOT_DIR = Path(__file__).parent
MAINTENANCE_FLAG = ROOT_DIR / "maintenance.flag"

# Track server start time for uptime calculation
_SERVER_START = datetime.now(timezone.utc)

# Optional psutil for CPU/MEM metrics
try:
    import psutil as _psutil
    _PSUTIL_OK = True
except ImportError:
    _PSUTIL_OK = False

# â”€â”€ xlrd shim: lets _parse_excel_request_sheet work on legacy .xls files â”€â”€â”€â”€â”€â”€
class _XlrdCell:
    """Mimics openpyxl Cell.value so existing parsers work unchanged."""
    __slots__ = ('value',)
    def __init__(self, value): self.value = value

class _XlrdSheet:
    """Mimics openpyxl Worksheet with ws['B5'].value access and iter_rows."""
    def __init__(self, sheet, datemode):
        self._s = sheet
        self._dm = datemode

    def _cell_value(self, row, col):
        try:  
            if row >= self._s.nrows or col >= self._s.ncols:
                return None
            cell = self._s.cell(row, col)
            if _XLRD_AVAILABLE:
                if cell.ctype == _xlrd.XL_CELL_EMPTY or cell.ctype == _xlrd.XL_CELL_BLANK:
                    return None
                if cell.ctype == _xlrd.XL_CELL_DATE:
                    return _xlrd.xldate_as_datetime(cell.value, self._dm)
                if cell.ctype == _xlrd.XL_CELL_TEXT and cell.value == '':
                    return None
            return cell.value
        except (IndexError, TypeError):
            return None

    def __getitem__(self, cell_ref: str):
        m = _re.match(r'([A-Za-z]+)(\d+)', cell_ref)
        if not m:
            return _XlrdCell(None)
        col_str = m.group(1).upper()
        row = int(m.group(2)) - 1
        col = 0
        for ch in col_str:
            col = col * 26 + (ord(ch) - 64)
        col -= 1
        return _XlrdCell(self._cell_value(row, col))

    def iter_rows(self, values_only=False):
        """Yield rows as tuples of values (values_only=True) or _XlrdCell objects."""
        for r in range(self._s.nrows):
            if values_only:
                yield tuple(self._cell_value(r, c) for c in range(self._s.ncols))
            else:
                yield tuple(_XlrdCell(self._cell_value(r, c)) for c in range(self._s.ncols))

class _XlrdWb:
    """Mimics openpyxl Workbook with .active and .worksheets."""
    def __init__(self, book):
        self.worksheets = [_XlrdSheet(book.sheet_by_index(i), book.datemode)
                           for i in range(book.nsheets)]
        self.active = self.worksheets[0] if self.worksheets else None
load_dotenv(ROOT_DIR / '.env')

# SQLite database path
DB_PATH = os.environ.get('DB_PATH', str(ROOT_DIR / 'rel_database.db'))

# JWT Configuration
SECRET_KEY = os.environ.get('JWT_SECRET', 'your-secret-key-change-in-production')
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 60 * 24  # 24 hours

# Electrical test selectable conditions (defaults, can be overridden via settings)
ELECTRICAL_TEST_CONDITIONS = ["P4", "P1", "Customer Site", "Other 3rd Party"]
ELECTRICAL_TEST_ITEMS = ["E-Test"]

# Create the main app
app = FastAPI()
api_router = APIRouter(prefix="/api")
security = HTTPBearer(auto_error=False)

# ----- Database Setup -----
async def get_db():
    """Get a database connection."""
    db = await aiosqlite.connect(DB_PATH)
    db.row_factory = aiosqlite.Row
    await db.execute("PRAGMA journal_mode=WAL")
    await db.execute("PRAGMA foreign_keys=ON")
    return db

async def init_db():
    """Initialize database tables."""
    db = await aiosqlite.connect(DB_PATH)
    await db.executescript("""
        CREATE TABLE IF NOT EXISTS users (
            id TEXT PRIMARY KEY,
            email TEXT UNIQUE NOT NULL,
            username TEXT NOT NULL,
            password TEXT NOT NULL,
            role TEXT NOT NULL,
            approved INTEGER DEFAULT 0,
            created_at TEXT NOT NULL
        );

        CREATE TABLE IF NOT EXISTS login_logs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id TEXT NOT NULL,
            email TEXT NOT NULL,
            username TEXT NOT NULL,
            role TEXT NOT NULL,
            login_at TEXT NOT NULL,
            ip_address TEXT
        );

        CREATE TABLE IF NOT EXISTS requests (
            id TEXT PRIMARY KEY,
            request_number TEXT UNIQUE NOT NULL,
            request_type TEXT DEFAULT 'REL',
            classification TEXT DEFAULT '',
            originator TEXT DEFAULT '',
            plant TEXT DEFAULT '',
            device_name TEXT DEFAULT '',
            lot_no TEXT DEFAULT '',
            customer TEXT DEFAULT '',
            pkg_info TEXT DEFAULT '',
            automotive INTEGER DEFAULT 0,
            date_ltc TEXT,
            product_hierarchy TEXT,
            pdl TEXT,
            body_size_x REAL,
            body_size_y REAL,
            package_thickness REAL,
            ball_pitch REAL,
            ball_count INTEGER,
            lead_pitch REAL,
            lead_count INTEGER,
            total_ss TEXT,
            purpose TEXT DEFAULT '',
            engineer_special_instruction TEXT,
            deadline TEXT,
            created_by TEXT NOT NULL,
            created_by_username TEXT NOT NULL,
            created_at TEXT NOT NULL,
            updated_at TEXT NOT NULL,
            status TEXT DEFAULT 'pending',
            current_step INTEGER DEFAULT 1
        );

        CREATE TABLE IF NOT EXISTS process_steps (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            request_id TEXT NOT NULL REFERENCES requests(id) ON DELETE CASCADE,
            leg INTEGER DEFAULT 1,
            step_number INTEGER NOT NULL,
            step_name TEXT NOT NULL,
            status TEXT DEFAULT 'pending',
            started_at TEXT,
            completed_at TEXT,
            machine_no TEXT,
            rack_no TEXT,
            operator_id TEXT,
            tray_no TEXT,
            qty_in INTEGER,
            qty_out INTEGER,
            notes TEXT,
            attachments TEXT DEFAULT '[]',
            custom_fields TEXT DEFAULT '{}',
            UNIQUE(request_id, leg, step_number)
        );

        -- Triggers: automatically NULL out any non-ISO datetime values written to process_steps
        -- This is a database-level safety net that runs regardless of application code.
        CREATE TRIGGER IF NOT EXISTS trg_clean_started_at_insert
        AFTER INSERT ON process_steps
        WHEN NEW.started_at IS NOT NULL AND NEW.started_at NOT LIKE '____-__-__%'
        BEGIN
            UPDATE process_steps SET started_at = NULL WHERE id = NEW.id;
        END;

        CREATE TRIGGER IF NOT EXISTS trg_clean_completed_at_insert
        AFTER INSERT ON process_steps
        WHEN NEW.completed_at IS NOT NULL AND NEW.completed_at NOT LIKE '____-__-__%'
        BEGIN
            UPDATE process_steps SET completed_at = NULL WHERE id = NEW.id;
        END;

        CREATE TRIGGER IF NOT EXISTS trg_clean_started_at_update
        AFTER UPDATE OF started_at ON process_steps
        WHEN NEW.started_at IS NOT NULL AND NEW.started_at NOT LIKE '____-__-__%'
        BEGIN
            UPDATE process_steps SET started_at = NULL WHERE id = NEW.id;
        END;

        CREATE TRIGGER IF NOT EXISTS trg_clean_completed_at_update
        AFTER UPDATE OF completed_at ON process_steps
        WHEN NEW.completed_at IS NOT NULL AND NEW.completed_at NOT LIKE '____-__-__%'
        BEGIN
            UPDATE process_steps SET completed_at = NULL WHERE id = NEW.id;
        END;

        CREATE TABLE IF NOT EXISTS settings (
            id INTEGER PRIMARY KEY CHECK (id = 1),
            app_name TEXT DEFAULT 'Rel Request Process Flow',
            app_logo TEXT,
            company_name TEXT,
            contact_email TEXT,
            process_steps TEXT,
            custom_fields TEXT DEFAULT '{}',
            updated_at TEXT
        );

        CREATE TABLE IF NOT EXISTS role_permissions (
            role TEXT NOT NULL,
            permission TEXT NOT NULL,
            granted INTEGER DEFAULT 1,
            PRIMARY KEY (role, permission)
        );

        CREATE TABLE IF NOT EXISTS backup_tracking (
            id INTEGER PRIMARY KEY CHECK (id = 1),
            last_critical_backup_at TEXT,
            last_backup_request_count INTEGER DEFAULT 0,
            critical_backup_required INTEGER DEFAULT 0,
            last_backup_downloaded INTEGER DEFAULT 0
        );

        CREATE TABLE IF NOT EXISTS machines (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            machine_no TEXT NOT NULL,
            description TEXT NOT NULL
        );

        CREATE TABLE IF NOT EXISTS employees (
            id TEXT PRIMARY KEY,
            name TEXT NOT NULL,
            position TEXT NOT NULL
        );

        CREATE TABLE IF NOT EXISTS relmon_sheet_data (
            site TEXT NOT NULL,
            sheet TEXT NOT NULL,
            rows_json TEXT NOT NULL,
            merges_json TEXT,
            form_json TEXT DEFAULT '{}',
            created_at TEXT NOT NULL,
            updated_at TEXT NOT NULL,
            updated_by TEXT,
            PRIMARY KEY (site, sheet)
        );
    """)
    await db.commit()

    # Initialize backup tracking record
    cursor = await db.execute("SELECT id FROM backup_tracking WHERE id = 1")
    if not await cursor.fetchone():
        await db.execute(
            "INSERT INTO backup_tracking (id, critical_backup_required, last_backup_downloaded) VALUES (1, 0, 0)"
        )
        await db.commit()

    # --- Migrations: add columns if missing ---
    try:
        await db.execute("ALTER TABLE requests ADD COLUMN note TEXT")
        await db.commit()
    except Exception:
        pass  # column already exists

    try:
        await db.execute("ALTER TABLE users ADD COLUMN approved INTEGER DEFAULT 0")
        await db.commit()
    except Exception:
        pass  # column already exists

    for col_def in [
        "ALTER TABLE users ADD COLUMN position TEXT DEFAULT ''",
        "ALTER TABLE users ADD COLUMN contact_email TEXT DEFAULT ''",
        "ALTER TABLE users ADD COLUMN plant TEXT DEFAULT ''",
        "ALTER TABLE users ADD COLUMN manager TEXT DEFAULT ''",
        "ALTER TABLE users ADD COLUMN last_seen TEXT",
        "ALTER TABLE users ADD COLUMN blocked INTEGER DEFAULT 0",
        "ALTER TABLE users ADD COLUMN security_question TEXT",
        "ALTER TABLE users ADD COLUMN security_answer TEXT",
        "ALTER TABLE users ADD COLUMN user_status TEXT DEFAULT 'pending'",
        "ALTER TABLE users ADD COLUMN declined_at TEXT DEFAULT NULL",
    ]:
        try:
            await db.execute(col_def)
            await db.commit()
        except Exception:
            pass  # column already exists

    try:
        await db.execute("ALTER TABLE settings ADD COLUMN tech_auth_code TEXT DEFAULT '735522'")
        await db.commit()
    except Exception:
        pass  # column already exists

    try:
        await db.execute("ALTER TABLE settings ADD COLUMN process_presets TEXT")
        await db.commit()
    except Exception:
        pass  # column already exists

    try:
        await db.execute("ALTER TABLE requests ADD COLUMN retention_details TEXT DEFAULT NULL")
        await db.commit()
    except Exception:
        pass  # column already exists

    try:
        await db.execute("ALTER TABLE requests ADD COLUMN analysis_notes TEXT DEFAULT NULL")
        await db.commit()
    except Exception:
        pass  # column already exists

    try:
        await db.execute("ALTER TABLE users ADD COLUMN avatar TEXT DEFAULT NULL")
        await db.commit()
    except Exception:
        pass  # column already exists

    try:
        await db.execute("ALTER TABLE requests ADD COLUMN approved_at TEXT DEFAULT NULL")
        await db.commit()
    except Exception:
        pass  # column already exists

    try:
        await db.execute("ALTER TABLE requests ADD COLUMN planner_est_start TEXT DEFAULT NULL")
        await db.execute("ALTER TABLE requests ADD COLUMN planner_est_end TEXT DEFAULT NULL")
        await db.execute("ALTER TABLE requests ADD COLUMN planner_note TEXT DEFAULT NULL")
        await db.commit()
    except Exception:
        pass  # columns already exist

    try:
        await db.execute("ALTER TABLE requests ADD COLUMN discontinued_at TEXT DEFAULT NULL")
        await db.execute("ALTER TABLE requests ADD COLUMN discontinued_by TEXT DEFAULT NULL")
        await db.execute("ALTER TABLE requests ADD COLUMN discontinued_reason TEXT DEFAULT NULL")
        await db.commit()
    except Exception:
        pass  # columns already exist

    try:
        await db.execute("ALTER TABLE requests ADD COLUMN original_rr_number TEXT DEFAULT NULL")
        await db.commit()
    except Exception:
        pass  # column already exists

    # Masterlist planner fields on requests
    try:
        await db.execute("ALTER TABLE requests ADD COLUMN ww TEXT DEFAULT NULL")
        await db.execute("ALTER TABLE requests ADD COLUMN lc_bc TEXT DEFAULT NULL")
        await db.execute("ALTER TABLE requests ADD COLUMN test_level TEXT DEFAULT NULL")
        await db.execute("ALTER TABLE requests ADD COLUMN ml_qty TEXT DEFAULT NULL")
        await db.execute("ALTER TABLE requests ADD COLUMN num_days TEXT DEFAULT NULL")
        await db.execute("ALTER TABLE requests ADD COLUMN num_legs TEXT DEFAULT NULL")
        await db.execute("ALTER TABLE requests ADD COLUMN recommit TEXT DEFAULT NULL")
        await db.commit()
    except Exception:
        pass  # columns already exist

    # Add request_type column if missing
    try:
        await db.execute("ALTER TABLE requests ADD COLUMN request_type TEXT DEFAULT 'REL'")
        await db.commit()
    except Exception:
        pass  # column already exists

    # Seed machines table from defaults if empty
    cursor = await db.execute("SELECT COUNT(*) FROM machines")
    row = await cursor.fetchone()
    if row[0] == 0:
        default_machines = [
            ('RXN-001','3D XRAY'),('RPL-001','AUTO POLISHER'),('RMD-002','B.SHEAR / W.PULL'),
            ('RDL-001','DEGASSER'),('RTT-001','DIGITAL SERIAL ANALYZER'),('RDJ-001','DRILLING/MILLING M/C'),
            ('RDC-001','DRY BOX'),('REH-001','EMISSION MICROSCOPE'),('RRV-001','REFLOW'),
            ('RFN-003','FTIR'),('RHE-001','HAST'),('RHH-001','HAST'),('RHH-002','HAST'),
            ('RHH-004','HAST'),('RHH-005','HAST'),('RHH-006','HAST'),('RHT-001','HAST'),
            ('RHT-002','HAST'),('RMN-001','HIGH OPTICAL MICROSCOPE'),('RMO-010','HIGH OPTICAL MICROSCOPE'),
            ('ROE-001','HTS'),('ROE-002','HTS'),('ROE-003','HTS'),('ROE-004','HTS'),
            ('ROE-005','HTS'),('ROE-006','HTS'),('ROE-007','HTS'),
            ('RPJ-001','ION BEAM POLISHER'),('RIC-001','ION COATER'),('RIT-001','ION COATER'),
            ('RIH-002','ION MILL'),('RIH-001','ION SPUTTER'),('RIH-004','ION SPUTTER'),
            ('RIS-001','IONIZER'),('RIS-002','IONIZER'),('RIS-003','IONIZER'),('RIS-004','IONIZER'),
            ('RIS-005','IONIZER'),('RIS-006','IONIZER'),('RIS-007','IONIZER'),
            ('RNL-001','LASER AUTO-DECAPSULATOR'),('RCL-001','LASER AUTO-DECAPSULATOR'),
            ('RLT-001','LOCK-IN TOMOGRAPHY'),('RMO-012','LOW MAG SCOPE'),
            ('RME-001','LOW MAG SCOPE'),('RME-002','LOW MAG SCOPE'),
            ('RMO-007','LOW OPT. MICROSCOPE'),('RMO-008','LOW OPT. MICROSCOPE'),
            ('RMO-004','LOW OPT. MICROSCOPE'),('RMO-005','LOW OPT. MICROSCOPE'),
            ('RMO-006','LOW OPT. MICROSCOPE'),('RMO-009','LOW OPT. MICROSCOPE'),
            ('RMV-001','LOW POWER SCOPE (ICAPS)'),('RSO-001','MEASURING MIC.'),
            ('RMM-001','MILLI OHMS RESISTANCE TESTER'),
            ('ROH-001','O/S TESTER'),('ROH-002','O/S TESTER'),('ROH-003','O/S TESTER'),
            ('ROH-004','O/S TESTER'),('RTI-001','O/S TESTER'),('ROC-001','OS TESTER HANDLER'),
            ('RPB-004','PLASMA DECAPPER'),('RMJ-001','PLASMA ETCHING MACHINE'),
            ('RPB-005','POLISHER'),('RPB-007','POLISHER'),('RRH-001','REFLOW'),
            ('RSS-003','SAT'),('RSS-005','SAT'),('RSS-006','SAT'),
            ('RSH-002','SEM'),('RSF-001','SEM/EDX/FIB'),('RSH-003','SEM/EDX'),
            ('RCE-002','TC'),('RCE-003','TC'),('RCE-004','TC'),('RCY-001','TC'),
            ('RCW-002','TC'),('RCW-003','TC'),
            ('RTE-001','TH'),('RTE-003','TH'),('RTE-005','TH'),('RTE-006','TH'),
            ('RTE-007','TH'),('RTE-008','TH'),('RTW-001','TH'),('RTW-002','TH'),
            ('RTW-003','TH'),('RTW-005','TH'),('RTW-007','TH'),('RTV-001','TH'),('RTV-002','TH'),
            ('RUS-001','ULTRA SLICE'),('RUA-001','ULTRASONIC CLEANER'),
            ('RPE-001','VAR. GRIN. POLISHER'),('RWM-002','WEIGHING BALANCE'),
            ('RMD-003','WIREPULL / BALLSHEAR TESTER'),
            ('RHP-001','HOT PLATE'),('RHP-002','HOT PLATE'),('RMK-001','3D SCOPE'),
        ]
        await db.executemany("INSERT INTO machines (machine_no, description) VALUES (?,?)", default_machines)
        await db.commit()

    # Seed employees table from defaults if empty
    cursor = await db.execute("SELECT COUNT(*) FROM employees")
    row = await cursor.fetchone()
    if row[0] == 0:
        default_employees = [
            ('947241','Celia Corpuz','Manager'),
            ('105445','Conrado Hidalgo','Sr. FA Engr'),
            ('240097','Pamela Satur','Rel Engr'),
            ('240167','Shelah Mae Perez','Rel Engr'),
            ('240168','Clarence Joshua Ramirez','FA Engr'),
            ('250296','Allyza Nicole Humirang','Rel Engr'),
            ('960853','Loreta Veran','Sr. Rel Engr'),
            ('993404','Lea Dalanon','FA Operation Engr'),
            ('982308','Esmeria, Erwin','FA ES P3'),
            ('175081','Hatulan, Irving','FA ES P3'),
            ('175075','Delos Santos, Charito','FA ES P3'),
            ('105294','Bermiso, Ricky','FA ES P3'),
            ('240427','Monterosa, Shaira','FA ES P3'),
            ('175083','Supapo, Bryane','FA ES P3'),
            ('175087','Ortiz, Van Joven','FA ES P3'),
            ('175198','Del Rosario, Wowie','FA ES P3'),
            ('175082','Foronda, Georjan','FA ES P3'),
            ('202544','Salazar, Jeronel','FA ES P3'),
            ('250125','Dela Rosa, Rowell','FA ES P3'),
            ('250158','Remigio, Alcen','FA ES P3'),
            ('250135','Trinidad, Maricel','REL ES'),
            ('155253','Delos Santos, Chlarissa','REL ES'),
            ('145087','Santiago, Kimberly Rose','REL ES'),
            ('155252','De Mesa, Rosemarie','REL ES'),
            ('175088','Velitario, Madelyn','REL ES'),
            ('145084','Arcega, Johnrey','REL ES'),
            ('155420','Reig, Leonito','REL ES'),
            ('175089','Barrera, Marissa','REL ES'),
            ('175074','Cruz, Jasthine Mae','REL ES'),
            ('230076','Rizano, Jan Mark','REL ES'),
            ('250136','Semillano, Adrian','REL ES'),
            ('252523','Balcita, Jeriel','REL ES'),
            ('981931','Reggie Quito','REL ES'),
            ('155389','Roy Tiquis','REL ES'),
            ('180966','Eduardo Visca','REL ES'),
        ]
        await db.executemany("INSERT OR IGNORE INTO employees (id, name, position) VALUES (?,?,?)", default_employees)
        await db.commit()

    try:
        await db.execute("ALTER TABLE process_steps ADD COLUMN qty_in INTEGER")
        await db.execute("ALTER TABLE process_steps ADD COLUMN qty_out INTEGER")
        await db.execute("ALTER TABLE process_steps ADD COLUMN rack_no TEXT")
        await db.commit()
    except Exception:
        pass  # columns already exist

    try:
        await db.execute("ALTER TABLE process_steps ADD COLUMN updated_by TEXT")
        await db.commit()
    except Exception:
        pass  # column already exists

    # Migration: add leg column and fix UNIQUE constraint
    try:
        cursor = await db.execute("PRAGMA table_info(process_steps)")
        cols = [row[1] for row in await cursor.fetchall()]
        if 'leg' not in cols:
            await db.executescript("""
                CREATE TABLE process_steps_new (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    request_id TEXT NOT NULL REFERENCES requests(id) ON DELETE CASCADE,
                    leg INTEGER DEFAULT 1,
                    step_number INTEGER NOT NULL,
                    step_name TEXT NOT NULL,
                    status TEXT DEFAULT 'pending',
                    started_at TEXT,
                    completed_at TEXT,
                    machine_no TEXT,
                    operator_id TEXT,
                    tray_no TEXT,
                    qty_in INTEGER,
                    qty_out INTEGER,
                    notes TEXT,
                    attachments TEXT DEFAULT '[]',
                    custom_fields TEXT DEFAULT '{}',
                    UNIQUE(request_id, leg, step_number)
                );
                INSERT INTO process_steps_new (id, request_id, leg, step_number, step_name, status,
                    started_at, completed_at, machine_no, operator_id, tray_no, qty_in, qty_out,
                    notes, attachments, custom_fields)
                SELECT id, request_id, 1, step_number, step_name, status,
                    started_at, completed_at, machine_no, operator_id, tray_no, qty_in, qty_out,
                    notes, attachments, custom_fields
                FROM process_steps;
                DROP TABLE process_steps;
                ALTER TABLE process_steps_new RENAME TO process_steps;
            """)
            await db.commit()
    except Exception as e:
        logging.warning(f"Leg column migration: {e}")

    # --- Seed default admin account ---
    cursor = await db.execute("SELECT id FROM users WHERE email = ?", ("admin@amkor.com",))
    if not await cursor.fetchone():
        admin_id = str(uuid.uuid4())
        admin_pw = hash_password("Adminn")
        now = datetime.now(timezone.utc).isoformat()
        await db.execute(
            "INSERT INTO users (id, email, username, password, role, approved, created_at) VALUES (?,?,?,?,?,?,?)",
            (admin_id, "admin@amkor.com", "Admin", admin_pw, "Admin", 1, now)
        )
        await db.commit()
        logging.info("Default admin account created: admin@amkor.com")

    # --- Seed default role permissions ---
    ALL_PERMISSIONS = [
        'create_request', 'edit_request', 'delete_request',
        'update_steps', 'manage_steps',
        'manage_users', 'manage_settings', 'import_requests', 'manage_backups',
    ]
    DEFAULT_ROLE_PERMISSIONS = {
        'Reliability Engineer': ALL_PERMISSIONS.copy(),
        'Failure Analysis': [],
        'Technician': ['update_steps'],
        'Planner': ['update_steps'],
    }
    for role, perms in DEFAULT_ROLE_PERMISSIONS.items():
        # Only seed rows that don't exist yet â€” never overwrite admin-saved permissions
        for p in ALL_PERMISSIONS:
            granted = 1 if p in perms else 0
            try:
                await db.execute(
                    "INSERT INTO role_permissions (role, permission, granted) VALUES (?, ?, ?) "
                    "ON CONFLICT(role, permission) DO NOTHING",
                    (role, p, granted)
                )
            except Exception:
                pass
    await db.commit()

    # masterlist_2026 table
    try:
        await db.execute("""
            CREATE TABLE IF NOT EXISTS masterlist_2026 (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                ww TEXT,
                date_received TEXT,
                rrs_no TEXT,
                purpose TEXT,
                qual_type TEXT,
                customer TEXT,
                pkg_type TEXT,
                lc_bc TEXT,
                rr_agile_no TEXT,
                test_level TEXT,
                qty TEXT,
                num_days TEXT,
                num_legs TEXT,
                est_start TEXT,
                est_completion TEXT,
                recommit TEXT,
                planner_remarks TEXT,
                uploaded_at TEXT
            )
        """)
        await db.commit()
    except Exception:
        pass

    # Migrate: drop old training_masterlist if it exists (replaced by masterlist_2026)
    try:
        await db.execute("DROP TABLE IF EXISTS training_masterlist")
        await db.commit()
    except Exception:
        pass

    # Add employee tracking columns to login_logs
    for _col_def in [
        "ALTER TABLE login_logs ADD COLUMN employee_id TEXT DEFAULT ''",
        "ALTER TABLE login_logs ADD COLUMN employee_name TEXT DEFAULT ''",
    ]:
        try:
            await db.execute(_col_def)
            await db.commit()
        except Exception:
            pass  # column already exists

    # Technician sessions table (tracks active guest logins)
    try:
        await db.execute("""
            CREATE TABLE IF NOT EXISTS technician_sessions (
                employee_id       TEXT PRIMARY KEY,
                employee_name     TEXT NOT NULL DEFAULT '',
                employee_position TEXT NOT NULL DEFAULT '',
                last_active       TEXT NOT NULL,
                login_at          TEXT NOT NULL
            )
        """)
        await db.commit()
    except Exception:
        pass

    # Add employee tracking columns to login_logs
    for _col_def in [
        "ALTER TABLE login_logs ADD COLUMN employee_id TEXT DEFAULT ''",
        "ALTER TABLE login_logs ADD COLUMN employee_name TEXT DEFAULT ''",
    ]:
        try:
            await db.execute(_col_def)
            await db.commit()
        except Exception:
            pass  # column already exists

    # Technician sessions table (tracks active guest logins)
    try:
        await db.execute("""
            CREATE TABLE IF NOT EXISTS technician_sessions (
                employee_id       TEXT PRIMARY KEY,
                employee_name     TEXT NOT NULL DEFAULT '',
                employee_position TEXT NOT NULL DEFAULT '',
                last_active       TEXT NOT NULL,
                login_at          TEXT NOT NULL
            )
        """)
        await db.commit()
    except Exception:
        pass

    await db.close()
    logging.info(f"SQLite database initialized at {DB_PATH}")

# ----- Enums -----
class UserRole(str, Enum):
    ADMIN = "Admin"
    RELIABILITY_ENGINEER = "Reliability Engineer"
    FAILURE_ANALYSIS = "Failure Analysis"
    TECHNICIAN = "Technician"
    PLANNER = "Planner"

class StepStatus(str, Enum):
    PENDING = "pending"
    IN_PROGRESS = "in_progress"
    COMPLETED = "completed"
    FAILED = "failed"
    HOLD = "hold"

# ----- Pydantic Models -----
class User(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    email: EmailStr
    username: str
    role: UserRole
    approved: bool = False
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    position: Optional[str] = ""
    contact_email: Optional[str] = ""
    plant: Optional[str] = ""
    manager: Optional[str] = ""
    is_guest: bool = False
    last_seen: Optional[datetime] = None
    blocked: bool = False
    user_status: Optional[str] = 'pending'  # pending, approved, hold, lock, declined
    avatar: Optional[str] = None  # base64 encoded profile picture
    employee_id: Optional[str] = None
    employee_name: Optional[str] = None

class UserCreate(BaseModel):
    email: EmailStr
    username: str
    password: str
    role: UserRole

class ProfileUpdate(BaseModel):
    username: Optional[str] = None
    position: Optional[str] = None
    contact_email: Optional[str] = None
    plant: Optional[str] = None
    manager: Optional[str] = None
    avatar: Optional[str] = None  # base64 encoded profile picture (pass empty string to clear)

class UserLogin(BaseModel):
    email: EmailStr
    password: str

class Token(BaseModel):
    access_token: str
    token_type: str
    user: User

class ProcessStep(BaseModel):
    step_number: int
    step_name: str
    leg: int = 1
    status: StepStatus = StepStatus.PENDING
    started_at: Optional[datetime] = None
    completed_at: Optional[datetime] = None
    machine_no: Optional[str] = None
    rack_no: Optional[str] = None
    operator_id: Optional[str] = None
    tray_no: Optional[str] = None
    qty_in: Optional[int] = None
    qty_out: Optional[int] = None
    notes: Optional[str] = None
    attachments: Optional[Any] = Field(default_factory=dict)  # dict of category->list[str], or legacy list[str]
    custom_fields: Optional[Dict[str, Any]] = Field(default_factory=dict)

class RelRequest(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    request_number: str
    request_type: str = 'REL'
    classification: Optional[str] = ""
    originator: Optional[str] = ""
    plant: Optional[str] = ""
    device_name: Optional[str] = ""
    lot_no: Optional[str] = ""
    customer: Optional[str] = ""
    pkg_info: Optional[str] = ""
    automotive: bool = False
    date_ltc: Optional[str] = None
    product_hierarchy: Optional[str] = None
    pdl: Optional[str] = None
    body_size_x: Optional[float] = None
    body_size_y: Optional[float] = None
    package_thickness: Optional[float] = None
    ball_pitch: Optional[float] = None
    ball_count: Optional[int] = None
    lead_pitch: Optional[float] = None
    lead_count: Optional[int] = None
    total_ss: Optional[str] = None
    purpose: Optional[str] = ""
    engineer_special_instruction: Optional[str] = None
    deadline: Optional[str] = None
    created_by: str
    created_by_username: str
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    updated_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    status: str = "pending"
    current_step: int = 1
    note: Optional[str] = None
    retention_details: Optional[str] = None
    analysis_notes: Optional[str] = None
    approved_at: Optional[str] = None
    planner_est_start: Optional[str] = None
    planner_est_end: Optional[str] = None
    planner_note: Optional[str] = None
    discontinued_at: Optional[str] = None
    discontinued_by: Optional[str] = None
    discontinued_reason: Optional[str] = None
    original_rr_number: Optional[str] = None
    steps: List[ProcessStep]

class RequestCreate(BaseModel):
    request_number: Optional[str] = None
    request_type: Optional[str] = 'REL'
    status: Optional[str] = None
    classification: Optional[str] = None
    originator: Optional[str] = None
    plant: Optional[str] = None
    device_name: Optional[str] = None
    lot_no: Optional[str] = None
    customer: Optional[str] = None
    pkg_info: Optional[str] = None
    automotive: Optional[bool] = False
    date_ltc: Optional[str] = None
    product_hierarchy: Optional[str] = None
    pdl: Optional[str] = None
    body_size_x: Optional[float] = None
    body_size_y: Optional[float] = None
    package_thickness: Optional[float] = None
    ball_pitch: Optional[float] = None
    ball_count: Optional[int] = None
    lead_pitch: Optional[float] = None
    lead_count: Optional[int] = None
    total_ss: Optional[str] = None
    purpose: Optional[str] = None
    engineer_special_instruction: Optional[str] = None
    deadline: Optional[str] = None
    note: Optional[str] = None
    retention_details: Optional[str] = None
    analysis_notes: Optional[str] = None
    planner_est_start: Optional[str] = None
    planner_est_end: Optional[str] = None
    planner_note: Optional[str] = None
    original_rr_number: Optional[str] = None
    custom_steps: Optional[List[str]] = None  # e.g. ["Incoming Inspection", "Visual", "SAT"]

class StepUpdate(BaseModel):
    step_name: Optional[str] = None
    status: Optional[StepStatus] = None
    started_at: Optional[str] = None
    completed_at: Optional[str] = None
    machine_no: Optional[str] = None
    rack_no: Optional[str] = None
    operator_id: Optional[str] = None
    tray_no: Optional[str] = None
    qty_in: Optional[int] = None
    qty_out: Optional[int] = None
    notes: Optional[str] = None
    attachments: Optional[Any] = None  # dict of category->list[str], or legacy list[str]
    custom_fields: Optional[Dict[str, Any]] = None

class DashboardStats(BaseModel):
    total_requests: int
    active_requests: int
    completed_requests: int
    pending_requests: int
    ongoing_requests: int
    delayed_requests: int
    upcoming_deadline_requests: int
    incoming_inspection_count: int = 0
    visual_count: int = 0
    sat_count: int = 0
    bake_count: int = 0
    hts_count: int = 0
    hold_count: int = 0
    hold_requests_list: List[Dict[str, Any]] = []
    step_progress: List[Dict[str, Any]] = []
    recent_activity: List[Dict[str, Any]]
    delayed_requests_list: List[Dict[str, Any]]
    upcoming_deadline_list: List[Dict[str, Any]]
    noticed_requests_list: List[Dict[str, Any]] = []
    review_requests: int = 0
    approval_requests: int = 0
    testing_requests: int = 0
    analysis_requests: int = 0
    backup_warning: Optional[str] = None
    backup_warning_level: Optional[str] = None  # 'info', 'warning', 'critical'
    requires_critical_backup: bool = False

class AppSettings(BaseModel):
    model_config = ConfigDict(extra="ignore")
    app_name: str = "Rel Request Process Flow"
    app_logo: Optional[str] = None
    company_name: Optional[str] = None
    contact_email: Optional[str] = None
    process_steps: List[str] = Field(default_factory=lambda: [step["step_name"] for step in DEFAULT_STEPS])
    process_presets: Optional[List[Any]] = None
    custom_fields: Dict[str, Any] = {}
    tech_auth_code: str = "735522"
    updated_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class SettingsUpdate(BaseModel):
    model_config = ConfigDict(extra="ignore")
    app_name: Optional[str] = None
    app_logo: Optional[str] = None
    company_name: Optional[str] = None
    contact_email: Optional[str] = None
    process_steps: Optional[List[str]] = None
    process_presets: Optional[List[Any]] = None
    custom_fields: Optional[Dict[str, Any]] = None
    tech_auth_code: Optional[str] = None

# ----- Helper functions -----
def hash_password(password: str) -> str:
    return bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')

def verify_password(plain_password: str, hashed_password: str) -> bool:
    return bcrypt.checkpw(plain_password.encode('utf-8'), hashed_password.encode('utf-8'))

def create_access_token(data: dict) -> str:
    to_encode = data.copy()
    expire = datetime.now(timezone.utc) + timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    to_encode.update({"exp": expire})
    return jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)

async def get_current_user(
    request: Request,
    credentials: Optional[HTTPAuthorizationCredentials] = Depends(security),
) -> User:
    # Try Bearer header first, then fall back to httpOnly cookie
    token = None
    if credentials:
        token = credentials.credentials
    if not token or token in ('null', 'undefined'):
        token = request.cookies.get("access_token")
    if not token:
        raise HTTPException(status_code=401, detail="Not authenticated")
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        user_id: str = payload.get("sub")
        if user_id is None:
            raise HTTPException(status_code=401, detail="Invalid authentication credentials")
    except HTTPException:
        raise
    except jwt.exceptions.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token has expired")
    except jwt.exceptions.PyJWTError:
        raise HTTPException(status_code=401, detail="Could not validate credentials")
    except Exception:
        raise HTTPException(status_code=401, detail="Could not validate credentials")

    # Guest token â€” no DB lookup or pydantic validation needed
    if payload.get("is_guest"):
        emp_id   = payload.get("employee_id", "") or ""
        emp_name = payload.get("employee_name", "") or ""
        display  = emp_name or "Technician"
        return User.model_construct(
            id="guest",
            email="guest@technician.local",
            username=display,
            role=UserRole.TECHNICIAN,
            approved=True,
            is_guest=True,
            created_at=datetime.now(timezone.utc),
            position="",
            contact_email="",
            plant="",
            manager="",
            employee_id=emp_id,
            employee_name=emp_name,
        )

    db = await get_db()
    try:
        cursor = await db.execute(
            "SELECT id, email, username, role, created_at, approved, position, contact_email, plant, manager, blocked, avatar FROM users WHERE id = ?", (user_id,)
        )
        row = await cursor.fetchone()
        if row is None:
            raise HTTPException(status_code=401, detail="User not found")
        if row[10]:
            raise HTTPException(status_code=403, detail="Your account has been blocked. Please contact an administrator.")
        return User(id=row[0], email=row[1], username=row[2], role=row[3],
                    created_at=datetime.fromisoformat(row[4]),
                    approved=bool(row[5]) if row[5] is not None else False,
                    position=row[6] or '', contact_email=row[7] or '',
                    plant=row[8] or '', manager=row[9] or '',
                    blocked=bool(row[10]) if row[10] is not None else False,
                    avatar=row[11] or None)
    finally:
        await db.close()

class _HTMLAuthRequired(Exception):
    """Raised when an HTML page route requires authentication but the cookie is missing/invalid."""
    pass

async def get_current_user_html(request: Request) -> User:
    """Auth dependency for HTML page routes — reads JWT from the access_token cookie.
    Raises _HTMLAuthRequired (handled globally as a /login redirect) on failure."""
    token = request.cookies.get("access_token")
    if not token:
        raise _HTMLAuthRequired()
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        user_id: str = payload.get("user_id") or payload.get("sub")
        if not user_id:
            raise _HTMLAuthRequired()
    except (jwt.exceptions.ExpiredSignatureError, jwt.exceptions.PyJWTError):
        raise _HTMLAuthRequired()
    db = await get_db()
    try:
        cursor = await db.execute(
            "SELECT id, email, username, role, created_at, approved, position, contact_email, plant, manager, blocked, avatar FROM users WHERE id = ?",
            (user_id,)
        )
        row = await cursor.fetchone()
        if row is None or row[10]:
            raise _HTMLAuthRequired()
        return User(
            id=row[0], email=row[1], username=row[2], role=row[3],
            created_at=datetime.fromisoformat(row[4]),
            approved=bool(row[5]) if row[5] is not None else False,
            position=row[6] or '', contact_email=row[7] or '',
            plant=row[8] or '', manager=row[9] or '',
            blocked=bool(row[10]) if row[10] is not None else False,
            avatar=row[11] or None
        )
    finally:
        await db.close()

def require_role(allowed_roles: List[UserRole]):
    async def role_checker(current_user: User = Depends(get_current_user)) -> User:
        if current_user.role not in allowed_roles:
            raise HTTPException(status_code=403, detail="Insufficient permissions")
        return current_user
    return role_checker

# All available permissions
ALL_PERMISSIONS = [
    'create_request', 'edit_request', 'delete_request',
    'update_steps', 'manage_steps',
    'manage_users', 'manage_settings', 'import_requests', 'manage_backups',
]

def require_permission(permission: str):
    """Check if the current user's role has the given permission.
    Admin always has all permissions. Guest Technicians are allowed update_steps."""
    async def permission_checker(current_user: User = Depends(get_current_user)) -> User:
        # Admin always has all permissions
        if current_user.role == UserRole.ADMIN:
            return current_user
        # Guest technician: allow read-related + update_steps
        if current_user.is_guest and permission in ('update_steps',):
            return current_user
        db = await get_db()
        try:
            cursor = await db.execute(
                "SELECT granted FROM role_permissions WHERE role = ? AND permission = ?",
                (current_user.role, permission)
            )
            row = await cursor.fetchone()
            if not row or not row[0]:
                raise HTTPException(status_code=403, detail="Insufficient permissions")
            return current_user
        finally:
            await db.close()
    return permission_checker

# Initialize default steps
DEFAULT_STEPS = [
    {"step_number": 1,  "step_name": "Incoming Inspection", "status": StepStatus.PENDING},
    {"step_number": 2,  "step_name": "Visual",              "status": StepStatus.PENDING},
    {"step_number": 3,  "step_name": "Serialize Samples",   "status": StepStatus.PENDING},
    {"step_number": 4,  "step_name": "O/S",                 "status": StepStatus.PENDING},
    {"step_number": 5,  "step_name": "SAT",                 "status": StepStatus.PENDING},
    {"step_number": 6,  "step_name": "Bake",                "status": StepStatus.PENDING},
    {"step_number": 7,  "step_name": "T & H Soak",          "status": StepStatus.PENDING},
    {"step_number": 8,  "step_name": "Reflow",              "status": StepStatus.PENDING},
    {"step_number": 9,  "step_name": "Electrical Test",     "status": StepStatus.PENDING},
    {"step_number": 10, "step_name": "SAT",                 "status": StepStatus.PENDING},
    {"step_number": 11, "step_name": "O/S",                 "status": StepStatus.PENDING},
    {"step_number": 12, "step_name": "Visual",              "status": StepStatus.PENDING},
    {"step_number": 13, "step_name": "Reliability Test",    "status": StepStatus.PENDING},
    {"step_number": 14, "step_name": "SAT",                 "status": StepStatus.PENDING},
    {"step_number": 15, "step_name": "O/S",                 "status": StepStatus.PENDING},
    {"step_number": 16, "step_name": "Visual",              "status": StepStatus.PENDING},
]

# Available step types that users can choose from
AVAILABLE_STEP_NAMES = [
    "Incoming Inspection", "Visual", "Serialize Samples",
    "O/S", "SAT", "Bake", "Dry Bake", "T & H Soak", "HTS", "Reliability Test", "Reflow",
    "Electrical Test",
    "Moisture Resistance Test", "Preconditioning (Precon)", "Forced Convection Reflow (FCR)",
    "Whisker Test", "Staging",
]

# Column order for SELECT * FROM requests
REQ_COLS = [
    'id', 'request_number', 'request_type', 'classification', 'originator', 'plant', 'device_name',
    'lot_no', 'customer', 'pkg_info', 'automotive', 'date_ltc', 'product_hierarchy',
    'pdl', 'body_size_x', 'body_size_y', 'package_thickness', 'ball_pitch', 'ball_count',
    'lead_pitch', 'lead_count', 'total_ss', 'purpose', 'engineer_special_instruction',
    'deadline', 'created_by', 'created_by_username', 'created_at', 'updated_at',
    'status', 'current_step', 'note', 'retention_details', 'analysis_notes', 'approved_at',
    'planner_est_start', 'planner_est_end', 'planner_note',
    'discontinued_at', 'discontinued_by', 'discontinued_reason',
    'ww', 'lc_bc', 'test_level', 'ml_qty', 'num_days', 'num_legs', 'recommit',
    'original_rr_number',
]

def _safe_isoparse(value):
    """Parse an ISO datetime string, returning None for missing or malformed values."""
    if not value:
        return None
    try:
        return datetime.fromisoformat(value)
    except (ValueError, TypeError):
        logging.warning(f"Ignoring malformed datetime value: {value!r}")
        return None

async def _row_to_request_dict(db, row):
    """Convert a request row + its steps into a dict suitable for the RelRequest model."""
    req = {}
    row_keys = row.keys() if hasattr(row, 'keys') else []
    for col in REQ_COLS:
        try:
            req[col] = row[col]
        except (IndexError, KeyError):
            req[col] = None
    req['automotive'] = bool(req['automotive'])
    req['created_at'] = _safe_isoparse(req['created_at']) or datetime.now(timezone.utc)
    req['updated_at'] = _safe_isoparse(req['updated_at']) or datetime.now(timezone.utc)

    cursor = await db.execute(
        "SELECT id, step_number, step_name, leg, status, started_at, completed_at, machine_no, rack_no, operator_id, "
        "tray_no, qty_in, qty_out, notes, attachments, custom_fields FROM process_steps "
        "WHERE request_id = ? ORDER BY leg, step_number", (req['id'],)
    )
    step_rows = await cursor.fetchall()
    steps = []
    for s in step_rows:
        steps.append({
            'id': s[0],
            'step_number': s[1], 'step_name': s[2], 'leg': s[3],
            'status': s[4],
            'started_at': _safe_isoparse(s[5]),
            'completed_at': _safe_isoparse(s[6]),
            'machine_no': s[7], 'rack_no': s[8], 'operator_id': s[9], 'tray_no': s[10],
            'qty_in': s[11], 'qty_out': s[12],
            'notes': s[13],
            'attachments': json.loads(s[14]) if s[14] else [],
            'custom_fields': json.loads(s[15]) if s[15] else {},
        })
    req['steps'] = steps
    return req


# ========================
# Auth Routes
# ========================
@api_router.post("/auth/register", response_model=User)
async def register(user_create: UserCreate):
    db = await get_db()
    try:
        cursor = await db.execute("SELECT id FROM users WHERE email = ?", (user_create.email,))
        if await cursor.fetchone():
            raise HTTPException(status_code=400, detail="Email already registered")

        hashed_password = hash_password(user_create.password)
        user = User(email=user_create.email, username=user_create.username, role=user_create.role, approved=False)
        await db.execute(
            "INSERT INTO users (id, email, username, password, role, approved, created_at) VALUES (?,?,?,?,?,?,?)",
            (user.id, user.email, user.username, hashed_password, user.role.value, 0, user.created_at.isoformat())
        )
        await db.commit()
        return user
    finally:
        await db.close()

@api_router.post("/auth/login", response_model=Token)
async def login(user_login: UserLogin):
    db = await get_db()
    try:
        cursor = await db.execute(
            "SELECT id, email, username, password, role, created_at, approved, blocked FROM users WHERE email = ?",
            (user_login.email,)
        )
        row = await cursor.fetchone()
        if not row:
            raise HTTPException(status_code=401, detail="Invalid email or password")
        if not verify_password(user_login.password, row[3]):
            raise HTTPException(status_code=401, detail="Invalid email or password")

        # Check if user is blocked
        blocked = bool(row[7]) if row[7] is not None else False
        if blocked:
            raise HTTPException(status_code=403, detail="Your account has been blocked. Please contact an administrator.")

        # Check if user is approved
        approved = bool(row[6]) if row[6] is not None else False
        if not approved:
            raise HTTPException(status_code=403, detail="Your account is pending admin approval. Please contact your administrator.")

        access_token = create_access_token(data={"sub": row[0]})
        # Fetch extended profile (position, avatar, etc.)
        cursor2 = await db.execute(
            "SELECT position, contact_email, plant, manager, avatar FROM users WHERE id = ?", (row[0],)
        )
        profile_row = await cursor2.fetchone()
        user = User(id=row[0], email=row[1], username=row[2], role=row[4],
                    approved=approved,
                    created_at=datetime.fromisoformat(row[5]),
                    position=profile_row[0] or '' if profile_row else '',
                    contact_email=profile_row[1] or '' if profile_row else '',
                    plant=profile_row[2] or '' if profile_row else '',
                    manager=profile_row[3] or '' if profile_row else '',
                    avatar=profile_row[4] or None if profile_row else None)

        # Log login
        now = datetime.now(timezone.utc).isoformat()
        try:
            await db.execute(
                "INSERT INTO login_logs (user_id, email, username, role, login_at, ip_address) VALUES (?,?,?,?,?,?)",
                (row[0], row[1], row[2], row[4], now, "local")
            )
            await db.commit()
        except Exception as e:
            logging.warning(f"Failed to log login: {e}")

        return Token(access_token=access_token, token_type="bearer", user=user)
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Login error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Internal server error: {str(e)}")
    finally:
        await db.close()

@api_router.post("/auth/guest-token")
async def guest_token(body: dict = None):
    """Issue a short-lived JWT for a guest Technician - no credentials required."""
    body = body or {}
    emp_id   = (body.get("employee_id") or "").strip()
    emp_name = (body.get("employee_name") or "").strip()
    emp_pos  = (body.get("employee_position") or "").strip()
    display  = emp_name or "Technician"
    token = create_access_token({"sub": "guest", "is_guest": True,
                                  "employee_id": emp_id, "employee_name": emp_name})
    now = datetime.now(timezone.utc).isoformat()
    db = await get_db()
    try:
        await db.execute(
            "INSERT INTO login_logs (user_id, email, username, role, login_at, ip_address, employee_id, employee_name) "
            "VALUES (?,?,?,?,?,?,?,?)",
            ("guest", "guest@technician.local", display, "Technician", now, "local", emp_id, emp_name)
        )
        if emp_id:
            await db.execute("""
                INSERT INTO technician_sessions (employee_id, employee_name, employee_position, last_active, login_at)
                VALUES (?, ?, ?, ?, ?)
                ON CONFLICT(employee_id) DO UPDATE SET
                    employee_name=excluded.employee_name,
                    employee_position=excluded.employee_position,
                    last_active=excluded.last_active,
                    login_at=excluded.login_at
            """, (emp_id, emp_name, emp_pos, now, now))
        await db.commit()
    finally:
        await db.close()
    return {
        "access_token": token,
        "token_type": "bearer",
        "user": {
            "id": "guest",
            "email": "guest@technician.local",
            "username": display,
            "role": "Technician",
            "approved": True,
            "is_guest": True,
            "employee_id": emp_id,
            "employee_name": emp_name,
        }
    }

@api_router.get("/auth/me")
async def get_me(current_user: User = Depends(get_current_user)):
    return {
        "id": current_user.id,
        "email": current_user.email,
        "username": current_user.username,
        "role": current_user.role,
        "approved": current_user.approved,
        "is_guest": current_user.is_guest,
        "position": current_user.position or "",
        "contact_email": current_user.contact_email or "",
        "plant": current_user.plant or "",
        "manager": current_user.manager or "",
        "avatar": current_user.avatar or None,
    }

@api_router.patch("/auth/profile", response_model=User)
async def update_profile(data: ProfileUpdate, current_user: User = Depends(get_current_user)):
    fields, values = [], []
    if data.username is not None and data.username.strip():
        fields.append("username = ?")
        values.append(data.username.strip())
    if data.position is not None:
        fields.append("position = ?")
        values.append(data.position.strip())
    if data.contact_email is not None:
        fields.append("contact_email = ?")
        values.append(data.contact_email.strip())
    if data.plant is not None:
        fields.append("plant = ?")
        values.append(data.plant.strip())
    if data.manager is not None:
        fields.append("manager = ?")
        values.append(data.manager.strip())
    if data.avatar is not None:
        fields.append("avatar = ?")
        values.append(data.avatar if data.avatar else None)
    if not fields:
        return current_user
    values.append(current_user.id)
    db = await get_db()
    try:
        await db.execute(f"UPDATE users SET {', '.join(fields)} WHERE id = ?", values)
        await db.commit()
        cursor = await db.execute(
            "SELECT id, email, username, role, created_at, approved, position, contact_email, plant, manager, avatar FROM users WHERE id = ?",
            (current_user.id,)
        )
        row = await cursor.fetchone()
        return User(id=row[0], email=row[1], username=row[2], role=row[3],
                    created_at=datetime.fromisoformat(row[4]),
                    approved=bool(row[5]) if row[5] is not None else False,
                    position=row[6] or '', contact_email=row[7] or '',
                    plant=row[8] or '', manager=row[9] or '',
                    avatar=row[10] or None)
    finally:
        await db.close()


# ========================
# User Management Routes
# ========================
@api_router.get("/users", response_model=List[User])
async def get_users(current_user: User = Depends(get_current_user)):
    # Guest users (unauthenticated Technician token) should not see the user directory
    if current_user.is_guest:
        raise HTTPException(status_code=403, detail="Guest access not permitted")
    db = await get_db()
    try:
        cursor = await db.execute("SELECT id, email, username, role, created_at, approved, position, contact_email, plant, manager, last_seen, blocked, user_status, avatar FROM users")
        rows = await cursor.fetchall()
        result = []
        for r in rows:
            _approved = bool(r[5]) if r[5] is not None else False
            _blocked = bool(r[11]) if r[11] is not None else False
            _user_status = r[12] or ('approved' if _approved else ('lock' if _blocked else 'pending'))
            result.append(User(id=r[0], email=r[1], username=r[2], role=r[3],
                     created_at=datetime.fromisoformat(r[4]),
                     approved=_approved,
                     position=r[6] or '', contact_email=r[7] or '',
                     plant=r[8] or '', manager=r[9] or '',
                     last_seen=datetime.fromisoformat(r[10]) if r[10] else None,
                     blocked=_blocked, user_status=_user_status,
                     avatar=r[13] or None))
        return result
    finally:
        await db.close()

@api_router.post("/auth/heartbeat")
async def heartbeat(current_user: User = Depends(get_current_user)):
    """Update the user's last_seen timestamp so others can see they're online."""
    now = datetime.now(timezone.utc).isoformat()
    if current_user.is_guest:
        emp_id = current_user.employee_id or ""
        if emp_id:
            db = await get_db()
            try:
                await db.execute(
                    "UPDATE technician_sessions SET last_active=? WHERE employee_id=?",
                    (now, emp_id)
                )
                await db.commit()
            finally:
                await db.close()
        return {"ok": True}
    db = await get_db()
    try:
        await db.execute("UPDATE users SET last_seen = ? WHERE id = ?", (now, current_user.id))
        await db.commit()
    finally:
        await db.close()
    return {"ok": True}

@api_router.get("/auth/online-users")
async def get_online_users(current_user: User = Depends(get_current_user)):
    """Return list of users who have been seen in the last 2 minutes."""
    two_min_ago = (datetime.now(timezone.utc) - timedelta(minutes=2)).isoformat()
    db = await get_db()
    try:
        cursor = await db.execute(
            "SELECT id, username, avatar, role FROM users WHERE last_seen > ? AND approved = 1 AND blocked = 0",
            (two_min_ago,)
        )
        rows = await cursor.fetchall()
        return {"users": [{"id": r[0], "username": r[1], "avatar": r[2] or None, "role": r[3]} for r in rows]}
    finally:
        await db.close()

@api_router.delete("/users/{user_id}")
async def delete_user(user_id: str, current_user: User = Depends(require_permission('manage_users'))):
    db = await get_db()
    try:
        cursor = await db.execute("SELECT approved, blocked, user_status FROM users WHERE id = ?", (user_id,))
        user_row = await cursor.fetchone()
        if not user_row:
            raise HTTPException(status_code=404, detail="User not found")
        _approved = bool(user_row[0]) if user_row[0] is not None else False
        _status = user_row[2] or ('approved' if _approved else 'pending')
        # Can only delete users who are pending, hold, lock, or declined â€” not fully approved ones
        if _status == 'approved':
            raise HTTPException(
                status_code=403,
                detail="Cannot delete an approved user. Use Lock or Declined status instead."
            )
        await db.execute("DELETE FROM users WHERE id = ?", (user_id,))
        await db.commit()
        return {"message": "User deleted successfully"}
    finally:
        await db.close()


class UpdateUserStatusRequest(BaseModel):
    status: str  # pending, approved, hold, lock, declined

@api_router.patch("/users/{user_id}/status")
async def update_user_status(user_id: str, request: UpdateUserStatusRequest, current_user: User = Depends(require_role([UserRole.ADMIN]))):
    """Update user account status (pending/approved/hold/lock/declined). Admin only."""
    valid_statuses = ['pending', 'approved', 'hold', 'lock', 'declined']
    if request.status not in valid_statuses:
        raise HTTPException(status_code=400, detail=f"Invalid status. Must be one of: {', '.join(valid_statuses)}")
    if user_id == current_user.id:
        raise HTTPException(status_code=400, detail="You cannot change your own status")
    db = await get_db()
    try:
        approved = 1 if request.status == 'approved' else 0
        blocked = 1 if request.status == 'lock' else 0
        declined_at = "datetime('now')" if request.status == 'declined' else None
        if declined_at:
            cursor = await db.execute(
                "UPDATE users SET approved = ?, blocked = ?, user_status = ?, declined_at = datetime('now') WHERE id = ?",
                (approved, blocked, request.status, user_id)
            )
        else:
            cursor = await db.execute(
                "UPDATE users SET approved = ?, blocked = ?, user_status = ?, declined_at = NULL WHERE id = ?",
                (approved, blocked, request.status, user_id)
            )
        await db.commit()
        if cursor.rowcount == 0:
            raise HTTPException(status_code=404, detail="User not found")
        return {"message": f"User status updated to {request.status}"}
    finally:
        await db.close()

@api_router.patch("/users/{user_id}/approve")
async def approve_user(user_id: str, current_user: User = Depends(require_role([UserRole.ADMIN]))):
    db = await get_db()
    try:
        cursor = await db.execute("UPDATE users SET approved = 1 WHERE id = ?", (user_id,))
        await db.commit()
        if cursor.rowcount == 0:
            raise HTTPException(status_code=404, detail="User not found")
        return {"message": "User approved successfully"}
    finally:
        await db.close()

@api_router.patch("/users/{user_id}/reject")
async def reject_user(user_id: str, current_user: User = Depends(require_role([UserRole.ADMIN]))):
    db = await get_db()
    try:
        cursor = await db.execute("UPDATE users SET approved = 0 WHERE id = ?", (user_id,))
        await db.commit()
        if cursor.rowcount == 0:
            raise HTTPException(status_code=404, detail="User not found")
        return {"message": "User approval revoked"}
    finally:
        await db.close()

class UpdateUserRoleRequest(BaseModel):
    role: str

@api_router.patch("/users/{user_id}/role")
async def update_user_role(user_id: str, request: UpdateUserRoleRequest, current_user: User = Depends(require_role([UserRole.ADMIN]))):
    """Update the role of an existing user. Only admins can perform this action."""
    # Validate that the role is one of the valid roles
    valid_roles = [r.value for r in UserRole]
    if request.role not in valid_roles:
        raise HTTPException(status_code=400, detail=f"Invalid role. Must be one of: {', '.join(valid_roles)}")
    
    db = await get_db()
    try:
        cursor = await db.execute("UPDATE users SET role = ? WHERE id = ?", (request.role, user_id))
        await db.commit()
        if cursor.rowcount == 0:
            raise HTTPException(status_code=404, detail="User not found")
        return {"message": "User role updated successfully"}
    finally:
        await db.close()


class UpdateUsernameRequest(BaseModel):
    username: str

@api_router.patch("/users/{user_id}/username")
async def update_user_username(user_id: str, request: UpdateUsernameRequest, current_user: User = Depends(require_role([UserRole.ADMIN]))):
    """Update the display name of a user. Admin only."""
    if not request.username.strip():
        raise HTTPException(status_code=400, detail="Username cannot be empty")
    db = await get_db()
    try:
        cursor = await db.execute("UPDATE users SET username = ? WHERE id = ?", (request.username.strip(), user_id))
        await db.commit()
        if cursor.rowcount == 0:
            raise HTTPException(status_code=404, detail="User not found")
        return {"message": "Username updated", "username": request.username.strip()}
    finally:
        await db.close()

@api_router.patch("/users/{user_id}/block")
async def toggle_block_user(user_id: str, current_user: User = Depends(require_role([UserRole.ADMIN]))):
    """Block or unblock a user. Admin only."""
    if user_id == current_user.id:
        raise HTTPException(status_code=400, detail="You cannot block your own account")
    db = await get_db()
    try:
        cursor = await db.execute("SELECT blocked FROM users WHERE id = ?", (user_id,))
        row = await cursor.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="User not found")
        new_blocked = 0 if (row[0] or 0) else 1
        await db.execute("UPDATE users SET blocked = ? WHERE id = ?", (new_blocked, user_id))
        await db.commit()
        return {"blocked": bool(new_blocked), "message": "User blocked" if new_blocked else "User unblocked"}
    finally:
        await db.close()


class ForgotPasswordRequest(BaseModel):
    email: EmailStr
    new_password: str

@api_router.post("/auth/forgot-password")
async def forgot_password(body: ForgotPasswordRequest):
    """Reset password for the given email. Math verification is handled client-side."""
    if len(body.new_password) < 6:
        raise HTTPException(status_code=400, detail="New password must be at least 6 characters")
    db = await get_db()
    try:
        cursor = await db.execute("SELECT id FROM users WHERE email = ?", (body.email,))
        row = await cursor.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="No account found with that email address")
        new_hashed = hash_password(body.new_password)
        await db.execute("UPDATE users SET password = ? WHERE id = ?", (new_hashed, row[0]))
        await db.commit()
        return {"message": "Password reset successfully. You can now log in with your new password."}
    finally:
        await db.close()


class ChangePasswordRequest(BaseModel):
    current_password: str
    new_password: str

@api_router.post("/auth/change-password")
async def change_password(body: ChangePasswordRequest, current_user: User = Depends(get_current_user)):
    """Change password for the currently authenticated user."""
    if current_user.is_guest:
        raise HTTPException(status_code=403, detail="Guest accounts cannot change password")
    if len(body.new_password) < 6:
        raise HTTPException(status_code=400, detail="New password must be at least 6 characters")
    db = await get_db()
    try:
        cursor = await db.execute("SELECT password FROM users WHERE id = ?", (current_user.id,))
        row = await cursor.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="User not found")
        if not verify_password(body.current_password, row[0]):
            raise HTTPException(status_code=400, detail="Current password is incorrect")
        new_hashed = hash_password(body.new_password)
        await db.execute("UPDATE users SET password = ? WHERE id = ?", (new_hashed, current_user.id))
        await db.commit()
        return {"message": "Password changed successfully"}
    finally:
        await db.close()


@api_router.get("/login-logs")
async def get_login_logs(current_user: User = Depends(require_role([UserRole.ADMIN]))):
    db = await get_db()
    try:
        cursor = await db.execute(
            "SELECT id, user_id, email, username, role, login_at, ip_address, "
            "COALESCE(employee_id,'') AS employee_id, COALESCE(employee_name,'') AS employee_name "
            "FROM login_logs ORDER BY login_at DESC LIMIT 200"
        )
        rows = await cursor.fetchall()
        return [{"id": r[0], "user_id": r[1], "email": r[2], "username": r[3],
                 "role": r[4], "login_at": r[5], "ip_address": r[6],
                 "employee_id": r[7], "employee_name": r[8]} for r in rows]
    finally:
        await db.close()

@api_router.get("/active-technicians")
async def get_active_technicians(current_user: User = Depends(get_current_user)):
    """Return technician sessions active in the last 5 minutes."""
    five_min_ago = (datetime.now(timezone.utc) - timedelta(minutes=5)).isoformat()
    db = await get_db()
    try:
        cursor = await db.execute(
            "SELECT employee_id, employee_name, employee_position, last_active, login_at "
            "FROM technician_sessions WHERE last_active >= ? ORDER BY last_active DESC",
            (five_min_ago,)
        )
        rows = await cursor.fetchall()
        return [{"employee_id": r[0], "employee_name": r[1], "employee_position": r[2],
                 "last_active": r[3], "login_at": r[4]} for r in rows]
    finally:
        await db.close()


# ========================
# Task Manager Stats
# ========================
@api_router.get("/task-manager/stats")
async def get_task_manager_stats(current_user: User = Depends(require_role([UserRole.ADMIN]))):
    """Website Task Manager â€” performance metrics, online users, recent activity."""
    db = await get_db()
    try:
        now = datetime.now(timezone.utc)
        five_min_ago = (now - timedelta(minutes=5)).isoformat()

        # Users online
        cursor = await db.execute("SELECT COUNT(*) FROM users WHERE last_seen > ?", (five_min_ago,))
        users_online = (await cursor.fetchone())[0]

        # Total users
        cursor = await db.execute("SELECT COUNT(*) FROM users")
        total_users = (await cursor.fetchone())[0]

        # Recent requests (last 10 added)
        cursor = await db.execute(
            "SELECT request_number, device_name, customer, created_by_username, created_at, status "
            "FROM requests ORDER BY created_at DESC LIMIT 10"
        )
        recent_requests = [
            {"request_number": r[0], "device_name": r[1] or '', "customer": r[2] or '',
             "created_by": r[3], "created_at": r[4], "status": r[5]}
            for r in await cursor.fetchall()
        ]

        # Recently edited process steps (top 10 most recently updated requests)
        cursor = await db.execute(
            """SELECT DISTINCT r.request_number, ps.step_name, ps.status, r.updated_at,
                      COALESCE(ps.updated_by, r.created_by_username), r.device_name
               FROM process_steps ps
               JOIN requests r ON ps.request_id = r.id
               WHERE ps.status IN ('in_progress', 'completed', 'hold')
               ORDER BY r.updated_at DESC LIMIT 10"""
        )
        recent_step_edits = [
            {"request_number": r[0], "step_name": r[1], "status": r[2],
             "updated_at": r[3], "by": r[4], "device_name": r[5] or ''}
            for r in await cursor.fetchall()
        ]

        # Request status counts
        cursor = await db.execute("SELECT status, COUNT(*) FROM requests GROUP BY status")
        status_counts = {r[0]: r[1] for r in await cursor.fetchall()}

        # Total requests
        cursor = await db.execute("SELECT COUNT(*) FROM requests")
        total_requests = (await cursor.fetchone())[0]

        # Recent login history (last 5 logins)
        cursor = await db.execute(
            "SELECT username, role, login_at FROM login_logs ORDER BY login_at DESC LIMIT 5"
        )
        recent_logins = [{"username": r[0], "role": r[1], "login_at": r[2]} for r in await cursor.fetchall()]

        return {
            "users_online": users_online,
            "total_users": total_users,
            "total_requests": total_requests,
            "recent_requests": recent_requests,
            "recent_step_edits": recent_step_edits,
            "status_counts": status_counts,
            "recent_logins": recent_logins,
            "server_time": now.isoformat(),
        }
    finally:
        await db.close()

@api_router.get("/system/health")
async def get_system_health(
    period: str = "24H",
    current_user: User = Depends(require_role([UserRole.ADMIN]))
):
    """System health metrics for the Settings page health dashboard."""
    now = datetime.now(timezone.utc)
    uptime_seconds = int((now - _SERVER_START).total_seconds())

    # CPU / Memory / Load
    if _PSUTIL_OK:
        cpu = round(_psutil.cpu_percent(interval=0.1), 1)
        mem = round(_psutil.virtual_memory().percent, 1)
        try:
            load_1, _, _ = _psutil.getloadavg()
            load = round(load_1, 2)
        except AttributeError:
            # Windows has no getloadavg; approximate from CPU
            load = round(_psutil.cpu_percent(interval=0) / 25.0, 2)
    else:
        cpu, mem, load = 0.0, 0.0, 0.0

    db = await get_db()
    try:
        # Active tests
        cursor = await db.execute(
            "SELECT COUNT(*) FROM requests WHERE status IN ('testing','in_progress')"
        )
        active_tests = (await cursor.fetchone())[0]

        # Total & completed counts (all time)
        cursor = await db.execute("SELECT COUNT(*) FROM requests")
        active_db = (await cursor.fetchone())[0]
        cursor = await db.execute(
            "SELECT COUNT(*) FROM requests WHERE status = 'completed'"
        )
        # note: completed are removed from DB after backup; approximate ratio
        completed_approx = active_db  # placeholder total denominator

        # Build time-bucketed chart data
        if period == "7D":
            num_buckets, bucket_hours = 7, 24
            def lbl(i): return (now - timedelta(days=(num_buckets - 1 - i))).strftime("%a")
        elif period == "30D":
            num_buckets, bucket_hours = 15, 48
            def lbl(i): return (now - timedelta(days=int((num_buckets - 1 - i) * 2))).strftime("%b %d")
        else:  # 24H
            num_buckets, bucket_hours = 12, 2
            def lbl(i): return f"{(i * bucket_hours):02d}:00"

        chart_data = []
        for i in range(num_buckets):
            b_end = now - timedelta(hours=(num_buckets - 1 - i) * bucket_hours)
            b_start = b_end - timedelta(hours=bucket_hours)
            b_s, b_e = b_start.isoformat(), b_end.isoformat()

            cursor = await db.execute(
                "SELECT COUNT(*) FROM requests WHERE status='completed' AND updated_at BETWEEN ? AND ?",
                (b_s, b_e)
            )
            success = (await cursor.fetchone())[0]

            cursor = await db.execute(
                "SELECT COUNT(*) FROM requests "
                "WHERE status IN ('testing','in_progress','review','approval') AND updated_at BETWEEN ? AND ?",
                (b_s, b_e)
            )
            running = (await cursor.fetchone())[0]

            cursor = await db.execute(
                "SELECT COUNT(*) FROM requests WHERE updated_at BETWEEN ? AND ? "
                "AND id IN (SELECT DISTINCT request_id FROM process_steps WHERE status='hold')",
                (b_s, b_e)
            )
            failed = (await cursor.fetchone())[0]

            chart_data.append({"time": lbl(i), "success": success, "running": running, "failed": failed})

        # Overall success rate against all requests touched today
        since_24h = (now - timedelta(hours=24)).isoformat()
        cursor = await db.execute(
            "SELECT COUNT(*) FROM requests WHERE updated_at > ?", (since_24h,)
        )
        touched_24h = (await cursor.fetchone())[0]
        cursor = await db.execute(
            "SELECT COUNT(*) FROM requests WHERE status='completed' AND updated_at > ?", (since_24h,)
        )
        done_24h = (await cursor.fetchone())[0]
        success_rate = round((done_24h / touched_24h * 100), 1) if touched_24h > 0 else 0.0

        # Throughput
        throughput_per_hr = done_24h / 24.0
        if throughput_per_hr >= 1000:
            throughput_str = f"{throughput_per_hr / 1000:.1f}K/h"
        elif throughput_per_hr >= 1:
            throughput_str = f"{throughput_per_hr:.1f}/h"
        else:
            throughput_str = f"{done_24h}/24h"

        # Avg simulated response (heavier load â†’ slightly slower)
        avg_response_ms = max(45, min(600, 65 + active_tests * 4))

        overall_status = "DEGRADED" if (cpu > 90 or mem > 95) else "OPERATIONAL"

        return {
            "status": overall_status,
            "active_tests": active_tests,
            "uptime_pct": 99.97,
            "uptime_seconds": uptime_seconds,
            "cpu_pct": cpu,
            "mem_pct": mem,
            "load_avg": load,
            "avg_response_ms": avg_response_ms,
            "success_rate": success_rate,
            "throughput": throughput_str,
            "chart_data": chart_data,
            "server_time": now.isoformat(),
        }
    finally:
        await db.close()


# ========================
# Role Permissions Routes
# ========================
@api_router.get("/role-permissions")
async def get_role_permissions(current_user: User = Depends(get_current_user)):
    """Get permissions for all roles (Admin only), or the current user's own permissions."""
    db = await get_db()
    try:
        if current_user.role == UserRole.ADMIN:
            # Admin sees all role permissions
            cursor = await db.execute("SELECT role, permission, granted FROM role_permissions")
            rows = await cursor.fetchall()
            result = {}
            # Initialize all roles with all permissions = false
            configurable_roles = ['Reliability Engineer', 'Failure Analysis', 'Technician', 'Planner']
            for role in configurable_roles:
                result[role] = {p: False for p in ALL_PERMISSIONS}
            for r in rows:
                role_name = r[0]
                if role_name in result:
                    result[role_name][r[1]] = bool(r[2])
            return {"permissions": result, "all_permissions": ALL_PERMISSIONS}
        else:
            # Non-admin: return own permissions
            cursor = await db.execute(
                "SELECT permission, granted FROM role_permissions WHERE role = ?",
                (current_user.role,)
            )
            rows = await cursor.fetchall()
            perms = {p: False for p in ALL_PERMISSIONS}
            for r in rows:
                perms[r[0]] = bool(r[1])
            # Admin always has all permissions
            return {"permissions": perms, "role": current_user.role}
    finally:
        await db.close()

class RolePermissionsUpdate(BaseModel):
    permissions: Dict[str, Dict[str, bool]]  # { role: { permission: true/false } }

@api_router.put("/role-permissions")
async def update_role_permissions(
    update: RolePermissionsUpdate,
    current_user: User = Depends(require_role([UserRole.ADMIN]))
):
    """Update permissions for roles. Admin only."""
    db = await get_db()
    try:
        configurable_roles = ['Reliability Engineer', 'Failure Analysis', 'Technician', 'Planner']
        for role, perms in update.permissions.items():
            if role not in configurable_roles:
                continue
            for perm, granted in perms.items():
                if perm not in ALL_PERMISSIONS:
                    continue
                await db.execute(
                    "INSERT INTO role_permissions (role, permission, granted) VALUES (?, ?, ?) "
                    "ON CONFLICT(role, permission) DO UPDATE SET granted = ?",
                    (role, perm, int(granted), int(granted))
                )
        await db.commit()
        return {"message": "Permissions updated successfully"}
    finally:
        await db.close()

# ========================
# Maintenance Mode
# ========================

class MaintenanceToggle(BaseModel):
    active: bool
    message: Optional[str] = "System is currently under maintenance. Please check back later."

@api_router.get("/maintenance")
async def get_maintenance_status():
    """Public endpoint: returns current maintenance mode status."""
    if MAINTENANCE_FLAG.exists():
        try:
            data = json.loads(MAINTENANCE_FLAG.read_text())
            return {"active": True, "message": data.get("message", ""), "started_at": data.get("started_at", "")}
        except Exception:
            return {"active": True, "message": "", "started_at": ""}
    return {"active": False, "message": "", "started_at": ""}

@api_router.post("/maintenance")
async def set_maintenance_mode(
    body: MaintenanceToggle,
    current_user: User = Depends(require_role([UserRole.ADMIN]))
):
    """Toggle maintenance mode on/off. Admin only."""
    if body.active:
        MAINTENANCE_FLAG.write_text(json.dumps({
            "message": body.message or "System is currently under maintenance. Please check back later.",
            "started_at": datetime.now(timezone.utc).isoformat(),
            "by": current_user.username,
        }))
        logging.info(f"Maintenance mode ENABLED by {current_user.username}")
        return {"active": True, "message": body.message}
    else:
        if MAINTENANCE_FLAG.exists():
            MAINTENANCE_FLAG.unlink()
        logging.info(f"Maintenance mode DISABLED by {current_user.username}")
        return {"active": False}

# ========================
# Server Controls (Admin)
# ========================

@api_router.post("/admin/restart-backend")
async def restart_backend(
    current_user: User = Depends(require_role([UserRole.ADMIN]))
):
    """Restart the backend process. Admin only."""
    logging.info(f"Backend restart requested by {current_user.username}")
    def _do_restart():
        os.execv(sys.executable, [sys.executable] + sys.argv)
    threading.Timer(0.8, _do_restart).start()
    return {"message": "Backend is restartingâ€¦"}

@api_router.get("/my-permissions")
async def get_my_permissions(current_user: User = Depends(get_current_user)):
    """Get current user's computed permissions."""
    if current_user.role == UserRole.ADMIN:
        return {"permissions": {p: True for p in ALL_PERMISSIONS}}
    db = await get_db()
    try:
        cursor = await db.execute(
            "SELECT permission, granted FROM role_permissions WHERE role = ?",
            (current_user.role,)
        )
        rows = await cursor.fetchall()
        perms = {p: False for p in ALL_PERMISSIONS}
        for r in rows:
            perms[r[0]] = bool(r[1])
        return {"permissions": perms}
    finally:
        await db.close()


# ========================
# Request Routes
# ========================
@api_router.get("/requests/next-number")
async def get_next_request_number(
    request_type: Optional[str] = 'REL',
    current_user: User = Depends(get_current_user)
):
    """Return the next request number for the given request type (REL or RMS)."""
    request_type = (request_type or 'REL').strip().upper()
    if request_type not in ('REL', 'RMS'):
        raise HTTPException(status_code=400, detail="Invalid request_type: must be 'REL' or 'RMS'.")
    db = await get_db()
    try:
        year = datetime.now(timezone.utc).year
        prefix = request_type
        cursor = await db.execute(
            "SELECT MAX(CAST(SUBSTR(request_number, ?) AS INTEGER)) FROM requests WHERE request_number LIKE ?",
            (len(f"{prefix}{year}") + 1, f"{prefix}{year}%")
        )
        max_row = await cursor.fetchone()
        next_number = f"{prefix}{year}{(max_row[0] or 0) + 1:05d}"
        return {"next_number": next_number}
    finally:
        await db.close()

@api_router.post("/requests", response_model=RelRequest)
async def create_request(
    request_create: RequestCreate,
    current_user: User = Depends(require_permission('create_request'))
):
    db = await get_db()
    try:
        request_type = (request_create.request_type or 'REL').strip().upper()
        if request_type not in ('REL', 'RMS'):
            raise HTTPException(status_code=400, detail="Invalid request_type: must be 'REL' or 'RMS'.")

        request_number = request_create.request_number
        if not request_number:
            year = datetime.now(timezone.utc).year
            prefix = request_type
            # Lock row set to avoid parallel allocation collisions (SQLite immediate transaction)
            await db.execute("BEGIN IMMEDIATE")
            cursor = await db.execute(
                "SELECT MAX(CAST(SUBSTR(request_number, ?) AS INTEGER)) FROM requests WHERE request_number LIKE ?",
                (len(f"{prefix}{year}") + 1, f"{prefix}{year}%")
            )
            max_row = await cursor.fetchone()
            request_number = f"{prefix}{year}{(max_row[0] or 0) + 1:05d}"

        # Build steps from custom_steps or use defaults
        if request_create.custom_steps and len(request_create.custom_steps) > 0:
            duplicate_steps = [name for name in request_create.custom_steps if request_create.custom_steps.count(name) > 1]
            if duplicate_steps:
                raise HTTPException(status_code=400, detail="Duplicate steps are not allowed when defining custom steps.")
            invalid_steps = [name for name in request_create.custom_steps if name not in AVAILABLE_STEP_NAMES]
            if invalid_steps:
                raise HTTPException(status_code=400, detail=f"Invalid step names: {', '.join(invalid_steps)}")
            step_defs = [
                {"step_number": i + 1, "step_name": name, "status": StepStatus.PENDING}
                for i, name in enumerate(request_create.custom_steps)
            ]
        else:
            step_defs = DEFAULT_STEPS

        steps = [ProcessStep(**step) for step in step_defs]
        request_data = request_create.model_dump(exclude_none=True)
        request_data.pop('custom_steps', None)  # Remove custom_steps from request data
        request_data['request_number'] = request_number
        request_data['request_type'] = request_type
        request_data['status'] = request_data.get('status') or 'incoming'

        request_obj = RelRequest(
            **request_data, created_by=current_user.id,
            created_by_username=current_user.username, steps=steps
        )

        await db.execute(
            """INSERT INTO requests (id, request_number, request_type, classification, originator, plant,
               device_name, lot_no, customer, pkg_info, automotive, date_ltc,
               product_hierarchy, pdl, body_size_x, body_size_y, package_thickness,
               ball_pitch, ball_count, lead_pitch, lead_count, total_ss, purpose,
               engineer_special_instruction, deadline, created_by, created_by_username,
               created_at, updated_at, status, current_step)
               VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
            (request_obj.id, request_obj.request_number, request_obj.request_type,
             request_obj.classification or '', request_obj.originator or '',
             request_obj.plant or '', request_obj.device_name or '',
             request_obj.lot_no or '', request_obj.customer or '',
             request_obj.pkg_info or '', int(request_obj.automotive),
             request_obj.date_ltc, request_obj.product_hierarchy,
             request_obj.pdl, request_obj.body_size_x, request_obj.body_size_y,
             request_obj.package_thickness, request_obj.ball_pitch, request_obj.ball_count,
             request_obj.lead_pitch, request_obj.lead_count, request_obj.total_ss,
             request_obj.purpose or '', request_obj.engineer_special_instruction,
             request_obj.deadline, request_obj.created_by, request_obj.created_by_username,
             request_obj.created_at.isoformat(), request_obj.updated_at.isoformat(),
             request_obj.status, request_obj.current_step)
        )

        # Default test items and conditions for step names
        DEFAULT_STEP_ITEMS = {
            'Electrical Test': 'E-Test',
            'T & H Soak': None,  # handled by frontend
            # Add more mappings as needed
        }
        DEFAULT_STEP_CONDITIONS = {
            'Electrical Test': 'P4',
            'T & H Soak': None,  # handled by frontend
            # Add more mappings as needed
        }
        for step in steps:
            # Set default custom_fields if not present
            cf = dict(step.custom_fields) if step.custom_fields else {}
            if 'test_item' not in cf and step.step_name in DEFAULT_STEP_ITEMS and DEFAULT_STEP_ITEMS[step.step_name]:
                cf['test_item'] = DEFAULT_STEP_ITEMS[step.step_name]
            if 'test_condition' not in cf and step.step_name in DEFAULT_STEP_CONDITIONS and DEFAULT_STEP_CONDITIONS[step.step_name]:
                cf['test_condition'] = DEFAULT_STEP_CONDITIONS[step.step_name]
            await db.execute(
                """INSERT INTO process_steps (request_id, leg, step_number, step_name, status,
                   started_at, completed_at, machine_no, rack_no, operator_id, tray_no, notes, attachments, custom_fields)
                   VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                (request_obj.id, 1, step.step_number, step.step_name, step.status.value,
                 step.started_at.isoformat() if step.started_at else None,
                 step.completed_at.isoformat() if step.completed_at else None,
                 step.machine_no, step.rack_no, step.operator_id, step.tray_no, step.notes,
                 json.dumps(step.attachments or []), json.dumps(cf))
            )

        await db.commit()
        return request_obj
    except aiosqlite.IntegrityError as e:
        if "UNIQUE constraint failed" in str(e):
            raise HTTPException(status_code=409, detail=f"A request with number '{request_number}' already exists")
        raise HTTPException(status_code=400, detail=str(e))
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Unexpected error in create_request: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        await db.close()

@api_router.get("/requests", response_model=List[RelRequest])
async def get_requests(
    search: Optional[str] = None, status: Optional[str] = None,
    date_from: Optional[str] = None, date_to: Optional[str] = None,
    created_by: Optional[str] = None,
    current_user: User = Depends(get_current_user)
):
    db = await get_db()
    try:
        conditions = []
        params: list = []

        if search:
            search_fields = [
                'request_number', 'device_name', 'customer', 'lot_no',
                'classification', 'originator', 'plant', 'pkg_info',
                'product_hierarchy', 'pdl', 'total_ss', 'purpose',
                'engineer_special_instruction', 'status',
            ]
            or_clauses = ' OR '.join(f"{f} LIKE ?" for f in search_fields)
            conditions.append(f"({or_clauses})")
            like = f"%{search}%"
            params.extend([like] * len(search_fields))

        if date_from:
            conditions.append("DATE(created_at) >= DATE(?)")
            params.append(date_from)

        if date_to:
            conditions.append("DATE(created_at) <= DATE(?)")
            params.append(date_to)

        if created_by:
            conditions.append("created_by_username LIKE ?")
            params.append(f"%{created_by}%")

        if status:
            if status == "delayed":
                today = datetime.now(timezone.utc).isoformat()
                conditions.append("deadline IS NOT NULL AND deadline != '' AND status NOT IN ('completed','discontinued') AND deadline < ?")
                params.append(today)
            elif status == "upcoming":
                today = datetime.now(timezone.utc).isoformat()
                three_days = (datetime.now(timezone.utc) + timedelta(days=3)).isoformat()
                conditions.append("deadline IS NOT NULL AND deadline != '' AND status NOT IN ('completed','discontinued') AND deadline >= ? AND deadline <= ?")
                params.extend([today, three_days])
            else:
                conditions.append("status = ?")
                params.append(status)

        where = ("WHERE " + " AND ".join(conditions)) if conditions else ""
        cursor = await db.execute(
            f"SELECT * FROM requests {where} ORDER BY created_at DESC", params)
        rows = await cursor.fetchall()

        results = []
        for row in rows:
            try:
                req = await _row_to_request_dict(db, row)
                results.append(RelRequest(**req))
            except Exception as _row_err:
                req_num = row[1] if len(row) > 1 else '?'
                logging.error(f"Skipping malformed request {req_num}: {_row_err}")
        return results
    finally:
        await db.close()

@api_router.get("/requests/{request_id}", response_model=RelRequest)
async def get_request(request_id: str, current_user: User = Depends(get_current_user)):
    db = await get_db()
    try:
        cursor = await db.execute("SELECT * FROM requests WHERE id = ?", (request_id,))
        row = await cursor.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Request not found")
        req = await _row_to_request_dict(db, row)
        return RelRequest(**req)
    finally:
        await db.close()

# ================================================================
# Excel Report Generation
# ================================================================
def _generate_request_report_excel(req: dict) -> bytes:
    """Generate a Reliability Test Report .xlsx matching the standard template format."""
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    ws = wb.active
    ws.title = "Reliability Test Report"

    # â”€â”€ Column widths â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    col_widths = {
        1: 3,  2: 22, 3: 3,  4: 3,  5: 22, 6: 3,
        7: 3,  8: 3,  9: 20, 10: 3, 11: 3, 12: 22,
        13: 3, 14: 3, 15: 3, 16: 24, 17: 3, 18: 3,
    }
    for col_idx, w in col_widths.items():
        ws.column_dimensions[get_column_letter(col_idx)].width = w

    # â”€â”€ Style helpers â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    thin_side = Side(style='thin', color='A0A0A0')
    thin_border = Border(left=thin_side, right=thin_side, top=thin_side, bottom=thin_side)
    gray_fill  = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type='solid')
    hdr_fill   = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type='solid')

    def sc(row, col, value=None, bold=False, italic=False, sz=10,
           bg=None, halign='left', wrap=False, color=None, border=False):
        c = ws.cell(row=row, column=col, value=value)
        fkw = {'size': sz}
        if bold:   fkw['bold'] = True
        if italic: fkw['italic'] = True
        if color:  fkw['color'] = color
        c.font = Font(**fkw)
        c.alignment = Alignment(horizontal=halign, vertical='top', wrap_text=wrap)
        if bg:     c.fill = PatternFill(start_color=bg, end_color=bg, fill_type='solid')
        if border: c.border = thin_border
        return c

    def mc(r1, c1, r2, c2):
        ws.merge_cells(start_row=r1, start_column=c1, end_row=r2, end_column=c2)

    def hdr_cell(row, col, text, span_end_col=None):
        c = sc(row, col, text, bold=True, sz=9, bg="F2F2F2", halign='center', border=True)
        c.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        if span_end_col:
            mc(row, col, row, span_end_col)
        return c

    # â”€â”€ Extract request fields â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    steps        = req.get('steps', [])
    rr           = req.get('request_number', '')
    classification = req.get('classification', '') or ''
    originator   = req.get('originator',   '') or ''
    plant        = req.get('plant',        '') or ''
    device_name  = req.get('device_name',  '') or ''
    lot_no       = req.get('lot_no',       '') or ''
    customer     = req.get('customer',     '') or ''
    pkg_info     = req.get('pkg_info',     '') or ''
    automotive   = 'Yes' if req.get('automotive') else 'No'
    prod_hier    = req.get('product_hierarchy', '') or ''
    pdl          = req.get('pdl',          '') or ''
    body_x       = req.get('body_size_x')
    body_y       = req.get('body_size_y')
    thickness    = req.get('package_thickness')
    ball_pitch   = req.get('ball_pitch')
    ball_count   = req.get('ball_count')
    lead_pitch   = req.get('lead_pitch')
    lead_count   = req.get('lead_count')
    total_ss     = req.get('total_ss',     '') or ''
    purpose      = req.get('purpose',      '') or ''
    conclusion   = req.get('engineer_special_instruction', '') or ''
    created_by   = req.get('created_by_username', '') or ''
    all_legs     = sorted(set(s.get('leg', 1) for s in steps))

    # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
    # COVER PAGE
    # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

    # Company info â€“ top right (col 16)
    company_lines = [
        "ATP3&4 Rel/FA Laboratory",
        "119 North Science Avenue",
        "Special Economic Processing Zone",
        "Laguna Technopark, Binan Laguna",
        "Philippines 4024 PHILIPPINES",
        "Tel: 632.884.3000",
        "Fax: 632.884.3160",
    ]
    for i, line in enumerate(company_lines, 1):
        c = sc(i, 16, line, sz=9, halign='center')
        mc(i, 16, i, 18)

    # Row 12: Report title
    c = sc(12, 2, "Reliability Test Report", bold=True, sz=14)
    mc(12, 2, 12, 14)

    # Row 13: Device / Package description
    desc = " / ".join(filter(None, [device_name, pkg_info]))
    c = sc(13, 2, desc, bold=True, sz=11)
    mc(13, 2, 13, 14)

    # Row 14: RRS No. / RR number
    sc(14, 6,  "     RRS No.-", sz=10)
    sc(14, 9,  rr, bold=True, sz=10)
    sc(14, 12, "- Final. 0", sz=10)

    # Row 15: Report date
    sc(15, 8, "Report date   : ", sz=10)
    sc(15, 9, datetime.now().strftime("%B %d, %Y"), sz=10)

    # Row 16: Classification
    sc(16, 8, classification, sz=10)

    # Row 26â€“27: Reported / Approved by
    sc(26, 6, "    Reported By : ", sz=10)
    sc(26, 9, created_by, sz=10)
    sc(27, 6, "    Approved By : ", sz=10)

    # Row 44: Distribution list label
    sc(44, 2, " Distribution List :", bold=True, sz=10)

    # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
    # MAIN REPORT BODY  (starts at row 56)
    # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
    R = 56

    # RRS ref top-right of body
    sc(R, 17, "    RRS No.", sz=9)
    sc(R, 18, rr, bold=True, sz=9)

    # "Reliability Test Report" heading
    sc(R+1, 2, "Reliability Test Report", bold=True, sz=12)
    mc(R+1, 2, R+1, 14)

    # â”€â”€ 1.0 TEST PURPOSE â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    row = R + 3
    sc(row, 2, "1.0 TEST PURPOSE", bold=True, sz=10)
    mc(row, 2, row, 14)
    row += 1
    c = sc(row, 3, purpose, sz=10, wrap=True)
    c.alignment = Alignment(horizontal='left', vertical='top', wrap_text=True)
    ws.row_dimensions[row].height = max(30, min(len(purpose) // 2, 150))
    mc(row, 3, row + 1, 14)
    row += 3

    # â”€â”€ 2.0 CONCLUSION â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    sc(row, 2, "2.0 CONCLUSION", bold=True, sz=10)
    mc(row, 2, row, 14)
    row += 1
    c = sc(row, 3, conclusion, sz=10, wrap=True)
    c.alignment = Alignment(horizontal='left', vertical='top', wrap_text=True)
    ws.row_dimensions[row].height = max(40, min(len(conclusion) // 2, 200))
    mc(row, 3, row + 2, 14)
    row += 4

    # â”€â”€ 3.0 SAMPLE INFORMATION â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    sc(row, 2, "3.0 SAMPLE INFORMATION", bold=True, sz=10)
    mc(row, 2, row, 18)
    row += 1

    # General Information sub-header
    c = sc(row, 2, " General Information", bold=True, sz=10, bg="D9E1F2")
    mc(row, 2, row, 18)
    row += 1

    def info_row(r, lbl_l, val_l, lbl_r="", val_r=""):
        c1 = sc(r, 2, lbl_l, bold=True, sz=9, border=True)
        mc(r, 2, r, 3)
        c2 = sc(r, 4, val_l, sz=9, border=True)
        mc(r, 4, r, 8)
        c3 = sc(r, 9, lbl_r, bold=True, sz=9, border=True)
        mc(r, 9, r, 10)
        c4 = sc(r, 11, val_r, sz=9, border=True)
        mc(r, 11, r, 18)

    def fmt_num(v): return str(v) if v is not None else ''

    gen_fields = [
        ("Request Number",  rr,           "Product Hierarchy", prod_hier),
        ("Classification",  classification,"PDL",               pdl),
        ("Originator",      originator,    "Body Size X (mm)",  fmt_num(body_x)),
        ("Plant",           plant,         "Body Size Y (mm)",  fmt_num(body_y)),
        ("Device Name",     device_name,   "Package Thickness (mm)", fmt_num(thickness)),
        ("Lot No",          lot_no,        "Ball Pitch (mm)",   fmt_num(ball_pitch)),
        ("Customer",        customer,      "Ball Count",        fmt_num(ball_count)),
        ("PKG Info",        pkg_info,      "Lead Pitch (mm)",   fmt_num(lead_pitch)),
        ("Automotive",      automotive,    "Lead Count",        fmt_num(lead_count)),
        ("",                "",            "Total S/S",         total_ss),
    ]
    for ll, vl, lr, vr in gen_fields:
        info_row(row, ll, vl, lr, vr)
        row += 1

    row += 1  # spacer

    # Material Information sub-header
    c = sc(row, 2, " Material Information", bold=True, sz=10, bg="D9E1F2")
    mc(row, 2, row, 18)
    row += 1

    mat_left = [
        "Die Attach Material", "Die coat after W/B", "Die Pad Config", "Die Pad Metal",
        "Die Pad Pitch(Î¼m)", "Die Passivation", "Die Size (mm)", "Die Thick (Î¼m)",
        "Down Bond", "EMC/Encap Material", "Heat Dissipation Mat'l", "LF Ag Option",
        "LF Etch/Stamp", "LF Inner Lead Pitch(Î¼m)", "LF/Sub Material", "LF/Sub Pad Size(Î¼m)",
        "LF/Sub Supplier", "LF/Sub Thickness(Î¼m)", "Lid Attach Epoxy", "Line Width",
    ]
    mat_right = [
        "Mfg Site", "Plating Option", "Rel Site", "Solder Ball Attach Paste",
        "Solder Ball Material", "Solder Ball Size(mm)", "Solder Mask Material", "Solder Paste Material",
        "Sub Layer", "Sub Pad Design", "Sub Pad Opening Size", "Sub Surface Treatment",
        "UBM Material", "UBM Opening Size (Î¼m)", "Underfill Material", "Wafer Type",
        "Wire Length Max (mm)", "Wire Material", "Wire Size(Î¼m)", "Wire Supplier",
    ]
    for i in range(max(len(mat_left), len(mat_right))):
        ll = mat_left[i]  if i < len(mat_left)  else ''
        lr = mat_right[i] if i < len(mat_right) else ''
        info_row(row, ll, '', lr, '')
        row += 1

    row += 2  # spacer

    # â”€â”€ 4.0 TEST PROCEDURE â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    sc(row, 2, "4.0 TEST PROCEDURE", bold=True, sz=10)
    mc(row, 2, row, 14)
    row += 1

    for leg in all_legs:
        leg_steps = [s for s in steps if (s.get('leg', 1)) == leg]

        c = sc(row, 2, f"    4.{leg} Rel Test Traveller - LEG {leg}", italic=True, sz=9)
        mc(row, 2, row, 14)
        row += 1

        # Step table header
        hdr_cell(row, 2,  "No",                    3)
        hdr_cell(row, 4,  "TEST ITEM",              7)
        hdr_cell(row, 8,  "CONDITION / READ POINT", 12)
        hdr_cell(row, 13, "Machine",                13)
        hdr_cell(row, 14, "Rack",                   14)
        hdr_cell(row, 15, "Operator",               15)
        hdr_cell(row, 16, "Qty In",                 16)
        hdr_cell(row, 17, "Qty Out",                17)
        ws.row_dimensions[row].height = 28
        row += 1

        for s in leg_steps:
            sc(row, 2,  s.get('step_number'), sz=9, halign='center', border=True)
            mc(row, 2, row, 3)
            sc(row, 4,  s.get('step_name', ''),   sz=9, border=True); mc(row, 4,  row, 7)
            sc(row, 8,  s.get('notes', '') or '', sz=9, border=True, wrap=True); mc(row, 8, row, 12)
            sc(row, 13, s.get('machine_no', '') or '', sz=9, halign='center', border=True)
            sc(row, 14, s.get('rack_no', '') or '', sz=9, halign='center', border=True)
            sc(row, 15, s.get('operator_id', '') or '', sz=9, halign='center', border=True)
            sc(row, 16, s.get('qty_in')  if s.get('qty_in')  is not None else '', sz=9, halign='center', border=True)
            sc(row, 17, s.get('qty_out') if s.get('qty_out') is not None else '', sz=9, halign='center', border=True)
            row += 1

        row += 1  # spacer between legs

    row += 1

    # â”€â”€ 5.0 TEST MATRIX & REL TEST ITEMS â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    sc(row, 2, "5.0 TEST MATRIX & REL TEST ITEMS", bold=True, sz=10)
    mc(row, 2, row, 18)
    row += 1

    mx_cols = [
        (2,  3,  "Leg"),
        (4,  5,  "Assy Lot No"),
        (6,  7,  "Other Info"),
        (8,  9,  "Test Type"),
        (10, 11, "Test Item"),
        (12, 13, "Test Condition"),
        (14, 15, "Reflow/R. Point"),
        (16, 16, "E/L"),
        (17, 17, "O/S"),
        (18, 18, "SAT"),
    ]
    for c1, c2, label in mx_cols:
        c = ws.cell(row=row, column=c1, value=label)
        c.font = Font(bold=True, size=9)
        c.fill = hdr_fill
        c.border = thin_border
        c.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        if c1 != c2: mc(row, c1, row, c2)
    ws.row_dimensions[row].height = 28
    row += 1

    for leg in all_legs:
        leg_steps = [s for s in steps if (s.get('leg', 1)) == leg]
        has_sat = any('sat' in s.get('step_name', '').lower() for s in leg_steps)
        has_os  = any(s.get('step_name', '').lower() in ('o/s', 'open/short') for s in leg_steps)
        mx_vals = [
            f"LEG{leg}", lot_no, device_name, "", "", "", "",
            "", "X" if has_os else "", "X" if has_sat else "",
        ]
        for (c1, c2, _), val in zip(mx_cols, mx_vals):
            c = ws.cell(row=row, column=c1, value=val)
            c.font = Font(size=9)
            c.border = thin_border
            c.alignment = Alignment(horizontal='center', vertical='center')
            if c1 != c2: mc(row, c1, row, c2)
        row += 1

    row += 2

    # â”€â”€ 6.0 TEST RESULTS â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    sc(row, 2, "6.0 TEST RESULTS", bold=True, sz=10)
    mc(row, 2, row, 18)
    row += 1

    # 6.1 SAT Test
    sc(row, 2, "    6.1 SAT Test", bold=True, italic=True, sz=9)
    mc(row, 2, row, 18)
    row += 1

    sat_cols = [
        (2, 3,   "Leg"),
        (4, 5,   "Test"),
        (6, 7,   "Reading Point"),
        (8, 8,   "SS"),
        (9, 9,   "Result"),
        (10, 10, "Before/After"),
        (11, 12, "T1"),
        (13, 13, "T2"),
        (14, 14, "T3"),
        (15, 15, "T4"),
        (16, 16, "T5"),
    ]
    for c1, c2, label in sat_cols:
        c = ws.cell(row=row, column=c1, value=label)
        c.font = Font(bold=True, size=8)
        c.fill = hdr_fill
        c.border = thin_border
        c.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        if c1 != c2: mc(row, c1, row, c2)
    ws.row_dimensions[row].height = 28
    row += 1

    for leg in all_legs:
        leg_steps = [s for s in steps if (s.get('leg', 1)) == leg]
        sat_step = next((s for s in leg_steps if 'sat' in s.get('step_name', '').lower()), None)
        if sat_step:
            result = 'Pass' if sat_step.get('status') == 'completed' else ''
            for ba in ['Before', 'After']:
                vals = [f"LEG{leg}", "SAT", "", total_ss, result if ba == 'After' else '', ba,
                        '0', '-', '0', '-', '0']
                for (c1, c2, _), val in zip(sat_cols, vals):
                    c = ws.cell(row=row, column=c1, value=val)
                    c.font = Font(size=8)
                    c.border = thin_border
                    c.alignment = Alignment(horizontal='center')
                    if c1 != c2: mc(row, c1, row, c2)
                row += 1

    # Notes
    row += 1
    for note in [
        "- Note",
        "T1 : Delamination at EMC or Encap / Die Top Surface",
        "T2 : Delamination at Die Attach Region",
        "T3 : Delamination at EMC or Encap / pad Top or Laminate Surface surrounding die",
        "T5 : Delamination at Lead finger / EMC",
        "",
        " Refer to the specification # 001-2531 for the Pass / Fail Criteria",
    ]:
        c = sc(row, 2, note, sz=8, italic=(note.startswith(' Refer')))
        mc(row, 2, row, 18)
        row += 1

    # 6.2 Open/Short Test
    row += 1
    sc(row, 2, "    6.2 Open/Short Test", bold=True, italic=True, sz=9)
    mc(row, 2, row, 18)
    row += 1

    os_cols = [
        (2,  3,  "Leg"),
        (4,  6,  "Test Item"),
        (7,  9,  "Test Condition"),
        (10, 12, "Reading Point"),
        (13, 13, "Result"),
        (14, 14, "SS"),
        (15, 15, "#Fail."),
        (16, 18, "Fail. Mode"),
    ]
    for c1, c2, label in os_cols:
        c = ws.cell(row=row, column=c1, value=label)
        c.font = Font(bold=True, size=8)
        c.fill = hdr_fill
        c.border = thin_border
        c.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        if c1 != c2: mc(row, c1, row, c2)
    ws.row_dimensions[row].height = 28
    row += 1

    for leg in all_legs:
        leg_steps = [s for s in steps if (s.get('leg', 1)) == leg]
        os_step = next((s for s in leg_steps if s.get('step_name', '').lower() in ('o/s', 'open/short')), None)
        if os_step:
            result = 'Pass' if os_step.get('status') == 'completed' else ''
            os_vals = [f"LEG{leg}", "O/S", "", "Open/Short", result, total_ss, 'n/a', 'n/a']
            for (c1, c2, _), val in zip(os_cols, os_vals):
                c = ws.cell(row=row, column=c1, value=val)
                c.font = Font(size=8)
                c.border = thin_border
                c.alignment = Alignment(horizontal='center')
                if c1 != c2: mc(row, c1, row, c2)
            row += 1

    row += 1
    sc(row, 2, "- Note", bold=True, sz=8)
    mc(row, 2, row, 18)
    row += 1
    sc(row, 2, " Refer to the specification # 001-2150 for the Pass / Fail Criteria",
       italic=True, sz=8)
    mc(row, 2, row, 18)
    row += 2

    # â”€â”€ Page setup â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    ws.print_area = f'A1:{get_column_letter(18)}{row}'
    ws.page_setup.orientation = ws.ORIENTATION_LANDSCAPE
    ws.page_setup.fitToPage = True
    ws.page_setup.fitToWidth = 1

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output.getvalue()


# ================================================================
# SAT Report Generation
# ================================================================
def _generate_sat_report_excel(req: dict) -> bytes:
    """Generate a SAT Observation Report .xlsx with per-leg sheets and embedded images."""
    import os
    from openpyxl.utils import get_column_letter
    from openpyxl.drawing.image import Image as XLImage

    wb = Workbook()

    steps         = req.get('steps', [])
    rr            = req.get('request_number', '')
    device_name   = req.get('device_name',  '') or ''
    lot_no        = req.get('lot_no',       '') or ''
    total_ss      = req.get('total_ss',     '') or ''
    classification = req.get('classification', '') or ''
    originator    = req.get('originator',   '') or ''
    pkg_info      = req.get('pkg_info',     '') or ''

    all_legs = sorted(set(s.get('leg', 1) for s in steps))

    uploads_dir = os.path.join(os.path.dirname(__file__), 'uploads')

    thin_side   = Side(style='thin', color='A0A0A0')
    thin_border = Border(left=thin_side, right=thin_side, top=thin_side, bottom=thin_side)
    blue_fill   = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type='solid')
    gray_fill   = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type='solid')

    def _sc(ws, row, col, value=None, bold=False, italic=False, sz=10,
            bg=None, halign='left', wrap=False, color=None, border=False):
        c = ws.cell(row=row, column=col, value=value)
        fkw = {'size': sz}
        if bold:   fkw['bold'] = True
        if italic: fkw['italic'] = True
        if color:  fkw['color'] = color
        c.font = Font(**fkw)
        c.alignment = Alignment(horizontal=halign, vertical='top', wrap_text=wrap)
        if bg:     c.fill = PatternFill(start_color=bg, end_color=bg, fill_type='solid')
        if border: c.border = thin_border
        return c

    def _mc(ws, r1, c1, r2, c2):
        ws.merge_cells(start_row=r1, start_column=c1, end_row=r2, end_column=c2)

    def _set_col_widths(ws, widths):
        for col_idx, w in widths.items():
            ws.column_dimensions[get_column_letter(col_idx)].width = w

    # â”€â”€ SUMMARY SHEET â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    ws_sum = wb.active
    ws_sum.title = "Summary"
    _set_col_widths(ws_sum, {1:3, 2:12, 3:22, 4:18, 5:10, 6:14, 7:14, 8:10, 9:10, 10:12, 11:12})

    r = 1
    # Company header (top-right)
    for i, line in enumerate([
        "ATP3&4 Rel/FA Laboratory", "119 North Science Avenue",
        "Special Economic Processing Zone", "Laguna Technopark, Binan Laguna",
        "Philippines 4024 PHILIPPINES", "Tel: 632.884.3000",
    ], 1):
        c = _sc(ws_sum, i, 9, line, sz=8, halign='center')
        _mc(ws_sum, i, 9, i, 11)

    # Title block
    _sc(ws_sum, 8, 2, "SAT Observation Report", bold=True, sz=14, halign='left')
    _mc(ws_sum, 8, 2, 8, 8)
    _sc(ws_sum, 9, 2, f"{device_name}  /  {pkg_info}", bold=True, sz=10)
    _mc(ws_sum, 9, 2, 9, 8)
    _sc(ws_sum, 10, 2, f"RRS No.:  {rr}", sz=10)
    _mc(ws_sum, 10, 2, 10, 5)
    _sc(ws_sum, 10, 6, f"Date:  {datetime.now().strftime('%Y/%m/%d')}", sz=10)
    _mc(ws_sum, 10, 6, 10, 8)
    _sc(ws_sum, 11, 2, f"Originator:  {originator}    Classification:  {classification}", sz=9, wrap=True)
    _mc(ws_sum, 11, 2, 11, 8)
    _sc(ws_sum, 12, 2, f"Lot No.:  {lot_no}    Total S/S:  {total_ss}", sz=9)
    _mc(ws_sum, 12, 2, 12, 8)

    r = 14
    # SAT results table header
    sat_hdr_cols = [
        (2, 3,  "Leg"),
        (4, 5,  "Step"),
        (6, 7,  "Test Condition"),
        (8, 8,  "Machine"),
        (9, 9,  "Operator"),
        (10, 10,"Qty In"),
        (11, 11,"Qty Out"),
    ]
    for c1, c2, label in sat_hdr_cols:
        c = ws_sum.cell(row=r, column=c1, value=label)
        c.font = Font(bold=True, size=9)
        c.fill = blue_fill
        c.border = thin_border
        c.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        if c1 != c2: _mc(ws_sum, r, c1, r, c2)
    ws_sum.row_dimensions[r].height = 24
    r += 1

    for leg in all_legs:
        leg_steps  = [s for s in steps if (s.get('leg', 1)) == leg]
        sat_steps  = [s for s in leg_steps if 'sat' in s.get('step_name', '').lower()]

        # Determine leg condition from Reliability Test step
        rel_step   = next((s for s in leg_steps if 'reliability' in s.get('step_name', '').lower()), None)
        leg_cond   = (rel_step.get('custom_fields', {}) or {}).get('test_condition', '') if rel_step else ''

        first_sat = True
        for s in sat_steps:
            cf    = s.get('custom_fields', {}) or {}
            step_cond = cf.get('test_condition', '') or leg_cond
            row_vals = [
                f"Leg {leg}" if first_sat else "",
                leg_cond    if first_sat else "",
                s.get('step_name', ''),
                step_cond,
                s.get('machine_no', '') or '',
                s.get('operator_id', '') or '',
                s.get('qty_in', '')  if s.get('qty_in')  is not None else '',
                s.get('qty_out', '') if s.get('qty_out') is not None else '',
            ]
            for (c1, c2, _), val in zip(sat_hdr_cols, row_vals):
                c = ws_sum.cell(row=r, column=c1, value=val)
                c.font = Font(size=9, bold=(first_sat and c1 == 2))
                c.border = thin_border
                c.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                if c1 != c2: _mc(ws_sum, r, c1, r, c2)
            r += 1
            first_sat = False

    # Freeze top rows
    ws_sum.freeze_panes = "B15"

    # â”€â”€ PER-LEG SHEETS â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    for leg in all_legs:
        leg_steps = [s for s in steps if (s.get('leg', 1)) == leg]
        sat_steps = [s for s in leg_steps if 'sat' in s.get('step_name', '').lower()]

        # Determine leg condition name for sheet title
        rel_step  = next((s for s in leg_steps if 'reliability' in s.get('step_name', '').lower()), None)
        leg_cond  = (rel_step.get('custom_fields', {}) or {}).get('test_condition', '') if rel_step else ''

        sheet_name = f"Leg {leg}"
        ws = wb.create_sheet(title=sheet_name)
        # Cols: 1=spacer, 2=T-Scan, 3=divider, 4=1.C-Scan, 5=divider, 6=2.C-Scan, 7=divider, 8=Attachments, 9=end
        _set_col_widths(ws, {1:3, 2:27, 3:2, 4:27, 5:2, 6:27, 7:2, 8:27, 9:3})

        # Company info top-right (col 6 = F, spans F-G)
        for i, line in enumerate([
            "ATP3&4 Rel/FA Laboratory", "Laguna Technopark, Binan Laguna",
            "Philippines 4024 PHILIPPINES",
        ], 1):
            c = _sc(ws, i, 6, line, sz=8, halign='center')
            _mc(ws, i, 6, i, 7)

        # Leg header (B to E = cols 2-5)
        _sc(ws, 1, 2, f"Leg {leg} â€” SAT Observation Report",  bold=True, sz=13)
        _mc(ws, 1, 2, 1, 5)
        _sc(ws, 2, 2, f"RRS No.: {rr}   |   Device: {device_name}   |   Condition: {leg_cond}", sz=9, wrap=True)
        _mc(ws, 2, 2, 2, 5)
        _sc(ws, 3, 2, f"Lot No.: {lot_no}   |   Total S/S: {total_ss}", sz=9)
        _mc(ws, 3, 2, 3, 5)

        r = 5
        # Step summary table â€” 3 zones matching image columns (B=2, D=4, F=6)
        # Zone headers
        for col, label in [(2, "Step # | Step Name"), (4, "Test Condition"), (6, "Machine  |  Operator  |  Qty In â†’ Out")]:
            c = ws.cell(row=r, column=col, value=label)
            c.font = Font(bold=True, size=8)
            c.fill = gray_fill
            c.border = thin_border
            c.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            _mc(ws, r, col, r, col + 1)   # each header spans its wide col + the divider col
        ws.row_dimensions[r].height = 20
        r += 1

        for s in leg_steps:
            cf      = s.get('custom_fields', {}) or {}
            snum    = s.get('step_number', '')
            sname   = s.get('step_name', '')
            tc      = cf.get('test_condition', '') or ''
            mach    = s.get('machine_no', '') or ''
            oper    = s.get('operator_id', '') or ''
            qin     = s.get('qty_in')
            qout    = s.get('qty_out')
            qty_str  = f"{qin if qin is not None else 'â€”'} â†’ {qout if qout is not None else 'â€”'}"
            mach_str = f"{mach}  |  {oper}  |  {qty_str}"
            for col, val in [(2, f"{snum}   {sname}"), (4, tc), (6, mach_str)]:
                c = ws.cell(row=r, column=col, value=val)
                c.font = Font(size=8)
                c.border = thin_border
                c.alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
                _mc(ws, r, col, r, col + 1)
            r += 1

        r += 1  # spacer

        # Observations section heading
        _sc(ws, r, 2, "7.0 OBSERVATION and ANALYSIS", bold=True, sz=10)
        _mc(ws, r, 2, r, 9)
        r += 2

        photo_no = 1
        # Each image: fits within one 27-unit col (~189px wide, same height)
        img_w    = 189
        img_h    = 165
        rows_per_img = max(1, round(img_h / 15))
        row_h_pt     = img_h / rows_per_img

        # 3 category groups â€” col anchors: T-Scanâ†’col2(B), 1.C-Scanâ†’col4(D), 2.C-Scanâ†’col6(F), Attachmentsâ†’col8(H)
        SAT_GROUPED = [
            ("1\u201324",  [("t_scan_1_24",    "T-Scan 1\u201324"),
                            ("c_scan_1_1_24",  "1. C-Scan 1\u201324"),
                            ("c_scan_2_1_24",  "2. C-Scan 1\u201324 (Opt.)")], "sat_files_1_24"),
            ("25\u201348", [("t_scan_25_48",   "T-Scan 25\u201348"),
                            ("c_scan_1_25_48", "1. C-Scan 25\u201348"),
                            ("c_scan_2_25_48", "2. C-Scan 25\u201348 (Opt.)")], "sat_files_25_48"),
            ("49\u201377", [("t_scan_49_77",   "T-Scan 49\u201377"),
                            ("c_scan_1_49_77", "1. C-Scan 49\u201377"),
                            ("c_scan_2_49_77", "2. C-Scan 49\u201377 (Opt.)")], "sat_files_49_77"),
        ]
        # Column start for each category index (0â†’col2 B, 1â†’col4 D, 2â†’col6 F); attachments at col8 H
        CAT_COLS = [2, 4, 6]
        ATT_COL   = 8

        for sat_step in sat_steps:
            raw_att   = sat_step.get('attachments') or {}
            cf        = sat_step.get('custom_fields', {}) or {}
            step_cond = cf.get('test_condition', '')

            # SAT step header
            step_label = f"SAT \u2014 {step_cond}" if step_cond else f"SAT Step {sat_step.get('step_number','')}"
            _sc(ws, r, 2, step_label, bold=True, sz=9, bg="E8EFFD")
            _mc(ws, r, 2, r, 9)
            r += 1

            if not isinstance(raw_att, dict):
                # Legacy flat list â€” render sequentially in col 2
                flat_list = raw_att if isinstance(raw_att, list) else []
                if not flat_list:
                    _sc(ws, r, 2, "(No SAT images attached)", italic=True, sz=8, color="888888")
                    _mc(ws, r, 2, r, 9)
                    r += 2
                else:
                    for att_url in flat_list:
                        filename = att_url.split('/')[-1]
                        img_path = os.path.join(uploads_dir, filename)
                        _sc(ws, r, 2, f"Photo {photo_no}: Leg {leg} SAT", bold=True, sz=9)
                        _mc(ws, r, 2, r, 9)
                        r += 1
                        if os.path.exists(img_path):
                            try:
                                xl_img = XLImage(img_path)
                                xl_img.width  = img_w * 3
                                xl_img.height = img_h * 2
                                ws.add_image(xl_img, f"B{r}")
                                rn = max(1, round(xl_img.height / 15))
                                for ir in range(r, r + rn):
                                    ws.row_dimensions[ir].height = xl_img.height / rn
                                r += rn
                            except Exception as e_img:
                                logging.warning(f"SAT report legacy embed: {e_img}")
                                _sc(ws, r, 2, f"[Image: {filename}]", italic=True, sz=8, color="888888")
                                _mc(ws, r, 2, r, 9)
                                r += 1
                        else:
                            _sc(ws, r, 2, f"[Not found: {filename}]", italic=True, sz=8, color="BB0000")
                            _mc(ws, r, 2, r, 9)
                            r += 1
                        photo_no += 1
                    r += 1
                continue

            # New dict format â€” render each range group as 4 side-by-side columns
            any_images = any(raw_att.get(k) for _, cats, _ in SAT_GROUPED for k, _ in cats) \
                         or any(raw_att.get(fk) for _, _, fk in SAT_GROUPED)
            if not any_images:
                _sc(ws, r, 2, "(No SAT images attached)", italic=True, sz=8, color="888888")
                _mc(ws, r, 2, r, 9)
                r += 2
                continue

            for range_label, group_cats, file_key in SAT_GROUPED:
                group_lists = [raw_att.get(k) or [] for k, _ in group_cats]
                extra_imgs  = [v for v in (raw_att.get(file_key) or []) if isinstance(v, str)]
                att_col_letter = get_column_letter(ATT_COL)

                # Group header spanning all 4 image columns
                _sc(ws, r, 2, f"Samples {range_label}", bold=True, sz=9, bg="DDE8F5")
                _mc(ws, r, 2, r, 9)
                r += 1

                # Column sub-headers: T-Scan | 1.C-Scan | 2.C-Scan | Attachments
                for ci, (_, cat_label) in enumerate(group_cats):
                    col = CAT_COLS[ci]
                    _sc(ws, r, col, cat_label, bold=True, sz=8, bg="EEF4FB")
                    _mc(ws, r, col, r, col + 1)
                _sc(ws, r, ATT_COL, "Attachments", bold=True, sz=8, bg="EEF4FB")
                _mc(ws, r, ATT_COL, r, ATT_COL + 1)
                r += 1

                # Image slots â€” stacked per column, up to max across all 4 cols
                max_slots = max(max((len(lst) for lst in group_lists), default=0), len(extra_imgs))
                for slot in range(max_slots):
                    slot_start = r
                    any_placed = False
                    for ci, (_, cat_label) in enumerate(group_cats):
                        lst = group_lists[ci]
                        col = CAT_COLS[ci]
                        col_letter = get_column_letter(col)
                        if slot >= len(lst):
                            # No image for this slot â€” leave blank
                            continue
                        url      = lst[slot]
                        filename = url.split('/')[-1]
                        img_path = os.path.join(uploads_dir, filename)
                        if os.path.exists(img_path):
                            try:
                                xl_img        = XLImage(img_path)
                                xl_img.width  = img_w
                                xl_img.height = img_h
                                ws.add_image(xl_img, f"{col_letter}{slot_start}")
                                any_placed = True
                                photo_no  += 1
                            except Exception as e_img:
                                logging.warning(f"SAT report grouped embed: {e_img}")
                                _sc(ws, slot_start, col, f"[Image: {filename}]",
                                    italic=True, sz=7, color="888888")
                                _mc(ws, slot_start, col, slot_start, col + 1)
                        else:
                            _sc(ws, slot_start, col, f"[Not found: {filename}]",
                                italic=True, sz=7, color="BB0000")
                            _mc(ws, slot_start, col, slot_start, col + 1)
                    # 4th column: extra attachment images
                    if slot < len(extra_imgs):
                        att_url  = extra_imgs[slot]
                        filename = att_url.split('/')[-1]
                        img_path = os.path.join(uploads_dir, filename)
                        if os.path.exists(img_path):
                            try:
                                xl_img        = XLImage(img_path)
                                xl_img.width  = img_w
                                xl_img.height = img_h
                                ws.add_image(xl_img, f"{att_col_letter}{slot_start}")
                                any_placed = True
                                photo_no  += 1
                            except Exception as e_img:
                                logging.warning(f"SAT report attachment embed: {e_img}")
                                _sc(ws, slot_start, ATT_COL, f"[Image: {filename}]",
                                    italic=True, sz=7, color="888888")
                                _mc(ws, slot_start, ATT_COL, slot_start, ATT_COL + 1)
                        else:
                            _sc(ws, slot_start, ATT_COL, f"[Not found: {filename}]",
                                italic=True, sz=7, color="BB0000")
                            _mc(ws, slot_start, ATT_COL, slot_start, ATT_COL + 1)
                    if any_placed:
                        for ir in range(slot_start, slot_start + rows_per_img):
                            ws.row_dimensions[ir].height = row_h_pt
                        r = slot_start + rows_per_img
                    else:
                        r = slot_start + 1
                r += 1  # spacer after range group

            r += 1  # spacer after SAT step

        ws.print_area = f"A1:{get_column_letter(9)}{r}"
        ws.page_setup.fitToPage  = True
        ws.page_setup.fitToWidth = 1
        ws.page_setup.orientation = ws.ORIENTATION_LANDSCAPE

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output.getvalue()


@api_router.get("/requests/{request_id}/sat-report")
async def download_sat_report(
    request_id: str,
    current_user: User = Depends(get_current_user)
):
    """Generate and download a SAT Observation Report Excel file for a request."""
    db = await get_db()
    try:
        cursor = await db.execute("SELECT * FROM requests WHERE id = ?", (request_id,))
        row = await cursor.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Request not found")
        req = await _row_to_request_dict(db, row)
    finally:
        await db.close()

    try:
        excel_bytes = _generate_sat_report_excel(req)
    except Exception as e:
        logging.error(f"SAT report generation error for {request_id}: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"SAT report generation failed: {str(e)}")

    rr       = req.get('request_number', request_id)
    filename = f"SATReport_{rr}.xlsx"
    return StreamingResponse(
        io.BytesIO(excel_bytes),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )


@api_router.get("/requests/{request_id}/report")
async def download_request_report(
    request_id: str,
    current_user: User = Depends(get_current_user)
):
    """Generate and download a Reliability Test Report Excel file for a request."""
    db = await get_db()
    try:
        cursor = await db.execute("SELECT * FROM requests WHERE id = ?", (request_id,))
        row = await cursor.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Request not found")
        req = await _row_to_request_dict(db, row)
    finally:
        await db.close()

    try:
        excel_bytes = _generate_request_report_excel(req)
    except Exception as e:
        logging.error(f"Report generation error for {request_id}: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Report generation failed: {str(e)}")

    rr = req.get('request_number', request_id)
    filename = f"ReliabilityReport_{rr}.xlsx"
    return StreamingResponse(
        io.BytesIO(excel_bytes),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )


# Additional report endpoint for step-level exports
@api_router.get('/reports/steps')
async def download_step_report(
    step_name: Optional[str] = None,
    status: Optional[str] = None,
    request_type: Optional[str] = None,
    current_user: User = Depends(get_current_user)
):
    db = await get_db()
    try:
        conditions = ['1=1']
        params = []

        if step_name:
            conditions.append('LOWER(ps.step_name) = ?')
            params.append(step_name.strip().lower())
        if status:
            conditions.append('ps.status = ?')
            params.append(status.strip())
        if request_type:
            request_type = request_type.strip().upper()
            if request_type not in ('REL', 'RMS'):
                raise HTTPException(status_code=400, detail="Invalid request_type: must be 'REL' or 'RMS'.")
            conditions.append('r.request_type = ?')
            params.append(request_type)

        query = (
            "SELECT r.request_number, r.request_type, r.device_name, r.lot_no, r.customer, r.classification, "
            "ps.step_name, ps.leg, ps.machine_no, ps.rack_no, ps.operator_id, COALESCE(e.name,'') as employee_name, ps.status, ps.started_at, ps.completed_at, ps.notes "
            "FROM process_steps ps "
            "JOIN requests r ON ps.request_id = r.id "
            "LEFT JOIN employees e ON ps.operator_id = e.id "
            f"WHERE {' AND '.join(conditions)} ORDER BY r.request_number, ps.leg, ps.step_number"
        )
        cursor = await db.execute(query, params)
        rows = await cursor.fetchall()
        step_rows = [{
            'request_number': r[0], 'request_type': r[1], 'device_name': r[2],
            'lot_no': r[3], 'customer': r[4], 'classification': r[5],
            'step_name': r[6], 'leg': r[7], 'machine_no': r[8], 'rack_no': r[9],
            'operator_id': r[10], 'employee_name': r[11], 'status': r[12], 'started_at': r[13], 'completed_at': r[14], 'notes': r[15]
        } for r in rows]
    finally:
        await db.close()

    try:
        wb = Workbook()
        ws = wb.active
        ws.title = 'Step Report'
        headers = [
            'Request Number', 'Request Type', 'Device Name', 'Lot No', 'Customer', 'Classification',
            'Step Name', 'Leg', 'Machine', 'Rack', 'Operator', 'Status', 'Start', 'End', 'Notes'
        ]
        for col, h in enumerate(headers, start=1):
            c = ws.cell(row=1, column=col, value=h)
            c.font = Font(bold=True)
            c.alignment = Alignment(horizontal='center', vertical='center')
            ws.column_dimensions[get_column_letter(col)].width = max(12, len(h) + 2)

        for idx, row_data in enumerate(step_rows, start=2):
            values = [
                row_data.get('request_number', ''), row_data.get('request_type', ''),
                row_data.get('device_name', ''), row_data.get('lot_no', ''),
                row_data.get('customer', ''), row_data.get('classification', ''),
                row_data.get('step_name', ''), row_data.get('leg', ''),
                row_data.get('machine_no', ''), row_data.get('rack_no', ''),
                row_data.get('operator_id', ''), row_data.get('status', ''),
                row_data.get('started_at', ''), row_data.get('completed_at', ''),
                row_data.get('notes', ''),
            ]
            for col, val in enumerate(values, start=1):
                ws.cell(row=idx, column=col, value=val)

        output = io.BytesIO()
        wb.save(output)
        output.seek(0)

        return StreamingResponse(
            output,
            media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={'Content-Disposition': 'attachment; filename="StepReport.xlsx"'}
        )
    except Exception as e:
        logging.error(f"Step report generation error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail='Step report generation failed')


@api_router.get('/reports/presentation')
async def download_presentation_report(
    current_user: User = Depends(get_current_user)
):
    """Generate and download the project timeline overview PowerPoint presentation."""
    db = await get_db()
    try:
        cursor = await db.execute(
            "SELECT request_number, created_at, deadline, status FROM requests ORDER BY created_at DESC LIMIT 50"
        )
        rows = await cursor.fetchall()
        timeline_events = [
            {
                'request_number': r['request_number'],
                'created_at': r['created_at'],
                'deadline': r['deadline'] or '',
                'status': r['status']
            }
            for r in rows
        ]
    finally:
        await db.close()

    try:
        pptx_path = build_powerpoint(timeline_events=timeline_events)
    except Exception as e:
        logging.error(f"Presentation generation failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail='Presentation generation failed')

    return FileResponse(
        pptx_path,
        media_type='application/vnd.openxmlformats-officedocument.presentationml.presentation',
        headers={'Content-Disposition': f'attachment; filename="Rel_Timeline_Presentation_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pptx"'}
    )


# Performance / Employee statistics endpoint
@api_router.get('/reports/employee-performance')
async def get_employee_performance(
    period: Optional[str] = 'day',
    employee_id: Optional[str] = None,
    date_from: Optional[str] = None,
    date_to: Optional[str] = None,
    current_user: User = Depends(get_current_user)
):
    if period not in ('day', 'week', 'month'):
        raise HTTPException(status_code=400, detail="Invalid period: must be 'day', 'week', or 'month'")

    POINTS_BY_STEP = {
        'incoming inspection': 1,
        'visual': 1,
        'serialize samples': 1,
        'o/s': 2,
        'sat': 2,
        'bake': 2,
        'dry bake': 2,
        't & h soak': 3,
        'forced convection reflow (fcr)': 3,
        'reflow': 3,
        'electrical test': 3,
        'reliability test': 4,
        'moisture resistance test': 4,
        'preconditioning (precon)': 3,
        'temperature cycle': 3,
        'whisker test': 3,
        'staging': 1,
    }

    def step_points(step_name):
        return POINTS_BY_STEP.get((step_name or '').strip().lower(), 1)

    db = await get_db()
    try:
        conditions = ["ps.status = 'completed'"]
        params = []

        if employee_id:
            conditions.append('ps.operator_id = ?')
            params.append(employee_id)
        if date_from:
            conditions.append("DATE(ps.completed_at) >= DATE(?)")
            params.append(date_from)
        if date_to:
            conditions.append("DATE(ps.completed_at) <= DATE(?)")
            params.append(date_to)

        query = (
            "SELECT ps.operator_id, COALESCE(e.name, '') as employee_name, ps.step_name, ps.completed_at "
            "FROM process_steps ps "
            "LEFT JOIN employees e ON ps.operator_id = e.id "
            f"WHERE {' AND '.join(conditions)} ORDER BY ps.completed_at DESC"
        )

        cursor = await db.execute(query, params)
        rows = await cursor.fetchall()

        records = []
        for r in rows:
            completed_at = r[3]
            if not completed_at:
                continue
            date_key = completed_at[:10]
            week_key = None
            month_key = completed_at[:7]
            if period == 'week':
                dt = datetime.fromisoformat(completed_at[:19])
                week_key = f"{dt.isocalendar()[0]}-W{dt.isocalendar()[1]:02d}"
            if period == 'day':
                period_key = date_key
            elif period == 'week':
                period_key = week_key
            else:
                period_key = month_key

            points = step_points(r[2])
            records.append({
                'employee_id': r[0],
                'employee_name': r[1],
                'step_name': r[2],
                'completed_at': completed_at,
                'period': period_key,
                'points': points,
            })

        by_employee_period = {}
        for rec in records:
            key = (rec['employee_id'], rec['period'])
            if key not in by_employee_period:
                by_employee_period[key] = {
                    'employee_id': rec['employee_id'],
                    'employee_name': rec['employee_name'],
                    'period': rec['period'],
                    'steps_completed': 0,
                    'points': 0,
                }
            by_employee_period[key]['steps_completed'] += 1
            by_employee_period[key]['points'] += rec['points']

        return {
            'period': period,
            'aggregation': list(by_employee_period.values()),
            'records': records,
        }
    finally:
        await db.close()


@api_router.patch("/requests/{request_id}/steps/{step_number}")
async def update_step(
    request_id: str, step_number: int, step_update: StepUpdate,
    leg: int = 1,
    current_user: User = Depends(require_permission('update_steps'))
):
    db = await get_db()
    try:
        cursor = await db.execute("SELECT id FROM requests WHERE id = ?", (request_id,))
        if not await cursor.fetchone():
            raise HTTPException(status_code=404, detail="Request not found")

        cursor = await db.execute(
            "SELECT id, status, operator_id, started_at FROM process_steps WHERE request_id = ? AND leg = ? AND step_number = ?",
            (request_id, leg, step_number))
        step_row = await cursor.fetchone()
        if not step_row:
            raise HTTPException(status_code=400, detail="Invalid step number")

        # Technicians can edit completed step fields; system tracks who edited in updated_by.
        update_data = step_update.model_dump(exclude_unset=True)

        # Validate required fields when setting status to 'completed'
        new_status_raw = update_data.get('status')
        if new_status_raw is not None:
            new_status_val = new_status_raw.value if hasattr(new_status_raw, 'value') else new_status_raw
            if new_status_val == 'completed':
                final_operator = update_data.get('operator_id') or step_row[2]
                final_started = update_data.get('started_at') or step_row[3]
                missing_fields = []
                if not final_operator:
                    missing_fields.append('Employee No.')
                if not final_started:
                    missing_fields.append('Start of Process')
                if missing_fields:
                    raise HTTPException(
                        status_code=400,
                        detail=f"Required to complete step: {', '.join(missing_fields)}"
                    )

        if not update_data:
            return {"message": "No updates provided"}

        set_clauses = []
        values: list = []

        if 'status' in update_data:
            val = update_data['status'].value if hasattr(update_data['status'], 'value') else update_data['status']
            set_clauses.append("status = ?")
            values.append(val)
            if val == 'in_progress':
                # Auto-set started_at if not already set
                cursor = await db.execute(
                    "SELECT started_at FROM process_steps WHERE request_id = ? AND leg = ? AND step_number = ?",
                    (request_id, leg, step_number))
                existing = await cursor.fetchone()
                if not existing[0]:
                    set_clauses.append("started_at = ?")
                    values.append(datetime.now(timezone.utc).isoformat())
            if val == 'completed':
                set_clauses.append("completed_at = ?")
                values.append(datetime.now(timezone.utc).isoformat())

        for field in ['started_at', 'completed_at']:
            if field in update_data and update_data[field] is not None:
                try:
                    datetime.fromisoformat(update_data[field])
                except (ValueError, TypeError):
                    raise HTTPException(status_code=400, detail=f"Invalid datetime format for '{field}': {update_data[field]!r}")
                set_clauses.append(f"{field} = ?")
                values.append(update_data[field])

        for field in ['step_name', 'machine_no', 'rack_no', 'operator_id', 'tray_no', 'notes']:
            if field in update_data:
                set_clauses.append(f"{field} = ?")
                values.append(update_data[field])

        for field in ['qty_in', 'qty_out']:
            if field in update_data:
                set_clauses.append(f"{field} = ?")
                values.append(update_data[field])

        if 'attachments' in update_data:
            set_clauses.append("attachments = ?")
            values.append(json.dumps(update_data['attachments']))

        if 'custom_fields' in update_data:
            set_clauses.append("custom_fields = ?")
            values.append(json.dumps(update_data['custom_fields']))

        # Always track who last edited the step
        set_clauses.append("updated_by = ?")
        values.append(current_user.username)

        if set_clauses:
            values.extend([request_id, leg, step_number])
            await db.execute(
                f"UPDATE process_steps SET {', '.join(set_clauses)} WHERE request_id = ? AND leg = ? AND step_number = ?",
                values)

        # Recalculate status
        cursor = await db.execute("SELECT status FROM requests WHERE id = ?", (request_id,))
        cur_req_status_row = await cursor.fetchone()
        cur_req_status = cur_req_status_row[0] if cur_req_status_row else 'incoming'

        cursor = await db.execute(
            "SELECT COUNT(*) FROM process_steps WHERE request_id = ? AND status = 'completed'",
            (request_id,))
        completed_count = (await cursor.fetchone())[0]
        cursor = await db.execute(
            "SELECT COUNT(*) FROM process_steps WHERE request_id = ?", (request_id,))
        total_steps = (await cursor.fetchone())[0]

        new_current = completed_count + 1
        cursor = await db.execute(
            "SELECT COUNT(*) FROM process_steps WHERE request_id = ? AND status IN ('in_progress', 'completed')",
            (request_id,))
        active_count = (await cursor.fetchone())[0]

        # Smart status recalculation â€” respect workflow statuses
        if cur_req_status in ('review', 'approval', 'completed'):
            new_status = cur_req_status  # Preserve â€” only explicit transitions change these
        elif cur_req_status == 'testing':
            new_status = 'analysis' if completed_count == total_steps else 'testing'
        elif cur_req_status == 'analysis':
            new_status = 'analysis'  # Stays until Complete Report button
        else:  # incoming/pending/in_progress â€” legacy behaviour
            new_status = 'completed' if completed_count == total_steps else ('in_progress' if active_count > 0 else 'incoming')
            # Auto-transition incoming â†’ review when custom_fields are being set
            if step_update.custom_fields and new_status == 'incoming':
                new_status = 'review'

        now = datetime.now(timezone.utc).isoformat()
        await db.execute(
            "UPDATE requests SET current_step=?, status=?, updated_at=? WHERE id=?",
            (new_current, new_status, now, request_id))
        await db.commit()

        cursor = await db.execute("SELECT * FROM requests WHERE id = ?", (request_id,))
        row = await cursor.fetchone()
        req = await _row_to_request_dict(db, row)
        return {"message": "Step updated successfully", "request": req}
    finally:
        await db.close()

@api_router.patch("/requests/{request_id}")
async def update_request(
    request_id: str, request_update: RequestCreate,
    current_user: User = Depends(require_permission('edit_request'))
):
    db = await get_db()
    try:
        cursor = await db.execute("SELECT id FROM requests WHERE id = ?", (request_id,))
        if not await cursor.fetchone():
            raise HTTPException(status_code=404, detail="Request not found")

        update_data = request_update.model_dump(exclude_unset=True)
        now = datetime.now(timezone.utc).isoformat()
        set_clauses = ["updated_at = ?"]
        values: list = [now]

        field_map = {
            'classification': 'classification', 'originator': 'originator',
            'plant': 'plant', 'device_name': 'device_name', 'lot_no': 'lot_no',
            'customer': 'customer', 'pkg_info': 'pkg_info', 'automotive': 'automotive',
            'date_ltc': 'date_ltc', 'product_hierarchy': 'product_hierarchy',
            'pdl': 'pdl', 'body_size_x': 'body_size_x', 'body_size_y': 'body_size_y',
            'package_thickness': 'package_thickness', 'ball_pitch': 'ball_pitch',
            'ball_count': 'ball_count', 'lead_pitch': 'lead_pitch', 'lead_count': 'lead_count',
            'total_ss': 'total_ss', 'purpose': 'purpose',
            'engineer_special_instruction': 'engineer_special_instruction',
            'deadline': 'deadline', 'status': 'status', 'request_number': 'request_number',
            'request_type': 'request_type',
            'note': 'note', 'retention_details': 'retention_details',
            'analysis_notes': 'analysis_notes',
            'planner_est_start': 'planner_est_start', 'planner_est_end': 'planner_est_end',
            'planner_note': 'planner_note',
        }

        for key, col in field_map.items():
            if key in update_data and update_data[key] is not None:
                set_clauses.append(f"{col} = ?")
                values.append(int(update_data[key]) if key == 'automotive' else update_data[key])

        values.append(request_id)
        await db.execute(f"UPDATE requests SET {', '.join(set_clauses)} WHERE id = ?", values)
        await db.commit()
        return {"message": "Request updated successfully", "request_id": request_id}
    finally:
        await db.close()

class PlannerEstimationUpdate(BaseModel):
    planner_est_start: Optional[str] = None
    planner_est_end: Optional[str] = None
    planner_note: Optional[str] = None

@api_router.patch("/requests/{request_id}/planner-estimation")
@api_router.put("/requests/{request_id}/planner-estimation")
async def update_planner_estimation(
    request_id: str, body: PlannerEstimationUpdate,
    current_user: User = Depends(require_role([UserRole.ADMIN, UserRole.PLANNER]))
):
    """Set or clear planner estimation dates/note for a request. Admin and Planner only."""
    db = await get_db()
    try:
        cursor = await db.execute("SELECT id FROM requests WHERE id = ?", (request_id,))
        if not await cursor.fetchone():
            raise HTTPException(status_code=404, detail="Request not found")
        now = datetime.now(timezone.utc).isoformat()
        await db.execute(
            "UPDATE requests SET planner_est_start = ?, planner_est_end = ?, planner_note = ?, updated_at = ? WHERE id = ?",
            (body.planner_est_start or None, body.planner_est_end or None,
             body.planner_note.strip() if body.planner_note else None, now, request_id)
        )
        await db.commit()
        return {"message": "Planner estimation updated", "request_id": request_id}
    finally:
        await db.close()

class NoteUpdate(BaseModel):
    note: Optional[str] = None  # empty string or None = clear the note

@api_router.patch("/requests/{request_id}/note")
async def update_request_note(
    request_id: str, body: NoteUpdate,
    current_user: User = Depends(require_permission('edit_request'))
):
    """Set or clear the public note for a request."""
    db = await get_db()
    try:
        cursor = await db.execute("SELECT id FROM requests WHERE id = ?", (request_id,))
        if not await cursor.fetchone():
            raise HTTPException(status_code=404, detail="Request not found")
        note_val = body.note.strip() if body.note else None
        now = datetime.now(timezone.utc).isoformat()
        await db.execute(
            "UPDATE requests SET note = ?, updated_at = ? WHERE id = ?",
            (note_val, now, request_id)
        )
        await db.commit()
        return {"message": "Note updated", "note": note_val}
    finally:
        await db.close()

@api_router.delete("/requests/{request_id}/note")
async def delete_request_note(
    request_id: str,
    current_user: User = Depends(require_permission('edit_request'))
):
    """Clear (delete) the public note for a request."""
    db = await get_db()
    try:
        cursor = await db.execute("SELECT id FROM requests WHERE id = ?", (request_id,))
        if not await cursor.fetchone():
            raise HTTPException(status_code=404, detail="Request not found")
        now = datetime.now(timezone.utc).isoformat()
        await db.execute("UPDATE requests SET note = NULL, updated_at = ? WHERE id = ?", (now, request_id))
        await db.commit()
        return {"message": "Note deleted"}
    finally:
        await db.close()


# ========================
# Workflow Transition Routes
# ========================

@api_router.post("/requests/{request_id}/submit-review")
async def submit_for_review(
    request_id: str,
    current_user: User = Depends(require_permission('edit_request'))
):
    """Transition a request from 'incoming'/'pending' â†’ 'review'."""
    db = await get_db()
    try:
        cursor = await db.execute("SELECT id, status FROM requests WHERE id = ?", (request_id,))
        row = await cursor.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Request not found")
        if row[1] not in ('incoming', 'pending', 'in_progress'):
            raise HTTPException(status_code=400, detail=f"Cannot submit for review from status '{row[1]}'")
        now = datetime.now(timezone.utc).isoformat()
        await db.execute("UPDATE requests SET status = 'review', updated_at = ? WHERE id = ?", (now, request_id))
        await db.commit()
        cursor = await db.execute("SELECT * FROM requests WHERE id = ?", (request_id,))
        req = await _row_to_request_dict(db, await cursor.fetchone())
        return {"message": "Request submitted for review", "request": req}
    finally:
        await db.close()


@api_router.post("/requests/{request_id}/submit-approval")
async def submit_for_approval(
    request_id: str,
    current_user: User = Depends(require_permission('edit_request'))
):
    """Transition a request from 'review' â†’ 'approval'."""
    db = await get_db()
    try:
        cursor = await db.execute("SELECT id, status FROM requests WHERE id = ?", (request_id,))
        row = await cursor.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Request not found")
        if row[1] not in ('review', 'incoming', 'pending', 'in_progress'):
            raise HTTPException(status_code=400, detail=f"Cannot submit for approval from status '{row[1]}'")
        now = datetime.now(timezone.utc).isoformat()
        await db.execute("UPDATE requests SET status = 'approval', updated_at = ? WHERE id = ?", (now, request_id))
        await db.commit()
        cursor = await db.execute("SELECT * FROM requests WHERE id = ?", (request_id,))
        req = await _row_to_request_dict(db, await cursor.fetchone())
        return {"message": "Request submitted for approval", "request": req}
    finally:
        await db.close()


@api_router.post("/requests/{request_id}/approve")
async def approve_request(
    request_id: str,
    current_user: User = Depends(get_current_user)
):
    """Planner or Admin approves a request â€” transitions 'approval' â†’ 'testing'."""
    if current_user.role not in (UserRole.ADMIN, UserRole.PLANNER):
        raise HTTPException(status_code=403, detail="Only Admin or Planner can approve requests")
    db = await get_db()
    try:
        cursor = await db.execute("SELECT id, status FROM requests WHERE id = ?", (request_id,))
        row = await cursor.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Request not found")
        if row[1] != 'approval':
            raise HTTPException(status_code=400, detail=f"Request is not pending approval (current status: '{row[1]}')")
        now = datetime.now(timezone.utc).isoformat()
        await db.execute("UPDATE requests SET status = 'testing', approved_at = ?, updated_at = ? WHERE id = ?", (now, now, request_id))
        # Auto-set the first step of every leg to 'in_progress' (In Queue)
        cursor = await db.execute(
            "SELECT DISTINCT leg FROM process_steps WHERE request_id = ?", (request_id,))
        legs = [r[0] for r in await cursor.fetchall()]
        for leg_num in legs:
            cursor = await db.execute(
                "SELECT id FROM process_steps WHERE request_id = ? AND leg = ? AND status = 'pending' ORDER BY step_number ASC LIMIT 1",
                (request_id, leg_num))
            first_step = await cursor.fetchone()
            if first_step:
                await db.execute(
                    "UPDATE process_steps SET status = 'in_progress' WHERE id = ?",
                    (first_step[0],))
        await db.commit()
        cursor = await db.execute("SELECT * FROM requests WHERE id = ?", (request_id,))
        req = await _row_to_request_dict(db, await cursor.fetchone())
        return {"message": "Request approved — testing can begin", "request": req}
    finally:
        await db.close()


@api_router.post("/requests/{request_id}/reject")
async def reject_request(
    request_id: str,
    current_user: User = Depends(get_current_user)
):
    """Planner or Admin rejects a request â€” transitions 'approval' â†’ 'review'."""
    if current_user.role not in (UserRole.ADMIN, UserRole.PLANNER):
        raise HTTPException(status_code=403, detail="Only Admin or Planner can reject requests")
    db = await get_db()
    try:
        cursor = await db.execute("SELECT id, status FROM requests WHERE id = ?", (request_id,))
        row = await cursor.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Request not found")
        if row[1] != 'approval':
            raise HTTPException(status_code=400, detail=f"Request is not pending approval (current status: '{row[1]}')")
        now = datetime.now(timezone.utc).isoformat()
        await db.execute("UPDATE requests SET status = 'review', updated_at = ? WHERE id = ?", (now, request_id))
        await db.commit()
        cursor = await db.execute("SELECT * FROM requests WHERE id = ?", (request_id,))
        req = await _row_to_request_dict(db, await cursor.fetchone())
        return {"message": "Request rejected â€” returned to review", "request": req}
    finally:
        await db.close()


class CompleteReportBody(BaseModel):
    notes: Optional[str] = None


class DiscontinueBody(BaseModel):
    reason: Optional[str] = None


@api_router.post("/requests/{request_id}/discontinue")
async def discontinue_request(
    request_id: str,
    body: DiscontinueBody = DiscontinueBody(),
    current_user: User = Depends(get_current_user)
):
    """Admin or Planner discontinues a request from any active status."""
    if current_user.role not in (UserRole.ADMIN, UserRole.PLANNER):
        raise HTTPException(status_code=403, detail="Only Admin or Planner can discontinue requests")
    db = await get_db()
    try:
        cursor = await db.execute("SELECT id, status FROM requests WHERE id = ?", (request_id,))
        row = await cursor.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Request not found")
        if row[1] in ('completed', 'discontinued'):
            raise HTTPException(status_code=400, detail=f"Cannot discontinue a request with status '{row[1]}'")
        now = datetime.now(timezone.utc).isoformat()
        await db.execute(
            "UPDATE requests SET status = 'discontinued', discontinued_at = ?, "
            "discontinued_by = ?, discontinued_reason = ?, updated_at = ? WHERE id = ?",
            (now, current_user.username, body.reason or None, now, request_id)
        )
        await db.commit()
        cursor = await db.execute("SELECT * FROM requests WHERE id = ?", (request_id,))
        req = await _row_to_request_dict(db, await cursor.fetchone())
        return {"message": "Request marked as discontinued", "request": req}
    finally:
        await db.close()


@api_router.post("/requests/{request_id}/complete-report")
async def complete_report(
    request_id: str,
    body: CompleteReportBody = CompleteReportBody(),
    current_user: User = Depends(require_permission('edit_request'))
):
    """Transition a request from 'analysis' â†’ 'completed' and save analysis notes."""
    db = await get_db()
    try:
        cursor = await db.execute("SELECT id, status FROM requests WHERE id = ?", (request_id,))
        row = await cursor.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Request not found")
        if row[1] not in ('analysis', 'testing', 'in_progress'):
            raise HTTPException(status_code=400, detail=f"Cannot complete report from status '{row[1]}'")

        # Enforce: ALL steps on ALL legs must be completed before closing the request
        cursor = await db.execute(
            "SELECT leg, step_number, step_name, status FROM process_steps "
            "WHERE request_id = ? AND status != 'completed' ORDER BY leg, step_number",
            (request_id,)
        )
        incomplete_steps = await cursor.fetchall()
        if incomplete_steps:
            items = ", ".join(
                f"Leg {s[0]} Step {s[1]} '{s[2]}' ({s[3]})" for s in incomplete_steps
            )
            raise HTTPException(
                status_code=400,
                detail=f"Cannot complete: the following steps are not yet done — {items}"
            )

        now = datetime.now(timezone.utc).isoformat()
        await db.execute(
            "UPDATE requests SET status = 'completed', analysis_notes = ?, updated_at = ? WHERE id = ?",
            (body.notes, now, request_id))
        await db.commit()
        cursor = await db.execute("SELECT * FROM requests WHERE id = ?", (request_id,))
        req = await _row_to_request_dict(db, await cursor.fetchone())
        return {"message": "Report completed", "request": req}
    finally:
        await db.close()


@api_router.delete("/requests/{request_id}")
async def delete_request(
    request_id: str,
    current_user: User = Depends(require_role([UserRole.ADMIN]))
):
    db = await get_db()
    try:
        cursor = await db.execute("SELECT id FROM requests WHERE id = ?", (request_id,))
        if not await cursor.fetchone():
            raise HTTPException(status_code=404, detail="Request not found")

        await db.execute("DELETE FROM process_steps WHERE request_id = ?", (request_id,))
        cursor = await db.execute("DELETE FROM requests WHERE id = ?", (request_id,))
        await db.commit()
        if cursor.rowcount == 0:
            raise HTTPException(status_code=404, detail="Request not found")
        return {"message": "Request deleted successfully"}
    finally:
        await db.close()


class UpdateStepsPayload(BaseModel):
    steps: List[str]  # list of step names in desired order
    leg: int = 1


@api_router.put("/requests/{request_id}/steps")
async def replace_request_steps(
    request_id: str,
    payload: UpdateStepsPayload,
    current_user: User = Depends(require_permission('manage_steps'))
):
    """Replace all process steps for a request with the given list. Preserves data for steps that
    still exist at the same position with the same name."""
    db = await get_db()
    try:
        cursor = await db.execute("SELECT id FROM requests WHERE id = ?", (request_id,))
        if not await cursor.fetchone():
            raise HTTPException(status_code=404, detail="Request not found")

        if len(payload.steps) == 0:
            raise HTTPException(status_code=400, detail="At least one step is required")

        # Fetch existing steps for this leg
        cursor = await db.execute(
            "SELECT step_number, step_name, status, started_at, completed_at, machine_no, rack_no, operator_id, "
            "tray_no, notes, attachments, custom_fields FROM process_steps "
            "WHERE request_id = ? AND leg = ? ORDER BY step_number", (request_id, payload.leg))
        old_rows = await cursor.fetchall()
        old_steps = {r[0]: {
            'step_name': r[1], 'status': r[2], 'started_at': r[3], 'completed_at': r[4],
            'machine_no': r[5], 'rack_no': r[6], 'operator_id': r[7], 'tray_no': r[8], 'notes': r[9],
            'attachments': r[10], 'custom_fields': r[11],
        } for r in old_rows}

        # Delete existing steps for this leg only
        await db.execute("DELETE FROM process_steps WHERE request_id = ? AND leg = ?", (request_id, payload.leg))

        # Insert new steps, preserving data when step_number+name match
        DEFAULT_STEP_ITEMS = {
            'Electrical Test': 'E-Test',
            'T & H Soak': None,
        }
        DEFAULT_STEP_CONDITIONS = {
            'Electrical Test': 'P4',
            'T & H Soak': None,
        }
        for i, step_name in enumerate(payload.steps):
            step_num = i + 1
            old = old_steps.get(step_num)
            if old and old['step_name'] == step_name:
                # Preserve existing data
                await db.execute(
                    """INSERT INTO process_steps (request_id, leg, step_number, step_name, status,
                       started_at, completed_at, machine_no, rack_no, operator_id, tray_no, notes, attachments, custom_fields)
                       VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                    (request_id, payload.leg, step_num, step_name, old['status'],
                     old['started_at'], old['completed_at'], old['machine_no'], old.get('rack_no'), old['operator_id'],
                     old['tray_no'], old['notes'], old['attachments'], old['custom_fields']))
            else:
                # New step with default custom_fields
                cf = {}
                if step_name in DEFAULT_STEP_ITEMS and DEFAULT_STEP_ITEMS[step_name]:
                    cf['test_item'] = DEFAULT_STEP_ITEMS[step_name]
                if step_name in DEFAULT_STEP_CONDITIONS and DEFAULT_STEP_CONDITIONS[step_name]:
                    cf['test_condition'] = DEFAULT_STEP_CONDITIONS[step_name]
                await db.execute(
                    """INSERT INTO process_steps (request_id, leg, step_number, step_name, status,
                       started_at, completed_at, machine_no, rack_no, operator_id, tray_no, notes, attachments, custom_fields)
                       VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                    (request_id, payload.leg, step_num, step_name, 'pending',
                     None, None, None, None, None, None, None, '[]', json.dumps(cf)))

        # Recalculate status
        cursor = await db.execute("SELECT status FROM requests WHERE id = ?", (request_id,))
        cur_req_status_row = await cursor.fetchone()
        cur_req_status = cur_req_status_row[0] if cur_req_status_row else 'incoming'

        cursor = await db.execute(
            "SELECT COUNT(*) FROM process_steps WHERE request_id = ? AND status = 'completed'", (request_id,))
        completed_count = (await cursor.fetchone())[0]
        cursor = await db.execute(
            "SELECT COUNT(*) FROM process_steps WHERE request_id = ?", (request_id,))
        total_steps = (await cursor.fetchone())[0]
        cursor = await db.execute(
            "SELECT COUNT(*) FROM process_steps WHERE request_id = ? AND status IN ('in_progress', 'completed')",
            (request_id,))
        active_count = (await cursor.fetchone())[0]

        # Smart status recalculation â€” respect workflow statuses
        if cur_req_status in ('review', 'approval', 'completed'):
            new_status = cur_req_status
        elif cur_req_status == 'testing':
            new_status = 'analysis' if completed_count == total_steps else 'testing'
        elif cur_req_status == 'analysis':
            new_status = 'analysis'
        else:
            new_status = 'completed' if completed_count == total_steps else ('in_progress' if active_count > 0 else 'incoming')

        now = datetime.now(timezone.utc).isoformat()
        await db.execute(
            "UPDATE requests SET current_step=?, status=?, updated_at=? WHERE id=?",
            (completed_count + 1, new_status, now, request_id))
        await db.commit()

        cursor = await db.execute("SELECT * FROM requests WHERE id = ?", (request_id,))
        row = await cursor.fetchone()
        req = await _row_to_request_dict(db, row)
        return {"message": "Steps updated successfully", "request": req}
    finally:
        await db.close()


# ========================
# LEG Management Routes
# ========================
@api_router.post("/requests/{request_id}/legs")
async def add_leg(
    request_id: str,
    current_user: User = Depends(require_permission('manage_steps'))
):
    """Add a new LEG to the request with default process steps."""
    db = await get_db()
    try:
        cursor = await db.execute("SELECT id FROM requests WHERE id = ?", (request_id,))
        if not await cursor.fetchone():
            raise HTTPException(status_code=404, detail="Request not found")

        # Find the highest leg number
        cursor = await db.execute(
            "SELECT COALESCE(MAX(leg), 0) FROM process_steps WHERE request_id = ?", (request_id,))
        max_leg = (await cursor.fetchone())[0]
        new_leg = max_leg + 1

        if new_leg > 50:
            raise HTTPException(status_code=400, detail="Maximum 50 LEGs allowed per request")

        # Get step names from LEG 1 (or use defaults)
        cursor = await db.execute(
            "SELECT step_name FROM process_steps WHERE request_id = ? AND leg = 1 ORDER BY step_number",
            (request_id,))
        leg1_rows = await cursor.fetchall()
        step_names = [r[0] for r in leg1_rows] if leg1_rows else [s["step_name"] for s in DEFAULT_STEPS]

        # Create steps for the new leg
        for i, step_name in enumerate(step_names):
            await db.execute(
                """INSERT INTO process_steps (request_id, leg, step_number, step_name, status,
                   started_at, completed_at, machine_no, rack_no, operator_id, tray_no, notes, attachments, custom_fields)
                   VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                (request_id, new_leg, i + 1, step_name, 'pending',
                 None, None, None, None, None, None, None, '[]', '{}'))

        now = datetime.now(timezone.utc).isoformat()
        await db.execute("UPDATE requests SET updated_at = ? WHERE id = ?", (now, request_id))
        await db.commit()

        cursor = await db.execute("SELECT * FROM requests WHERE id = ?", (request_id,))
        row = await cursor.fetchone()
        req = await _row_to_request_dict(db, row)
        return {"message": f"LEG {new_leg} added successfully", "request": req, "leg": new_leg}
    finally:
        await db.close()


@api_router.post("/requests/{request_id}/legs/{source_leg}/duplicate")
async def duplicate_leg(
    request_id: str,
    source_leg: int,
    current_user: User = Depends(require_permission('manage_steps'))
):
    """Duplicate a LEG: create a new LEG copying step names and all custom fields
    (test_item, test_condition, etc.) from the source leg. Status is reset to pending."""
    db = await get_db()
    try:
        cursor = await db.execute("SELECT id FROM requests WHERE id = ?", (request_id,))
        if not await cursor.fetchone():
            raise HTTPException(status_code=404, detail="Request not found")

        # Fetch source leg steps with full data
        cursor = await db.execute(
            "SELECT step_number, step_name, custom_fields FROM process_steps "
            "WHERE request_id = ? AND leg = ? ORDER BY step_number",
            (request_id, source_leg))
        source_rows = await cursor.fetchall()
        if not source_rows:
            raise HTTPException(status_code=404, detail=f"Source LEG {source_leg} not found or has no steps")

        # Determine new leg number
        cursor = await db.execute(
            "SELECT COALESCE(MAX(leg), 0) FROM process_steps WHERE request_id = ?", (request_id,))
        max_leg = (await cursor.fetchone())[0]
        new_leg = max_leg + 1

        if new_leg > 50:
            raise HTTPException(status_code=400, detail="Maximum 50 LEGs allowed per request")

        # Insert steps into new leg preserving step names and custom_fields
        for row in source_rows:
            step_num, step_name, custom_fields_raw = row[0], row[1], row[2]
            await db.execute(
                """INSERT INTO process_steps (request_id, leg, step_number, step_name, status,
                   started_at, completed_at, machine_no, rack_no, operator_id, tray_no, notes, attachments, custom_fields)
                   VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                (request_id, new_leg, step_num, step_name, 'pending',
                 None, None, None, None, None, None, None, '[]', custom_fields_raw or '{}'))

        now = datetime.now(timezone.utc).isoformat()
        await db.execute("UPDATE requests SET updated_at = ? WHERE id = ?", (now, request_id))
        await db.commit()

        cursor = await db.execute("SELECT * FROM requests WHERE id = ?", (request_id,))
        row = await cursor.fetchone()
        req = await _row_to_request_dict(db, row)
        return {"message": f"LEG {source_leg} duplicated as LEG {new_leg}", "request": req, "leg": new_leg}
    finally:
        await db.close()


@api_router.delete("/requests/{request_id}/legs/{leg_number}")
async def remove_leg(
    request_id: str, leg_number: int,
    current_user: User = Depends(require_permission('manage_steps'))
):
    """Remove a LEG and all its steps from the request."""
    db = await get_db()
    try:
        cursor = await db.execute("SELECT id FROM requests WHERE id = ?", (request_id,))
        if not await cursor.fetchone():
            raise HTTPException(status_code=404, detail="Request not found")

        # Don't allow removing last leg
        cursor = await db.execute(
            "SELECT COUNT(DISTINCT leg) FROM process_steps WHERE request_id = ?", (request_id,))
        leg_count = (await cursor.fetchone())[0]
        if leg_count <= 1:
            raise HTTPException(status_code=400, detail="Cannot remove the last LEG")

        # Check leg exists
        cursor = await db.execute(
            "SELECT COUNT(*) FROM process_steps WHERE request_id = ? AND leg = ?",
            (request_id, leg_number))
        if (await cursor.fetchone())[0] == 0:
            raise HTTPException(status_code=404, detail=f"LEG {leg_number} not found")

        # Delete the leg
        await db.execute(
            "DELETE FROM process_steps WHERE request_id = ? AND leg = ?",
            (request_id, leg_number))

        # Renumber remaining legs to maintain continuity
        cursor = await db.execute(
            "SELECT DISTINCT leg FROM process_steps WHERE request_id = ? ORDER BY leg",
            (request_id,))
        remaining_legs = [row[0] for row in await cursor.fetchall()]
        for new_num, old_num in enumerate(remaining_legs, 1):
            if new_num != old_num:
                await db.execute(
                    "UPDATE process_steps SET leg = ? WHERE request_id = ? AND leg = ?",
                    (new_num, request_id, old_num))

        # Recalculate status
        now = datetime.now(timezone.utc).isoformat()
        cursor = await db.execute("SELECT status FROM requests WHERE id = ?", (request_id,))
        cur_req_status_row = await cursor.fetchone()
        cur_req_status = cur_req_status_row[0] if cur_req_status_row else 'incoming'

        cursor = await db.execute(
            "SELECT COUNT(*) FROM process_steps WHERE request_id = ? AND status = 'completed'",
            (request_id,))
        completed_count = (await cursor.fetchone())[0]
        cursor = await db.execute(
            "SELECT COUNT(*) FROM process_steps WHERE request_id = ?", (request_id,))
        total_steps = (await cursor.fetchone())[0]
        cursor = await db.execute(
            "SELECT COUNT(*) FROM process_steps WHERE request_id = ? AND status IN ('in_progress', 'completed')",
            (request_id,))
        active_count = (await cursor.fetchone())[0]

        # Smart status recalculation â€” respect workflow statuses
        if cur_req_status in ('review', 'approval', 'completed'):
            new_status = cur_req_status
        elif cur_req_status == 'testing':
            new_status = 'analysis' if completed_count == total_steps else 'testing'
        elif cur_req_status == 'analysis':
            new_status = 'analysis'
        else:
            new_status = 'completed' if completed_count == total_steps else ('in_progress' if active_count > 0 else 'incoming')

        await db.execute(
            "UPDATE requests SET current_step=?, status=?, updated_at=? WHERE id=?",
            (completed_count + 1, new_status, now, request_id))
        await db.commit()

        cursor = await db.execute("SELECT * FROM requests WHERE id = ?", (request_id,))
        row = await cursor.fetchone()
        req = await _row_to_request_dict(db, row)
        return {"message": f"LEG {leg_number} removed successfully", "request": req}
    finally:
        await db.close()


# ========================
# Dashboard Routes
# ========================
@api_router.get("/dashboard/stats", response_model=DashboardStats)
async def get_dashboard_stats(current_user: User = Depends(get_current_user)):
    db = await get_db()
    try:
        async def cnt(st=None):
            if st:
                q = "SELECT COUNT(*) FROM requests WHERE status = ?"
                p = [st]
            else:
                q = "SELECT COUNT(*) FROM requests"
                p = []
            c = await db.execute(q, p)
            return (await c.fetchone())[0]

        total_all = await cnt()
        active = await cnt("in_progress")
        completed = await cnt("completed")
        pending = await cnt("incoming")
        review_count = await cnt("review")
        approval_count = await cnt("approval")
        testing_count = await cnt("testing")
        analysis_count = await cnt("analysis")
        active = active + testing_count  # Count both in_progress and testing as "active"
        total = total_all - completed  # Exclude completed from total (they have their own page)

        today = datetime.now(timezone.utc).isoformat()
        three_days = (datetime.now(timezone.utc) + timedelta(days=3)).isoformat()

        # Delayed requests
        dq = "SELECT * FROM requests WHERE deadline IS NOT NULL AND deadline != '' AND status != 'completed' AND deadline < ?"
        cursor = await db.execute(dq, [today])
        delayed_rows = await cursor.fetchall()
        delayed_list = [{
            "id": r[0], "request_number": r[1], "device_name": r[5] or '',
            "customer": r[7] or '', "deadline": r[23], "created_at": r[26], "status": r[28],
            "note": r['note'] if 'note' in r.keys() else None,
        } for r in delayed_rows]
        delayed_list.sort(key=lambda x: x['created_at'])

        # Upcoming deadlines
        uq = "SELECT * FROM requests WHERE deadline IS NOT NULL AND deadline != '' AND status != 'completed' AND deadline >= ? AND deadline <= ?"
        cursor = await db.execute(uq, [today, three_days])
        upcoming_rows = await cursor.fetchall()
        upcoming = len(upcoming_rows)
        upcoming_list = [{
            "id": r[0], "request_number": r[1], "device_name": r[5] or '',
            "customer": r[7] or '', "deadline": r[23], "created_at": r[26], "status": r[28],
            "note": r['note'] if 'note' in r.keys() else None,
        } for r in upcoming_rows]
        upcoming_list.sort(key=lambda x: x['deadline'])

        # Noticed requests (have a non-empty note)
        nq = "SELECT * FROM requests WHERE note IS NOT NULL AND note != '' AND status != 'completed' ORDER BY updated_at DESC"
        cursor = await db.execute(nq)
        noticed_rows = await cursor.fetchall()
        noticed_list = [{
            "id": r[0], "request_number": r[1], "device_name": r[5] or '',
            "customer": r[7] or '', "status": r[28], "deadline": r[23],
            "updated_at": r[27],
            "note": r['note'] if 'note' in r.keys() else None,
        } for r in noticed_rows]

        # Incoming Inspection steps that are pending or in_progress
        iiq = (
            "SELECT COUNT(*) FROM process_steps ps "
            "JOIN requests r ON ps.request_id = r.id "
            "WHERE LOWER(ps.step_name) = 'incoming inspection' "
            "AND ps.status IN ('pending', 'in_progress') "
            "AND r.status != 'completed'"
        )
        cursor = await db.execute(iiq)
        incoming_inspection_count = (await cursor.fetchone())[0]

        # Visual steps pending/in_progress
        cursor = await db.execute(
            "SELECT COUNT(*) FROM process_steps ps "
            "JOIN requests r ON ps.request_id = r.id "
            "WHERE LOWER(ps.step_name) = 'visual' "
            "AND ps.status IN ('pending', 'in_progress') "
            "AND r.status != 'completed'"
        )
        visual_count = (await cursor.fetchone())[0]

        # SAT steps pending/in_progress
        cursor = await db.execute(
            "SELECT COUNT(*) FROM process_steps ps "
            "JOIN requests r ON ps.request_id = r.id "
            "WHERE LOWER(ps.step_name) = 'sat' "
            "AND ps.status IN ('pending', 'in_progress') "
            "AND r.status != 'completed'"
        )
        sat_count = (await cursor.fetchone())[0]

        # Bake steps pending/in_progress
        cursor = await db.execute(
            "SELECT COUNT(*) FROM process_steps ps "
            "JOIN requests r ON ps.request_id = r.id "
            "WHERE LOWER(ps.step_name) = 'bake' "
            "AND ps.status IN ('pending', 'in_progress') "
            "AND r.status != 'completed'"
        )
        bake_count = (await cursor.fetchone())[0]

        # HTS steps pending/in_progress
        cursor = await db.execute(
            "SELECT COUNT(*) FROM process_steps ps "
            "JOIN requests r ON ps.request_id = r.id "
            "WHERE LOWER(ps.step_name) = 'hts' "
            "AND ps.status IN ('pending', 'in_progress') "
            "AND r.status != 'completed'"
        )
        hts_count = (await cursor.fetchone())[0]

        # Steps currently on Hold across all active requests
        cursor = await db.execute(
            "SELECT COUNT(*) FROM process_steps ps "
            "JOIN requests r ON ps.request_id = r.id "
            "WHERE ps.status = 'hold' "
            "AND r.status != 'completed'"
        )
        hold_count = (await cursor.fetchone())[0]

        # Requests with at least one step on Hold
        cursor = await db.execute(
            "SELECT r.id, r.request_number, r.device_name, r.customer, r.deadline, r.status, "
            "r.note, GROUP_CONCAT(ps.step_name, ', ') AS hold_steps "
            "FROM requests r "
            "JOIN process_steps ps ON ps.request_id = r.id "
            "WHERE ps.status = 'hold' AND r.status != 'completed' "
            "GROUP BY r.id "
            "ORDER BY r.updated_at DESC"
        )
        hold_req_rows = await cursor.fetchall()
        hold_requests_list = [{
            "id": row[0], "request_number": row[1], "device_name": row[2] or '',
            "customer": row[3] or '', "deadline": row[4], "status": row[5],
            "note": row[6], "hold_steps": row[7] or '',
        } for row in hold_req_rows]

        # Per-step status breakdown for Technician dashboard
        cursor = await db.execute(
            "SELECT CASE WHEN UPPER(ps.step_name) = 'SAT' THEN 'SAT' ELSE ps.step_name END AS step_name, "
            "SUM(CASE WHEN ps.status='pending' THEN 1 ELSE 0 END) AS pending, "
            "SUM(CASE WHEN ps.status='in_progress' THEN 1 ELSE 0 END) AS in_progress, "
            "SUM(CASE WHEN ps.status='completed' THEN 1 ELSE 0 END) AS completed, "
            "SUM(CASE WHEN ps.status='hold' THEN 1 ELSE 0 END) AS hold "
            "FROM process_steps ps "
            "JOIN requests r ON ps.request_id = r.id "
            "WHERE r.status != 'completed' "
            "GROUP BY CASE WHEN UPPER(ps.step_name) = 'SAT' THEN 'SAT' ELSE ps.step_name END "
            "ORDER BY MIN(ps.step_number)"
        )
        sp_rows = await cursor.fetchall()
        step_progress = [
            {"step_name": row[0], "pending": row[1], "in_progress": row[2], "completed": row[3], "hold": row[4]}
            for row in sp_rows
        ]

        # Recent activity
        rq = "SELECT * FROM requests ORDER BY updated_at DESC LIMIT 10"
        cursor = await db.execute(rq)
        recent_rows = await cursor.fetchall()
        recent_activity = [{
            "id": r[0], "request_number": r[1], "device_name": r[5] or '',
            "customer": r[7] or '', "status": r[28], "current_step": r[29],
            "updated_at": r[27], "deadline": r[23],
            "note": r['note'] if 'note' in r.keys() else None,
        } for r in recent_rows]

        # Check for backup warning based on total request count
        backup_warning = None
        backup_warning_level = None
        requires_critical_backup = False
        
        # Check backup tracking status
        cursor = await db.execute(
            "SELECT last_critical_backup_at, last_backup_request_count, critical_backup_required, last_backup_downloaded FROM backup_tracking WHERE id = 1"
        )
        backup_track = await cursor.fetchone()
        
        if total_all >= 1000:
            backup_warning = "âš ï¸ CRITICAL: Database has reached 1000+ requests! Create a backup immediately to prevent performance issues and data loss risk."
            backup_warning_level = "critical"
            
            # Check if critical backup has been completed
            if backup_track:
                last_count = backup_track[1] or 0
                already_required = backup_track[2] or 0
                was_downloaded = backup_track[3] or 0
                
                # Require backup if we've crossed 1000 threshold and haven't backed up yet
                # OR if we haven't downloaded the last backup
                if total_all >= 1000 and (last_count < 1000 or not was_downloaded or not already_required):
                    requires_critical_backup = True
                    # Update the flag in database
                    await db.execute(
                        "UPDATE backup_tracking SET critical_backup_required = 1 WHERE id = 1"
                    )
                    await db.commit()
            else:
                requires_critical_backup = True
                
        elif total_all >= 950:
            backup_warning = "âš ï¸ WARNING: Database is approaching 1000 requests (currently at {}). Please create a backup soon.".format(total_all)
            backup_warning_level = "warning"
        elif total_all >= 900:
            backup_warning = "â„¹ï¸ Info: Database has {} requests. Consider creating a backup when convenient.".format(total_all)
            backup_warning_level = "info"

        return DashboardStats(
            total_requests=total, active_requests=active,
            completed_requests=completed, pending_requests=pending,
            ongoing_requests=active, delayed_requests=len(delayed_list),
            upcoming_deadline_requests=upcoming,
            review_requests=review_count,
            approval_requests=approval_count,
            testing_requests=testing_count,
            analysis_requests=analysis_count,
            incoming_inspection_count=incoming_inspection_count,
            visual_count=visual_count,
            sat_count=sat_count,
            bake_count=bake_count,
            hts_count=hts_count,
            hold_count=hold_count,
            hold_requests_list=hold_requests_list,
            step_progress=step_progress,
            recent_activity=recent_activity, delayed_requests_list=delayed_list,
            upcoming_deadline_list=upcoming_list,
            noticed_requests_list=noticed_list,
            backup_warning=backup_warning,
            backup_warning_level=backup_warning_level,
            requires_critical_backup=requires_critical_backup,
        )
    finally:
        await db.close()


@api_router.get("/")
async def root():
    return {"message": "Rel Request Process Flow API"}

@api_router.get("/step-names")
async def get_available_step_names():
    """Return the list of available process step names that can be used when creating requests."""
    return {"step_names": AVAILABLE_STEP_NAMES}


# ========================
# Employee & Machine Directory
# ========================
EMPLOYEES = [
    {"id": "947241", "name": "Celia Corpuz", "position": "Manager"},
    {"id": "105445", "name": "Conrado Hidalgo", "position": "Sr. FA Engr"},
    {"id": "240097", "name": "Pamela Satur", "position": "Rel Engr"},
    {"id": "240167", "name": "Shelah Mae Perez", "position": "Rel Engr"},
    {"id": "240168", "name": "Clarence Joshua Ramirez", "position": "FA Engr"},
    {"id": "250296", "name": "Allyza Nicole Humirang", "position": "Rel Engr"},
    {"id": "960853", "name": "Loreta Veran", "position": "Sr. Rel Engr"},
    {"id": "993404", "name": "Lea Dalanon", "position": "FA Operation Engr"},
    {"id": "982308", "name": "Esmeria, Erwin", "position": "FA ES P3"},
    {"id": "175081", "name": "Hatulan, Irving", "position": "FA ES P3"},
    {"id": "175075", "name": "Delos Santos, Charito", "position": "FA ES P3"},
    {"id": "105294", "name": "Bermiso, Ricky", "position": "FA ES P3"},
    {"id": "240427", "name": "Monterosa, Shaira", "position": "FA ES P3"},
    {"id": "175083", "name": "Supapo, Bryane", "position": "FA ES P3"},
    {"id": "175087", "name": "Ortiz, Van Joven", "position": "FA ES P3"},
    {"id": "175198", "name": "Del Rosario, Wowie", "position": "FA ES P3"},
    {"id": "175082", "name": "Foronda, Georjan", "position": "FA ES P3"},
    {"id": "202544", "name": "Salazar, Jeronel", "position": "FA ES P3"},
    {"id": "250125", "name": "Dela Rosa, Rowell", "position": "FA ES P3"},
    {"id": "250158", "name": "Remigio, Alcen", "position": "FA ES P3"},
    {"id": "250135", "name": "Trinidad, Maricel", "position": "REL ES"},
    {"id": "155253", "name": "Delos Santos, Chlarissa", "position": "REL ES"},
    {"id": "145087", "name": "Santiago, Kimberly Rose", "position": "REL ES"},
    {"id": "155252", "name": "De Mesa, Rosemarie", "position": "REL ES"},
    {"id": "175088", "name": "Velitario, Madelyn", "position": "REL ES"},
    {"id": "145084", "name": "Arcega, Johnrey", "position": "REL ES"},
    {"id": "155420", "name": "Reig, Leonito", "position": "REL ES"},
    {"id": "175089", "name": "Barrera, Marissa", "position": "REL ES"},
    {"id": "175074", "name": "Cruz, Jasthine Mae", "position": "REL ES"},
    {"id": "230076", "name": "Rizano, Jan Mark", "position": "REL ES"},
    {"id": "250136", "name": "Semillano, Adrian", "position": "REL ES"},
    {"id": "252523", "name": "Balcita, Jeriel", "position": "REL ES"},
    {"id": "981931", "name": "Reggie Quito", "position": "REL ES"},
    {"id": "155389", "name": "Roy Tiquis", "position": "REL ES"},
    {"id": "180966", "name": "Eduardo Visca", "position": "REL ES"},
]

class EmployeeCreate(BaseModel):
    id: str
    name: str
    position: str = ""

class MachineCreate(BaseModel):
    machine_no: str
    description: str

@api_router.get("/employees")
async def get_employees():
    """Return the employee directory from the DB."""
    db = await get_db()
    try:
        cursor = await db.execute("SELECT id, name, position FROM employees ORDER BY name")
        rows = await cursor.fetchall()
        return {"employees": [{"id": r[0], "name": r[1], "position": r[2]} for r in rows]}
    finally:
        await db.close()

@api_router.post("/employees", status_code=201)
async def add_employee(data: EmployeeCreate, current_user: User = Depends(require_permission('manage_settings'))):
    db = await get_db()
    try:
        cursor = await db.execute("SELECT id FROM employees WHERE id = ?", (data.id.strip(),))
        if await cursor.fetchone():
            raise HTTPException(status_code=400, detail="Employee ID already exists")
        await db.execute("INSERT INTO employees (id, name, position) VALUES (?,?,?)",
                         (data.id.strip(), data.name.strip(), data.position.strip()))
        await db.commit()
        return {"id": data.id.strip(), "name": data.name.strip(), "position": data.position.strip()}
    finally:
        await db.close()

@api_router.delete("/employees/{emp_id}", status_code=204)
async def delete_employee(emp_id: str, current_user: User = Depends(require_permission('manage_settings'))):
    db = await get_db()
    try:
        await db.execute("DELETE FROM employees WHERE id = ?", (emp_id,))
        await db.commit()
    finally:
        await db.close()

@api_router.get("/machines")
async def get_machines():
    """Return the machine list from the DB."""
    db = await get_db()
    try:
        cursor = await db.execute("SELECT id, machine_no, description FROM machines ORDER BY machine_no")
        rows = await cursor.fetchall()
        return {"machines": [{"id": r[0], "machine_no": r[1], "description": r[2]} for r in rows]}
    finally:
        await db.close()

@api_router.post("/machines", status_code=201)
async def add_machine(data: MachineCreate, current_user: User = Depends(require_permission('manage_settings'))):
    db = await get_db()
    try:
        cursor = await db.execute(
            "INSERT INTO machines (machine_no, description) VALUES (?,?)",
            (data.machine_no.strip().upper(), data.description.strip().upper())
        )
        await db.commit()
        return {"id": cursor.lastrowid, "machine_no": data.machine_no.strip().upper(), "description": data.description.strip().upper()}
    finally:
        await db.close()

@api_router.delete("/machines/{machine_id}", status_code=204)
async def delete_machine(machine_id: int, current_user: User = Depends(require_permission('manage_settings'))):
    db = await get_db()
    try:
        await db.execute("DELETE FROM machines WHERE id = ?", (machine_id,))
        await db.commit()
    finally:
        await db.close()


# ========================
# Settings Routes
# ========================
@api_router.get("/public/stats")
async def public_stats():
    """Public endpoint — returns approved user count (no auth required)."""
    db = await get_db()
    try:
        cursor = await db.execute("SELECT COUNT(*) FROM users WHERE approved = 1")
        row = await cursor.fetchone()
        return {"approved_users": row[0] if row else 0}
    finally:
        await db.close()

@api_router.post("/verify-tech-code")
async def verify_tech_code(data: dict):
    code = data.get("code", "")
    db = await get_db()
    try:
        cursor = await db.execute("SELECT tech_auth_code FROM settings WHERE id = 1")
        row = await cursor.fetchone()
        stored = row[0] if row and row[0] else "735522"
        return {"valid": code == stored}
    finally:
        await db.close()

@api_router.get("/settings", response_model=AppSettings)
async def get_settings(current_user: User = Depends(get_current_user)):
    db = await get_db()
    try:
        cursor = await db.execute("SELECT * FROM settings WHERE id = 1")
        row = await cursor.fetchone()
        if not row:
            settings = AppSettings()
        else:
            settings = AppSettings(
                app_name=row[1] or "Rel Request Process Flow",
                app_logo=row[2], company_name=row[3], contact_email=row[4],
                process_steps=json.loads(row[5]) if row[5] else [s["step_name"] for s in DEFAULT_STEPS],
                process_presets=json.loads(row[9]) if len(row) > 9 and row[9] else None,
                custom_fields=json.loads(row[6]) if row[6] else {},
                tech_auth_code=row[8] if len(row) > 8 and row[8] else "735522",
                updated_at=datetime.fromisoformat(row[7]) if row[7] else datetime.now(timezone.utc),
            )
        # Inject Electrical Test items and conditions into custom_fields for frontend
        if "electrical_test_items" not in settings.custom_fields:
            settings.custom_fields["electrical_test_items"] = ELECTRICAL_TEST_ITEMS
        if "electrical_test_conditions" not in settings.custom_fields:
            settings.custom_fields["electrical_test_conditions"] = ELECTRICAL_TEST_CONDITIONS
        return settings
    finally:
        await db.close()

@api_router.patch("/settings")
async def update_settings(
    settings_update: SettingsUpdate,
    current_user: User = Depends(require_permission('manage_settings'))
):
    db = await get_db()
    try:
        update_data = settings_update.model_dump(exclude_unset=True)
        now = datetime.now(timezone.utc).isoformat()

        cursor = await db.execute("SELECT id FROM settings WHERE id = 1")
        exists = await cursor.fetchone()

        if exists:
            parts = ["updated_at = ?"]
            vals: list = [now]
            for k in ['app_name', 'app_logo', 'company_name', 'contact_email', 'tech_auth_code']:
                if k in update_data:
                    parts.append(f"{k} = ?"); vals.append(update_data[k])
            if 'process_steps' in update_data:
                parts.append("process_steps = ?"); vals.append(json.dumps(update_data['process_steps']))
            if 'process_presets' in update_data:
                parts.append("process_presets = ?"); vals.append(json.dumps(update_data['process_presets']))
            if 'custom_fields' in update_data:
                parts.append("custom_fields = ?"); vals.append(json.dumps(update_data['custom_fields']))
            await db.execute(f"UPDATE settings SET {', '.join(parts)} WHERE id = 1", vals)
        else:
            await db.execute(
                "INSERT INTO settings (id,app_name,app_logo,company_name,contact_email,custom_fields,tech_auth_code,updated_at) VALUES (1,?,?,?,?,?,?,?)",
                (update_data.get('app_name', 'Rel Request Process Flow'),
                 update_data.get('app_logo'), update_data.get('company_name'),
                 update_data.get('contact_email'),
                 json.dumps(update_data.get('custom_fields', {})),
                 update_data.get('tech_auth_code', '735522'), now))

        await db.commit()
        return {"message": "Settings updated successfully"}
    finally:
        await db.close()


# ========================
# Excel Import Routes - Reliability Test Request Sheets
# ========================
# Cell mapping for the standard "Reliability Test Request Sheet" form layout.
# Each file is one request. Data lives in fixed cells on Sheet1.

EXCEL_CELL_MAP = {
    # General Information (left column)
    'request_number': 'D8',       # e.g. 'RR00142272'
    'rrs_number':     'F8',       # e.g. 'RRS# 220260103'
    'classification': 'D9',       # e.g. 'E : Engineering Evaluation'
    'originator':     'D10',      # e.g. 'Paulyn Ysabel Aquino (162298)'
    'plant':          'D11',      # e.g. 'ATP|P3'
    'device_name':    'D12',      # e.g. 'MLX81300EBF_3'
    'lot_no':         'D13',      # e.g. 10367075
    'customer':       'D14',      # e.g. 'MELEXIS TECHNOLOGIES NV (385)'
    'pkg_info':       'D15',      # e.g. 'Mold'
    'automotive':     'D16',      # 'Yes' / 'No'
    'date_ltc':       'D17',      # datetime
    'purpose':        'D18',      # text
    # General Information (right column)
    'product_hierarchy': 'O8',    # e.g. 'Leadframe|QFN|MLF SAW'
    'pdl':               'O9',    # e.g. 'XB  05  032'
    'body_size_x':       'O10',   # number
    'body_size_y':       'O11',   # number
    'package_thickness': 'O12',   # number
    'ball_pitch':        'O13',   # number or '-'
    'ball_count':        'O14',   # number or '-'
    'lead_pitch':        'O15',   # number or '-'
    'lead_count':        'O16',   # number
    'total_ss':          'O17',   # text
}

# Engineer special instruction spans B21 and B22 (merged cells with multi-line text)
ENGINEER_INSTR_CELLS = ['B21', 'B22']


# ========================
# Word (.docx) Import Parser - Reliability Test Traveller
# ========================

def _parse_word_request_sheet(contents: bytes) -> dict:
    """Parse a Word Reliability Test Traveller (.docx) into a request data dict.

    Expected layout (3 tables):
      Table 0 â€“ General Information  (row-label / value pairs, 4 columns)
      Table 1 â€“ Rel Test Traveller   (test steps â€“ not imported)
      Table 2 â€“ Comments / Special Instructions
    """
    _WNS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    def _cell_text(tc):
        return ''.join(t.text or '' for t in tc.findall('.//w:t', _WNS)).strip()

    def _table_rows(tbl):
        return [[_cell_text(tc) for tc in tr.findall('w:tc', _WNS)]
                for tr in tbl.findall('w:tr', _WNS)]

    try:
        with zipfile.ZipFile(io.BytesIO(contents)) as z:
            xml_data = z.read('word/document.xml').decode('utf-8')
    except Exception as exc:
        return {'data': {}, 'errors': [f'Cannot read .docx file: {exc}']}

    try:
        root = _ET.fromstring(xml_data)
    except Exception as exc:
        return {'data': {}, 'errors': [f'Cannot parse .docx XML: {exc}']}

    body = root.find('.//w:body', _WNS)
    if body is None:
        return {'data': {}, 'errors': ['Invalid .docx: no body element found']}

    tables = body.findall('.//w:tbl', _WNS)
    if not tables:
        return {'data': {}, 'errors': ['No tables found â€“ is this a Reliability Test Traveller document?']}

    data = {}
    parse_errors = []

    # â”€â”€ Table 0: General Information â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    rows_t0 = _table_rows(tables[0])
    for row in rows_t0:
        # Each data row: [left_label, left_value, right_label, right_value]
        left_lbl  = (row[0].lower() if len(row) > 0 else '').strip()
        left_val  = row[1].strip()  if len(row) > 1 else ''
        right_lbl = (row[2].lower() if len(row) > 2 else '').strip()
        right_val = row[3].strip()  if len(row) > 3 else ''

        # â”€â”€ left column â”€â”€
        if 'report number' in left_lbl and left_val:
            data['request_number'] = f'RMS{left_val}'
        if 'classification' in left_lbl and left_val:
            data['classification'] = left_val
        if 'mfg site' in left_lbl and left_val:
            data['plant'] = left_val
        if left_lbl == 'customer' and left_val:
            data['customer'] = left_val
        if 'device#' in left_lbl and left_val:
            data['device_name'] = left_val
        if 'lot#' in left_lbl and left_val:
            data['lot_no'] = left_val

        # â”€â”€ right column â”€â”€
        if 'pkg type' in right_lbl and right_val:
            data['pkg_info'] = right_val
        if 'pkg size' in right_lbl and right_val:
            m = _re.match(r'([\d.]+)\s*[xXÃ—]\s*([\d.]+)', right_val)
            if m:
                try:
                    data['body_size_x'] = float(m.group(1))
                    data['body_size_y'] = float(m.group(2))
                except ValueError:
                    parse_errors.append(f"Could not parse Pkg Size: '{right_val}'")
        if 'pkg thick' in right_lbl and right_val:
            m = _re.match(r'([\d.]+)', right_val)
            if m:
                try:
                    data['package_thickness'] = float(m.group(1))
                except ValueError:
                    parse_errors.append(f"Could not parse Pkg Thick: '{right_val}'")
        if right_lbl == 'lead' and right_val:
            try:
                data['lead_count'] = int(float(right_val))
            except ValueError:
                parse_errors.append(f"Could not parse Lead count: '{right_val}'")
        if 'lead pitch' in right_lbl and right_val:
            s = right_val
            if s.lower() not in ('n/a', '-', ''):
                m = _re.match(r'([\d.]+)', s)
                if m:
                    try:
                        data['lead_pitch'] = float(m.group(1))
                    except ValueError:
                        parse_errors.append(f"Could not parse Lead Pitch: '{right_val}'")
        if 'submission date' in right_lbl and right_val:
            s = right_val.strip()
            for fmt in ('%m/%d/%y', '%m/%d/%Y', '%Y-%m-%d'):
                try:
                    dt = datetime.strptime(s, fmt)
                    data['date_ltc'] = dt.strftime('%Y-%m-%d')
                    break
                except ValueError:
                    continue
            else:
                data['date_ltc'] = s
        if 'total s/s' in right_lbl and right_val:
            data['total_ss'] = right_val

    # â”€â”€ Material info (EMC, LF, Die Attach, Wire) â†’ pkg notes â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    material_parts: list = []
    for row in rows_t0:
        left_lbl  = (row[0].lower() if row else '').strip()
        left_val  = row[1].strip()  if len(row) > 1 else ''
        right_lbl = (row[2].lower() if len(row) > 2 else '').strip()
        right_val = row[3].strip()  if len(row) > 3 else ''

        if left_lbl == 'emc' and left_val:
            material_parts.append(f'EMC: {left_val}')
        if 'lf material' in left_lbl and left_val:
            material_parts.append(f'LF: {left_val}')
        if 'die attach' in right_lbl and right_val:
            material_parts.append(f'Die Attach: {right_val}')
        if 'wire' in right_lbl and right_val:
            material_parts.append(f'Wire: {right_val}')
            if 'automotive' in right_val.lower():
                data['automotive'] = True

    # â”€â”€ Table 2: Comments / Special Instructions â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    # ── Leg Traveller Tables ────────────────────────────────────────────────────────
    # Tables[1:-1] are Rel Test Traveller tables (one per leg).
    # The last table (tables[-1], when >=3 tables) is Comments.
    TRAVELLER_RE = _re.compile(r'Rel\s+Test\s+Trav', _re.IGNORECASE)
    LEG_RE_TBL   = _re.compile(r'\bLEG\s*(\d+)\b', _re.IGNORECASE)

    TRAV_HEADER_NAMES = {
        'test item':          'item',
        'test condition':     'condition',
        'qty in':             'qty_in',
        'qty\nin':            'qty_in',
        'date and time in':   'time_in',
        'date and\ntime in':  'time_in',
        'date / time in':     'time_in',
        'date/time in':       'time_in',
        'date and time out':  'time_out',
        'date and\ntime out': 'time_out',
        'date / time out':    'time_out',
        'date/time out':      'time_out',
        'operator':           'operator',
        'machine':            'machine_no',
        'machine no':         'machine_no',
        'machine no.':        'machine_no',
        'result qty out':     'qty_out',
        'result\nqty out':    'qty_out',
        'qty out':            'qty_out',
        'qty out\nresult':    'qty_out',
        'tray':               'tray_no',
        'tray no':            'tray_no',
        'tray no.':           'tray_no',
    }

    leg_data = []
    leg_counter = 0

    # Middle tables are potential leg travellers; last table is Comments
    mid_tables = tables[1:-1] if len(tables) > 2 else []

    for tbl in mid_tables:
        tbl_rows = _table_rows(tbl)
        if not tbl_rows:
            continue

        # Check if the first few rows contain "Rel Test Traveller"
        trav_found = False
        trav_leg_num = None
        for ri, row in enumerate(tbl_rows[:4]):
            rt = ' '.join(row)
            if TRAVELLER_RE.search(rt):
                trav_found = True
                m = LEG_RE_TBL.search(rt)
                if m:
                    trav_leg_num = int(m.group(1))
                # Scan individual cells too (title may span columns)
                if trav_leg_num is None:
                    for cell in row:
                        m2 = LEG_RE_TBL.search(cell)
                        if m2:
                            trav_leg_num = int(m2.group(1))
                            break
                break

        if not trav_found:
            continue

        if trav_leg_num is None:
            leg_counter += 1
            trav_leg_num = leg_counter
        else:
            leg_counter = trav_leg_num

        # Find the column-header row (contains "Test Item" or similar)
        col_map = {}
        data_start = 2
        for ri, row in enumerate(tbl_rows[1:6], start=1):
            candidate = {}
            for ci, cell_val in enumerate(row):
                s = cell_val.lower().strip().replace('\n', ' ')
                for hname, hkey in TRAV_HEADER_NAMES.items():
                    if hname.replace('\n', ' ') == s and hkey not in candidate:
                        candidate[hkey] = ci
            if 'item' in candidate:
                col_map = candidate
                data_start = ri + 1
                break

        if 'item' not in col_map:
            continue

        # Capture col_map in closures with default argument
        def _sv_trav(row, key, _cm=col_map):
            ci = _cm.get(key)
            if ci is None or ci >= len(row):
                return None
            s = row[ci].strip()
            return s if s else None

        def _int_sv_trav(row, key, _cm=col_map):
            ci = _cm.get(key)
            if ci is None or ci >= len(row):
                return None
            v = row[ci].strip()
            if not v:
                return None
            try:
                return int(float(v))
            except (ValueError, TypeError):
                return None

        items_trav = []
        for row in tbl_rows[data_start:]:
            if not any(c.strip() for c in row):
                continue
            step_name = _sv_trav(row, 'item')
            if not step_name:
                continue
            items_trav.append({
                'step_name':      step_name,
                'test_condition': _sv_trav(row, 'condition'),
                'qty_in':         _int_sv_trav(row, 'qty_in'),
                'time_in':        None,
                'time_out':       None,
                'operator':       _sv_trav(row, 'operator'),
                'machine_no':     _sv_trav(row, 'machine_no'),
                'qty_out_result': _sv_trav(row, 'qty_out'),
                'tray_no':        _sv_trav(row, 'tray_no'),
            })

        if items_trav:
            leg_data.append({'leg': trav_leg_num, 'items': items_trav})

    # ── Last table: Comments / Special Instructions ──────────────────────────────────
    # Use tables[-1] when >=3 tables (supports variable number of leg tables).
    comments = ''
    if len(tables) >= 3:
        rows_last = _table_rows(tables[-1])
        comment_parts = []
        for row in rows_last[1:]:   # skip header row
            text = ' '.join(c for c in row if c).strip()
            if text:
                comment_parts.append(text)
        comments = '\n'.join(comment_parts)
        if 'automotive' in comments.lower():
            data['automotive'] = True

    # engineer_special_instruction ← comments only (last table)
    if comments:
        data['engineer_special_instruction'] = comments

    # purpose ← material information section
    if material_parts:
        data['purpose'] = ('Material Information\n'
                           + '\n'.join(material_parts))

    if 'automotive' not in data:
        data['automotive'] = False

    return {'data': data, 'errors': parse_errors, 'leg_data': leg_data}


def _parse_excel_request_sheet(wb) -> dict:
    """Parse a Reliability Test Request Sheet workbook into a request data dict."""
    ws = wb.active or wb.worksheets[0]

    data = {}
    parse_errors = []

    for key, cell_ref in EXCEL_CELL_MAP.items():
        val = ws[cell_ref].value
        if val is None:
            continue

        # Convert value based on field
        try:
            if key == 'automotive':
                data[key] = str(val).strip().lower() in ('yes', 'true', '1', 'y')
            elif key == 'date_ltc':
                if hasattr(val, 'strftime'):
                    data[key] = val.strftime('%Y-%m-%d')
                else:
                    data[key] = str(val).strip()
            elif key in ('body_size_x', 'body_size_y', 'package_thickness', 'ball_pitch', 'lead_pitch'):
                s = str(val).strip()
                if s and s != '-':
                    data[key] = float(s)
            elif key in ('ball_count', 'lead_count'):
                s = str(val).strip()
                if s and s != '-':
                    data[key] = int(float(s))
            elif key == 'lot_no':
                data[key] = str(val).strip()
            else:
                s = str(val).strip()
                if s:
                    data[key] = s
        except (ValueError, TypeError) as e:
            parse_errors.append(f"Cell {cell_ref} ({key}): could not parse '{val}'")

    # Engineer special instruction â€” combine B21 + B22
    instr_parts = []
    for cell_ref in ENGINEER_INSTR_CELLS:
        val = ws[cell_ref].value
        if val and str(val).strip():
            instr_parts.append(str(val).strip())
    if instr_parts:
        data['engineer_special_instruction'] = '\n'.join(instr_parts)

    return {'data': data, 'errors': parse_errors}


def _parse_excel_leg_travellers(ws) -> list:
    """Parse 'Rel Test Traveller' leg sections from a Reliability Request Sheet.

    Supports two header formats:
      - Old: single cell "Rel Test Traveller – LEG N"  (all in one merged cell)
      - New: "Rel Test Traveller - REL" in one cell AND "LEG N" in another cell
             of the same row (columns may differ)

    Column positions are discovered dynamically from the 'Test Item' header row
    that immediately follows each leg header, so the function works regardless
    of how columns are laid out in different template versions.

    Returns list of:
      { 'leg': int, 'items': [{ step_name, test_condition, qty_in,
                                time_in, time_out, operator,
                                machine_no, qty_out_result, tray_no }] }
    """
    # Match full 'Traveller' AND truncated variants like 'Travel-' that appear
    # in some Excel templates where the cell text is cut short.
    TRAVELLER_RE = _re.compile(r'Rel\s+Test\s+Travel', _re.IGNORECASE)
    LEG_RE       = _re.compile(r'\bLEG\s*(\d+)\b', _re.IGNORECASE)

    # Map normalised cell text → field key
    HEADER_NAMES = {
        'test item':           'item',
        'test condition':      'condition',
        'qty in':              'qty_in',
        'date / time in':      'time_in',
        'date/time in':        'time_in',
        'date / time out':     'time_out',
        'date/time out':       'time_out',
        'operator':            'operator',
        'machine no':          'machine_no',
        'machine no.':         'machine_no',
        'qty out':             'qty_out',
        'qty out result':      'qty_out',
        'qty out\nresult':     'qty_out',
        'tray no':             'tray_no',
        'tray no.':            'tray_no',
    }

    rows = list(ws.iter_rows(values_only=True))

    # ── Pass 1: locate every leg-header row and its leg number ─────────────
    leg_header_rows = []   # [(row_idx, leg_num), …]
    for i, row in enumerate(rows):
        row_strs = [str(c).strip() if c is not None else '' for c in row]
        row_text = ' '.join(row_strs)
        if not TRAVELLER_RE.search(row_text):
            continue
        m = LEG_RE.search(row_text)
        if m:
            leg_header_rows.append((i, int(m.group(1))))

    if not leg_header_rows:
        return []

    # ── Pass 2: for each leg discover columns then collect data rows ───────
    result = []

    for hi, (hdr_idx, leg_num) in enumerate(leg_header_rows):
        end_idx = leg_header_rows[hi + 1][0] if hi + 1 < len(leg_header_rows) else len(rows)

        # Find the "Test Item" column-header row (usually next row, ≤3 rows away)
        col_map = {}      # field key → column index
        data_start = hdr_idx + 2   # default: skip 1 header row

        for ri in range(hdr_idx + 1, min(hdr_idx + 5, end_idx)):
            candidate = {}
            for ci, cell in enumerate(rows[ri]):
                s = str(cell).strip().lower().replace('\n', ' ') if cell is not None else ''
                for hname, hkey in HEADER_NAMES.items():
                    if hname == s and hkey not in candidate:
                        candidate[hkey] = ci
            if 'item' in candidate:
                col_map = candidate
                data_start = ri + 1
                break

        if 'item' not in col_map:
            continue   # no recognisable header — skip this leg

        # ── Collect data rows ─────────────────────────────────────────────
        def _sv(row, key):
            ci = col_map.get(key)
            if ci is None or ci >= len(row) or row[ci] is None:
                return None
            s = str(row[ci]).strip()
            return s if s else None

        def _int_v(row, key):
            ci = col_map.get(key)
            if ci is None or ci >= len(row) or row[ci] is None:
                return None
            try:
                return int(float(str(row[ci])))
            except (ValueError, TypeError):
                return None

        def _dt_v(row, key):
            ci = col_map.get(key)
            if ci is None or ci >= len(row) or row[ci] is None:
                return None
            v = row[ci]
            if hasattr(v, 'isoformat'):
                return v.isoformat()
            s = str(v).strip()
            if not s:
                return None
            try:
                datetime.fromisoformat(s)
                return s
            except (ValueError, TypeError):
                return None  # Ignore non-datetime text in date columns

        items = []
        for ri in range(data_start, end_idx):
            row = rows[ri]
            if all(c is None for c in row):
                continue
            step_name = _sv(row, 'item')
            if not step_name:
                continue
            items.append({
                'step_name':      step_name,
                'test_condition': _sv(row, 'condition'),
                'qty_in':         _int_v(row, 'qty_in'),
                'time_in':        _dt_v(row, 'time_in'),
                'time_out':       _dt_v(row, 'time_out'),
                'operator':       _sv(row, 'operator'),
                'machine_no':     _sv(row, 'machine_no'),
                'qty_out_result': _sv(row, 'qty_out'),
                'tray_no':        _sv(row, 'tray_no'),
            })

        if items:
            result.append({'leg': leg_num, 'items': items})

    return result


def _parse_excel_test_matrix(ws) -> list:
    """Parse the Test Matrix section of a Reliability Request Sheet.

    Scans for the header row that contains 'Leg' and 'Test Type' and 'Test Items',
    then collects every data row until a fully-empty row is reached.

    Returns list of:
      { 'leg': int, 'items': [{ step_name, test_condition, qty_in, time_in,
                                time_out, operator, machine_no, qty_out_result,
                                tray_no }] }

    This is used as a fallback / supplement to _parse_excel_leg_travellers so
    that every leg found in the matrix is always represented, even when a
    Traveller section is missing or has an unrecognised header text.
    """
    LEG_RE = _re.compile(r'LEG\s*(\d+)', _re.IGNORECASE)

    rows = list(ws.iter_rows(values_only=True))

    # ── Find the Test Matrix header row ───────────────────────────────────
    header_idx = None
    col_leg = col_test_type = col_test_items = col_condition = col_reading = None

    for i, row in enumerate(rows):
        cells_lower = [str(c).strip().lower() if c is not None else '' for c in row]
        row_text = ' '.join(cells_lower)
        if 'leg' in cells_lower and 'test type' in row_text and 'test items' in row_text:
            header_idx = i
            for j, s in enumerate(cells_lower):
                if s == 'leg' and col_leg is None:
                    col_leg = j
                elif s == 'test type' and col_test_type is None:
                    col_test_type = j
                elif s == 'test items' and col_test_items is None:
                    col_test_items = j
                elif s == 'test condition' and col_condition is None:
                    col_condition = j
                elif ('reflow' in s or 'reading' in s) and col_reading is None:
                    col_reading = j
            break

    if header_idx is None or col_leg is None:
        return []

    # ── Collect data rows grouped by leg ──────────────────────────────────
    legs_dict = {}   # leg_num (int) -> [item, ...]
    current_leg = None

    for row in rows[header_idx + 1:]:
        if all(c is None for c in row):
            break

        def _sv(col):
            if col is None or col >= len(row) or row[col] is None:
                return None
            s = str(row[col]).strip()
            return s if s else None

        cell_leg = _sv(col_leg)
        if cell_leg:
            m = LEG_RE.match(cell_leg)
            if m:
                current_leg = int(m.group(1))
            # If cell_leg text doesn't look like a leg, keep current_leg

        if current_leg is None:
            continue

        test_type  = _sv(col_test_type)
        test_items = _sv(col_test_items)
        condition  = _sv(col_condition)
        reading    = _sv(col_reading)

        if not test_items:
            continue

        full_condition = ' | '.join(filter(None, [condition, reading]))
        if current_leg not in legs_dict:
            legs_dict[current_leg] = []
        legs_dict[current_leg].append({
            'step_name':      test_items,
            'test_condition': full_condition or None,
            'qty_in':         None,
            'time_in':        None,
            'time_out':       None,
            'operator':       None,
            'machine_no':     None,
            'qty_out_result': None,
            'tray_no':        None,
        })

    return [{'leg': k, 'items': v} for k, v in sorted(legs_dict.items())]


def _clean_dt(v):
    """Return v only if it is a valid ISO datetime/date string, otherwise None.
    Prevents non-datetime text (e.g. request numbers) from being stored in date columns."""
    if v is None:
        return None
    s = str(v).strip()
    if not s:
        return None
    try:
        datetime.fromisoformat(s)
        return s
    except (ValueError, TypeError):
        return None


@api_router.post("/requests/import")
async def import_requests_from_excel(
    files: List[UploadFile] = File(...),
    current_user: User = Depends(require_permission('import_requests'))
):
    """Import requests from Reliability Test Request Sheet Excel files.
    Each .xlsx file is one request. Multiple files can be uploaded at once."""

    results_created = []
    results_errors = []

    for file in files:
        fname = file.filename or 'unknown'

        if not fname.lower().endswith(('.xlsx', '.xls')):
            results_errors.append({'file': fname, 'errors': ['Not an Excel file (.xlsx)']})
            continue

        db = None
        try:
            contents = await file.read()
            if fname.lower().endswith('.xls'):
                if not _XLRD_AVAILABLE:
                    results_errors.append({'file': fname, 'errors': ['xlrd is required to read .xls files. Please install it.']})
                    continue
                book = _xlrd.open_workbook(file_contents=contents)
                wb = _XlrdWb(book)
            else:
                wb = load_workbook(filename=io.BytesIO(contents), data_only=True)

            ws_active = wb.active or wb.worksheets[0]
            parsed = _parse_excel_request_sheet(wb)
            request_data = parsed['data']
            file_errors = parsed['errors']

            # Parse leg traveller sections (detailed procedural steps per leg)
            leg_data = _parse_excel_leg_travellers(ws_active)

            # Also parse Test Matrix to discover ALL legs by number.
            # Any leg present in the matrix but missing from traveller data
            # (e.g. truncated/variant header text) is added using matrix items.
            matrix_legs = _parse_excel_test_matrix(ws_active)
            traveller_leg_nums = {lg['leg'] for lg in leg_data}
            for ml in matrix_legs:
                if ml['leg'] not in traveller_leg_nums:
                    leg_data.append(ml)
            leg_data.sort(key=lambda x: x['leg'])

            if not request_data or not any(request_data.get(k) for k in ('device_name', 'lot_no', 'customer', 'request_number')):
                file_errors.append('No recognizable data found in the file. Is this a Reliability Test Request Sheet?')
                results_errors.append({'file': fname, 'errors': file_errors})
                continue

            # Always auto-generate REL# — store the old RR# for reference
            original_rr = request_data.pop('request_number', None)
            rrs_number = request_data.pop('rrs_number', None)

            db = await get_db()
            year = datetime.now(timezone.utc).year
            cursor = await db.execute(
                "SELECT MAX(CAST(SUBSTR(request_number, ?) AS INTEGER)) FROM requests WHERE request_number LIKE ?",
                (len(f"REL{year}") + 1, f"REL{year}%")
            )
            max_row = await cursor.fetchone()
            request_number = f"REL{year}{(max_row[0] or 0) + 1:05d}"
            request_data['original_rr_number'] = original_rr or None

            # Check for duplicate request_number
            dup_cursor = await db.execute(
                "SELECT id FROM requests WHERE request_number = ?", (request_number,)
            )
            if await dup_cursor.fetchone():
                results_errors.append({
                    'file': fname,
                    'errors': [f"Request '{request_number}' already exists in the system."]
                })
                continue

            # Store RRS# in classification if present, keep classification clean
            if rrs_number:
                existing_class = request_data.get('classification', '')
                if existing_class:
                    request_data['classification'] = f"{existing_class} | {rrs_number}"
                else:
                    request_data['classification'] = rrs_number

            steps = [ProcessStep(**step) for step in DEFAULT_STEPS]
            request_data['request_number'] = request_number
            request_data['status'] = 'incoming'

            request_obj = RelRequest(
                **request_data, created_by=current_user.id,
                created_by_username=current_user.username, steps=steps
            )

            await db.execute(
                """INSERT INTO requests (id, request_number, classification, originator, plant,
                   device_name, lot_no, customer, pkg_info, automotive, date_ltc,
                   product_hierarchy, pdl, body_size_x, body_size_y, package_thickness,
                   ball_pitch, ball_count, lead_pitch, lead_count, total_ss, purpose,
                   engineer_special_instruction, deadline, created_by, created_by_username,
                   created_at, updated_at, status, current_step, original_rr_number)
                   VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                (request_obj.id, request_obj.request_number,
                 request_obj.classification or '', request_obj.originator or '',
                 request_obj.plant or '', request_obj.device_name or '',
                 request_obj.lot_no or '', request_obj.customer or '',
                 request_obj.pkg_info or '', int(request_obj.automotive),
                 request_obj.date_ltc, request_obj.product_hierarchy,
                 request_obj.pdl, request_obj.body_size_x, request_obj.body_size_y,
                 request_obj.package_thickness, request_obj.ball_pitch, request_obj.ball_count,
                 request_obj.lead_pitch, request_obj.lead_count, request_obj.total_ss,
                 request_obj.purpose or '', request_obj.engineer_special_instruction,
                 request_obj.deadline, request_obj.created_by, request_obj.created_by_username,
                 request_obj.created_at.isoformat(), request_obj.updated_at.isoformat(),
                 request_obj.status, request_obj.current_step, request_obj.original_rr_number)
            )

            if leg_data:
                # ── Use traveller leg items as process steps ────────────────
                for leg_info in leg_data:
                    leg_num = leg_info['leg']
                    for step_num, item in enumerate(leg_info['items'], start=1):
                        cf = {}
                        if item.get('test_condition'):
                            cf['test_condition'] = item['test_condition']
                        if item.get('qty_out_result'):
                            cf['qty_out_result'] = item['qty_out_result']
                        await db.execute(
                            """INSERT INTO process_steps
                               (request_id, leg, step_number, step_name, status,
                                started_at, completed_at, machine_no, operator_id,
                                tray_no, qty_in, notes, attachments, custom_fields)
                               VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                            (request_obj.id, leg_num, step_num, item['step_name'], 'pending',
                             _clean_dt(item.get('time_in')), _clean_dt(item.get('time_out')),
                             item.get('machine_no'), None,
                             item.get('tray_no'), item.get('qty_in'),
                             item.get('operator'),
                             json.dumps([]), json.dumps(cf))
                        )
            else:
                # ── Fall back to default process steps ─────────────────────
                for step in steps:
                    await db.execute(
                        """INSERT INTO process_steps
                           (request_id, leg, step_number, step_name, status,
                            started_at, completed_at, machine_no, operator_id,
                            tray_no, notes, attachments, custom_fields)
                           VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                        (request_obj.id, 1, step.step_number, step.step_name, step.status.value,
                         step.started_at.isoformat() if step.started_at else None,
                         step.completed_at.isoformat() if step.completed_at else None,
                         step.machine_no, step.operator_id, step.tray_no, step.notes,
                         json.dumps(step.attachments or []), json.dumps(step.custom_fields or {}))
                    )

            await db.commit()

            created_entry = {
                'file': fname,
                'request_number': request_number,
                'device_name': request_obj.device_name or '',
                'customer': request_obj.customer or '',
            }
            if leg_data:
                created_entry['legs'] = len(leg_data)
                created_entry['leg_items'] = sum(len(lg['items']) for lg in leg_data)
            if file_errors:
                created_entry['warnings'] = file_errors
            results_created.append(created_entry)
        except Exception as e:
            logging.error(f"Import error for {fname}: {e}", exc_info=True)
            results_errors.append({'file': fname, 'errors': [str(e)]})
        finally:
            if db:
                await db.close()

    return {
        'total_files': len(files),
        'created': len(results_created),
        'failed': len(results_errors),
        'created_requests': results_created,
        'errors': results_errors,
    }


@api_router.post("/requests/import-word")
async def import_requests_from_word(
    files: List[UploadFile] = File(...),
    current_user: User = Depends(require_permission('import_requests'))
):
    """Import requests from Word Reliability Test Traveller (.docx) files.
    Each .docx file is treated as one request. Multiple files can be uploaded at once."""

    results_created = []
    results_errors = []

    for file in files:
        fname = file.filename or 'unknown'

        if not fname.lower().endswith('.docx'):
            results_errors.append({'file': fname, 'errors': ['Not a Word file (.docx)']})
            continue

        db = None
        try:
            contents = await file.read()

            parsed = _parse_word_request_sheet(contents)
            request_data = parsed['data']
            file_errors = parsed['errors']
            leg_data = parsed.get('leg_data', [])

            if not request_data or not any(request_data.get(k) for k in ('device_name', 'lot_no', 'customer', 'request_number')):
                file_errors.append('No recognizable data found. Is this a Reliability Test Traveller document?')
                results_errors.append({'file': fname, 'errors': file_errors})
                continue

            original_rr = request_data.pop('request_number', None)

            db = await get_db()
            now = datetime.now(timezone.utc)
            year = now.year
            week = now.isocalendar()[1]
            rms_prefix = f"RMS{year}{week:02d}"
            cursor = await db.execute(
                "SELECT MAX(CAST(SUBSTR(request_number, ?) AS INTEGER)) FROM requests WHERE request_number LIKE ?",
                (len(rms_prefix) + 1, f"{rms_prefix}%")
            )
            max_row = await cursor.fetchone()
            request_number = f"{rms_prefix}{(max_row[0] or 0) + 1:02d}"
            request_data['original_rr_number'] = original_rr or None

            dup_cursor = await db.execute(
                "SELECT id FROM requests WHERE request_number = ?", (request_number,)
            )
            if await dup_cursor.fetchone():
                results_errors.append({
                    'file': fname,
                    'errors': [f"Request '{request_number}' already exists in the system."]
                })
                continue

            steps = [ProcessStep(**step) for step in DEFAULT_STEPS]
            request_data['request_number'] = request_number
            request_data['status'] = 'incoming'

            request_obj = RelRequest(
                **request_data, created_by=current_user.id,
                created_by_username=current_user.username, steps=steps
            )

            await db.execute(
                """INSERT INTO requests (id, request_number, classification, originator, plant,
                   device_name, lot_no, customer, pkg_info, automotive, date_ltc,
                   product_hierarchy, pdl, body_size_x, body_size_y, package_thickness,
                   ball_pitch, ball_count, lead_pitch, lead_count, total_ss, purpose,
                   engineer_special_instruction, deadline, created_by, created_by_username,
                   created_at, updated_at, status, current_step, original_rr_number)
                   VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                (request_obj.id, request_obj.request_number,
                 request_obj.classification or '', request_obj.originator or '',
                 request_obj.plant or '', request_obj.device_name or '',
                 request_obj.lot_no or '', request_obj.customer or '',
                 request_obj.pkg_info or '', int(request_obj.automotive),
                 request_obj.date_ltc, request_obj.product_hierarchy,
                 request_obj.pdl, request_obj.body_size_x, request_obj.body_size_y,
                 request_obj.package_thickness, request_obj.ball_pitch, request_obj.ball_count,
                 request_obj.lead_pitch, request_obj.lead_count, request_obj.total_ss,
                 request_obj.purpose or '', request_obj.engineer_special_instruction,
                 request_obj.deadline, request_obj.created_by, request_obj.created_by_username,
                 request_obj.created_at.isoformat(), request_obj.updated_at.isoformat(),
                 request_obj.status, request_obj.current_step, request_obj.original_rr_number)
            )

            if leg_data:
                # ── Use traveller leg items as process steps ──────────────────────────
                for leg_info in leg_data:
                    leg_num = leg_info['leg']
                    for step_num, item in enumerate(leg_info['items'], start=1):
                        cf = {}
                        if item.get('test_condition'):
                            cf['test_condition'] = item['test_condition']
                        if item.get('qty_out_result'):
                            cf['qty_out_result'] = item['qty_out_result']
                        await db.execute(
                            """INSERT INTO process_steps
                               (request_id, leg, step_number, step_name, status,
                                started_at, completed_at, machine_no, operator_id,
                                tray_no, qty_in, notes, attachments, custom_fields)
                               VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                            (request_obj.id, leg_num, step_num, item['step_name'], 'pending',
                             None, None,
                             item.get('machine_no'), None,
                             item.get('tray_no'), item.get('qty_in'),
                             item.get('operator'),
                             json.dumps([]), json.dumps(cf))
                        )
            else:
                # ── Fall back to default process steps ────────────────────────────────
                for step in steps:
                    await db.execute(
                        """INSERT INTO process_steps (request_id, leg, step_number, step_name, status,
                           started_at, completed_at, machine_no, operator_id, tray_no, notes, attachments, custom_fields)
                           VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                        (request_obj.id, 1, step.step_number, step.step_name, step.status.value,
                         step.started_at.isoformat() if step.started_at else None,
                         step.completed_at.isoformat() if step.completed_at else None,
                         step.machine_no, step.operator_id, step.tray_no, step.notes,
                         json.dumps(step.attachments or []), json.dumps(step.custom_fields or {}))
                    )

            await db.commit()

            created_entry = {
                'file': fname,
                'request_number': request_number,
                'device_name': request_obj.device_name or '',
                'customer': request_obj.customer or '',
            }
            if leg_data:
                created_entry['legs'] = len(leg_data)
                created_entry['leg_items'] = sum(len(lg['items']) for lg in leg_data)
            if file_errors:
                created_entry['warnings'] = file_errors
            results_created.append(created_entry)
        except Exception as e:
            logging.error(f"Word import error for {fname}: {e}", exc_info=True)
            results_errors.append({'file': fname, 'errors': [str(e)]})
        finally:
            if db:
                await db.close()

    return {
        'total_files': len(files),
        'created': len(results_created),
        'failed': len(results_errors),
        'created_requests': results_created,
        'errors': results_errors,
    }


# ========================
# Import Whisker Test
# ========================

def _parse_whisker_docx(contents: bytes) -> dict:
    """Parse a Whisker Test Request Form (.docx) and return structured data."""
    import re as _re
    from docx import Document as _Document

    doc = _Document(io.BytesIO(contents))
    tables = doc.tables

    result: dict = {'errors': [], 'legs': []}

    # ── Paragraphs: purpose / description ──────────────────────────────────
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    purpose_parts = []
    in_purpose = False
    for p in paragraphs:
        if 'DETAIL DESCRIPTION OF PURPOSE' in p.upper():
            in_purpose = True
            continue
        if in_purpose:
            purpose_parts.append(p)
    result['purpose'] = ' '.join(purpose_parts).strip()

    # Helpers
    def _cell(row_cells, idx, default=''):
        try:
            return row_cells[idx].strip() if idx < len(row_cells) else default
        except Exception:
            return default

    def _first_nonempty(cells, start=2):
        for c in cells[start:]:
            v = c.strip() if hasattr(c, 'strip') else str(c).strip()
            if v:
                return v
        return ''

    def _parse_num(s, cast=float):
        m = _re.search(r'[\d.]+', s.replace(',', ''))
        return cast(m.group()) if m else None

    # ── Table 1: Service Request No. ───────────────────────────────────────
    if tables:
        t1 = tables[0]
        for row in t1.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) >= 2 and 'service request' in cells[0].lower():
                result['service_request_no'] = cells[1]
                break

    # ── Table 2: Package & Process Info ────────────────────────────────────
    if len(tables) > 1:
        t2 = tables[1]
        rows = {i: [c.text.strip() for c in row.cells] for i, row in enumerate(t2.rows)}

        result['pkg_code']        = _first_nonempty(rows.get(5, []))  # Package Code
        result['pkg_type']        = _first_nonempty(rows.get(6, []))  # Package Type
        result['pkg_size']        = _first_nonempty(rows.get(7, []))  # Package Size
        result['pkg_thickness']   = _first_nonempty(rows.get(8, []))  # Package Thickness
        result['lead_count_raw']  = _first_nonempty(rows.get(9, []))  # Lead Count
        result['lead_pitch_raw']  = _first_nonempty(rows.get(10, [])) # Lead Pitch
        result['factory']         = _first_nonempty(rows.get(12, [])) # Factory / Assembly Site
        result['plating_finish']  = _first_nonempty(rows.get(14, [])) # Plating Finish
        result['plating_chem']    = _first_nonempty(rows.get(15, [])) # Plating Chemistry

        # Base Metal Alloy is in the second half of row 14
        r14 = rows.get(14, [])
        result['base_metal'] = _cell(r14, 10) or _cell(r14, 11)

        # Quantity: row 6, searching for a numeric-only cell after col 8
        r6 = rows.get(6, [])
        for r6_cell in r6[8:]:
            v = r6_cell.strip()
            if v and v.isdigit():
                result['quantity'] = v
                break
        else:
            result['quantity'] = ''

        # Date submitted: row 5, first non-empty cell after col 8
        r5 = rows.get(5, [])
        for r5_cell in r5[8:]:
            v = r5_cell.strip()
            if v and v not in ('', '-'):
                result['date_submitted'] = v
                break

        # Lot labels from row 17+ (Plating Date / Lot names)
        lots = []
        for ri in (17, 18, 19):
            r = rows.get(ri, [])
            lot_label = _cell(r, 2)
            if lot_label and lot_label.lower().startswith('lot'):
                lots.append(lot_label)
        result['lots'] = lots

    # ── Table 3: Test Matrix ────────────────────────────────────────────────
    if len(tables) > 2:
        t3 = tables[2]
        header_passed = False
        for row in t3.rows:
            cells = [c.text.strip() for c in row.cells]
            if not cells:
                continue
            if cells[0].lower() == 'leg no':
                header_passed = True
                continue
            if header_passed and cells[0].lower().startswith('leg'):
                m = _re.search(r'\d+', cells[0])
                leg_num = int(m.group()) if m else len(result['legs']) + 1
                test  = _cell(cells, 1)
                cond  = _cell(cells, 2)
                precon = _cell(cells, 3)
                ss    = _cell(cells, 4)
                rpts  = _cell(cells, 5)
                result['legs'].append({
                    'leg_num': leg_num,
                    'test': test,
                    'condition': cond,
                    'precon': precon,
                    'ss': ss,
                    'read_points': rpts,
                })

    # ── Table 4: Special Instructions ──────────────────────────────────────
    if len(tables) > 3:
        t4 = tables[3]
        for row in t4.rows:
            for cell in row.cells:
                txt = cell.text.strip()
                if txt and 'special instruction' in txt.lower():
                    # instruction text is everything after the colon
                    parts = txt.split('\n')
                    result['special_instruction'] = ' '.join(
                        p for p in parts if 'special instruction' not in p.lower()
                    ).strip()
                    break

    return result


def _map_whisker_test_item(test: str, condition: str) -> str:
    """Map raw test name + condition to one of the WHISKER_TEST_ITEMS values."""
    t = test.upper().strip()
    c = condition.lower()
    if t == 'TH':
        if '55' in c and '85' in c:
            return 'TH 55/85'
        return 'TH 30/60'
    if t == 'TC':
        return 'TC A -55/85'
    # Fallback: try to build a display name
    return test


# Whisker-specific short step template (without precon — added dynamically)
WHISKER_BASE_STEPS = [
    'Incoming Inspection',
    'Visual',
    'Serialize Samples',
    'O/S',
    'SAT',
]


@api_router.post("/requests/import-whisker")
async def import_whisker_request(
    file: UploadFile = File(...),
    current_user: User = Depends(require_permission('import_requests'))
):
    """
    Import a Whisker Test Request Form (.docx), auto-create a request with
    REL numbering, and create one leg per test matrix row using the whisker
    step template.  Step 6 of each leg is the Whisker Test step (renamed
    from the Test column), with auto-filled test_item, test_condition and
    read_points custom fields.  Preconditioning is only inserted when the
    document specifies a precon condition (not 'No Precon').
    """
    fname = file.filename or 'unknown'
    if not fname.lower().endswith('.docx'):
        raise HTTPException(status_code=400, detail='Please upload a .docx file')

    contents = await file.read()
    try:
        parsed = _parse_whisker_docx(contents)
    except Exception as e:
        logging.error(f"Whisker parse error: {e}", exc_info=True)
        raise HTTPException(status_code=422, detail=f'Could not parse document: {e}')

    legs_data = parsed.get('legs', [])
    if not legs_data:
        raise HTTPException(status_code=422, detail='No test matrix legs found in the document')

    # ── Build request fields ─────────────────────────────────────────────
    pkg_type  = parsed.get('pkg_type', '')
    pkg_code  = parsed.get('pkg_code', '')
    pkg_size  = parsed.get('pkg_size', '')
    lead_raw  = parsed.get('lead_count_raw', '')
    pitch_raw = parsed.get('lead_pitch_raw', '')
    factory   = parsed.get('factory', '')
    plating   = parsed.get('plating_finish', '')
    base_metal = parsed.get('base_metal', '')
    quantity  = parsed.get('quantity', '')

    def _parse_num_val(s, cast=float):
        import re as _r
        m = _r.search(r'[\d.]+', (s or '').replace(',', ''))
        try:
            return cast(m.group()) if m else None
        except Exception:
            return None

    lead_count  = _parse_num_val(lead_raw, int)
    lead_pitch  = _parse_num_val(pitch_raw, float)

    device_name = ' '.join(filter(None, [pkg_code, pkg_type, lead_raw])).strip()
    pkg_info    = ' '.join(filter(None, [pkg_type, pkg_size])).strip()

    # Build pkg_info with plating details
    if plating or base_metal:
        extras = ' / '.join(filter(None, [plating, base_metal]))
        pkg_info = f"{pkg_info} - {extras}" if pkg_info else extras

    db = await get_db()
    try:
        # ── Auto-generate REL number ────────────────────────────────────
        now = datetime.now(timezone.utc)
        year = now.year
        cursor = await db.execute(
            "SELECT MAX(CAST(SUBSTR(request_number, ?) AS INTEGER)) FROM requests WHERE request_number LIKE ?",
            (len(f"REL{year}") + 1, f"REL{year}%")
        )
        max_row = await cursor.fetchone()
        request_number = f"REL{year}{(max_row[0] or 0) + 1:05d}"

        # ── Check duplicate ─────────────────────────────────────────────
        dup = await db.execute("SELECT id FROM requests WHERE request_number = ?", (request_number,))
        if await dup.fetchone():
            raise HTTPException(status_code=409, detail=f"Request '{request_number}' already exists")

        # ── Build RelRequest object ───────────────────────────────────
        req_id = str(uuid.uuid4())
        status = 'incoming'
        num_legs_val = str(len(legs_data))

        await db.execute(
            """INSERT INTO requests
               (id, request_number, classification, originator, plant,
                device_name, lot_no, customer, pkg_info, automotive, date_ltc,
                product_hierarchy, pdl, body_size_x, body_size_y, package_thickness,
                ball_pitch, ball_count, lead_pitch, lead_count, total_ss, purpose,
                engineer_special_instruction, deadline, created_by, created_by_username,
                created_at, updated_at, status, current_step, num_legs)
               VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
            (req_id, request_number,
             '3.0 Whisker Monitor', '', factory,
             device_name, '', '',
             pkg_info, 0, None,
             None, None, None, None, None,
             None, None, lead_pitch, lead_count,
             quantity,
             parsed.get('purpose', ''),
             parsed.get('special_instruction', ''),
             None,
             current_user.id, current_user.username,
             now.isoformat(), now.isoformat(),
             status, 1, num_legs_val)
        )

        # ── Create process steps per leg ────────────────────────────────
        for leg_info in legs_data:
            leg_num  = leg_info['leg_num']
            test     = leg_info['test']          # e.g. "TH" or "TC"
            cond     = leg_info['condition']     # e.g. "30°C/60%rh"
            precon   = leg_info['precon']        # e.g. "No Precon" or "1x @ 215°C"
            ss       = leg_info['ss']            # sample size, e.g. "6"
            rpts     = leg_info['read_points']   # e.g. "0,1000,1500hrs"

            has_precon = precon and precon.strip().lower() != 'no precon'
            test_item  = _map_whisker_test_item(test, cond)

            # Build step list: base + optional precon + whisker test
            step_list = list(WHISKER_BASE_STEPS)
            if has_precon:
                step_list.append('Preconditioning (Precon)')
            step_list.append('Whisker Test')

            for step_num, step_name in enumerate(step_list, start=1):
                cf: dict = {}
                qty_in_val = None

                if step_name == 'Preconditioning (Precon)':
                    cf['test_condition'] = precon

                elif step_name == 'Whisker Test':
                    cf['test_item']      = test_item
                    cf['test_condition'] = cond
                    cf['read_points']    = rpts
                    if ss:
                        try:
                            qty_in_val = int(ss)
                        except ValueError:
                            pass

                await db.execute(
                    """INSERT INTO process_steps
                       (request_id, leg, step_number, step_name, status,
                        started_at, completed_at, machine_no, operator_id,
                        tray_no, qty_in, notes, attachments, custom_fields)
                       VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                    (req_id, leg_num, step_num, step_name, 'pending',
                     None, None, None, None,
                     None, qty_in_val, None,
                     json.dumps([]), json.dumps(cf))
                )

        await db.commit()

        # Build a preview summary for the response
        leg_summary = [
            {
                'leg': l['leg_num'],
                'test': l['test'],
                'condition': l['condition'],
                'precon': l['precon'],
                'ss': l['ss'],
                'read_points': l['read_points'],
                'test_item': _map_whisker_test_item(l['test'], l['condition']),
                'has_precon': bool(l['precon'] and l['precon'].strip().lower() != 'no precon'),
            }
            for l in legs_data
        ]

        return {
            'request_id': req_id,
            'request_number': request_number,
            'device_name': device_name,
            'pkg_info': pkg_info,
            'factory': factory,
            'num_legs': len(legs_data),
            'legs': leg_summary,
            'warnings': parsed.get('errors', []),
        }

    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Whisker import error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        await db.close()


# ========================
# File Upload Route
# ========================
@api_router.post("/upload")
async def upload_file(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user)
):
    try:
        upload_dir = ROOT_DIR / "uploads"
        upload_dir.mkdir(exist_ok=True)
        file_extension = Path(file.filename).suffix
        unique_filename = f"{uuid.uuid4()}{file_extension}"
        file_path = upload_dir / unique_filename
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        return {"filename": unique_filename, "original_filename": file.filename,
                "url": f"/uploads/{unique_filename}"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to upload file: {str(e)}")


# ========================
# Backup System
# ========================
# Create dedicated backup folder at project root for secure data storage
PROJECT_ROOT = ROOT_DIR.parent  # Go up from backend to main project
BACKUP_DIR = PROJECT_ROOT / "Rel_Request_Backups"
BACKUP_DIR.mkdir(exist_ok=True)

# Create organized subfolders for better management
AUTO_BACKUP_DIR = BACKUP_DIR / "Auto_Backups"
MANUAL_BACKUP_DIR = BACKUP_DIR / "Manual_Backups"
AUTO_BACKUP_DIR.mkdir(exist_ok=True)
MANUAL_BACKUP_DIR.mkdir(exist_ok=True)

# Create backup info file if it doesn't exist
backup_info_file = BACKUP_DIR / "BACKUP_INFO.txt"
if not backup_info_file.exists():
    with open(backup_info_file, 'w', encoding='utf-8') as f:
        f.write("""â•”â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•—
â•‘         REL REQUEST PROCESS FLOW - BACKUP STORAGE                 â•‘
â•‘                   SECURE DATA ARCHIVE                              â•‘
â•šâ•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

This folder contains critical backup files for the Reliability Request 
Process Flow system. These Excel files preserve all historical data 
including completed requests, process steps, and system configurations.

ðŸ“ FOLDER STRUCTURE:
â”œâ”€â”€ Manual_Backups/     - User-created backups
â”œâ”€â”€ Auto_Backups/       - System-generated monthly backups

âš ï¸  IMPORTANT NOTES:
â€¢ DO NOT DELETE these files - they contain irreplaceable historical data
â€¢ Backup files are in Excel format (.xlsx) for easy access and portability
â€¢ Each backup includes: Requests, Process Steps, Users, Settings, Login Logs
â€¢ Completed requests are automatically removed from active database after backup
â€¢ Keep backups secure - they may contain sensitive customer/device information

ðŸ“Š BACKUP DETAILS:
â€¢ Manual backups: Created on-demand by administrators
â€¢ Auto backups: Created monthly on the 1st at 2:00 AM
â€¢ Retention: Last 12 monthly backups are kept automatically
â€¢ File naming: rel_database_backup_YYYY-MM-DD_HHMMSS.xlsx

ðŸ”’ SECURITY:
â€¢ This folder should be backed up to external storage regularly
â€¢ Consider encrypting backups if stored on shared drives
â€¢ Ensure proper access controls are maintained

Last updated: {}""" .format(datetime.now().strftime("%Y-%m-%d %H:%M:%S")))
    logging.info(f"Backup directory initialized at: {BACKUP_DIR}")

# Create .gitignore to protect backup files
gitignore_file = BACKUP_DIR / ".gitignore"
if not gitignore_file.exists():
    with open(gitignore_file, 'w') as f:
        f.write("""# Ignore all backup files to prevent accidental commits
*.xlsx
*.xls
*.db
*.sqlite

# But keep the info file and structure
!BACKUP_INFO.txt
!.gitignore
""")


async def create_backup_excel(is_auto=False):
    """Create a timestamped backup as a ZIP containing the Excel database export
    plus all SAT attachment images from the uploads folder.
    Organizes backups into year/month subfolders for better management.
    
    Args:
        is_auto: If True, saves to Auto_Backups folder, else Manual_Backups
    """
    import zipfile
    db = await aiosqlite.connect(DB_PATH)
    db.row_factory = aiosqlite.Row
    
    try:
        now = datetime.now()
        timestamp = now.strftime("%Y-%m-%d_%H%M%S")
        excel_name  = f"rel_database_backup_{timestamp}.xlsx"
        backup_name = f"rel_database_backup_{timestamp}.zip"
        
        # Determine base directory based on backup type
        base_dir = AUTO_BACKUP_DIR if is_auto else MANUAL_BACKUP_DIR
        
        # Create year and month subfolders for organization
        year_folder = base_dir / str(now.year)
        month_folder = year_folder / f"{now.month:02d}_{now.strftime('%B')}"
        month_folder.mkdir(parents=True, exist_ok=True)
        
        backup_path = month_folder / backup_name
        excel_tmp   = month_folder / excel_name  # written temporarily, then zipped
        
        # Create Excel workbook
        wb = Workbook()
        wb.remove(wb.active)  # Remove default sheet
        
        # Export Users table
        cursor = await db.execute("SELECT id, email, username, role, approved, created_at FROM users")
        rows = await cursor.fetchall()
        if rows:
            ws = wb.create_sheet("Users")
            ws.append(["ID", "Email", "Username", "Role", "Approved", "Created At"])
            for row in rows:
                ws.append(list(row))
        
        # Export Requests table (only completed requests)
        cursor = await db.execute("SELECT * FROM requests WHERE status = 'completed'")
        rows = await cursor.fetchall()
        if rows:
            ws = wb.create_sheet("Requests")
            # Get column names
            cursor2 = await db.execute("PRAGMA table_info(requests)")
            columns = await cursor2.fetchall()
            headers = [col[1] for col in columns]
            ws.append(headers)
            for row in rows:
                ws.append(list(row))
        
        # Export Process Steps table (only for completed requests)
        cursor = await db.execute(
            "SELECT * FROM process_steps WHERE request_id IN (SELECT id FROM requests WHERE status = 'completed')"
        )
        rows = await cursor.fetchall()
        if rows:
            ws = wb.create_sheet("Process Steps")
            cursor2 = await db.execute("PRAGMA table_info(process_steps)")
            columns = await cursor2.fetchall()
            headers = [col[1] for col in columns]
            ws.append(headers)
            for row in rows:
                ws.append(list(row))
        
        # Export Settings table
        cursor = await db.execute("SELECT * FROM settings")
        rows = await cursor.fetchall()
        if rows:
            ws = wb.create_sheet("Settings")
            cursor2 = await db.execute("PRAGMA table_info(settings)")
            columns = await cursor2.fetchall()
            headers = [col[1] for col in columns]
            ws.append(headers)
            for row in rows:
                ws.append(list(row))
        
        # Export Role Permissions table
        cursor = await db.execute("SELECT * FROM role_permissions")
        rows = await cursor.fetchall()
        if rows:
            ws = wb.create_sheet("Role Permissions")
            ws.append(["Role", "Permission", "Granted"])
            for row in rows:
                ws.append(list(row))
        
        # Export Login Logs table
        cursor = await db.execute("SELECT * FROM login_logs ORDER BY login_at DESC LIMIT 1000")
        rows = await cursor.fetchall()
        if rows:
            ws = wb.create_sheet("Login Logs")
            ws.append(["ID", "User ID", "Email", "Username", "Role", "Login At", "IP Address"])
            for row in rows:
                ws.append(list(row))
        
        # Save workbook to a temp .xlsx file
        wb.save(str(excel_tmp))

        # Collect all SAT attachment image paths referenced in completed process steps
        uploads_dir = Path(os.path.dirname(DB_PATH)) / 'uploads'
        cursor = await db.execute(
            "SELECT attachments FROM process_steps "
            "WHERE request_id IN (SELECT id FROM requests WHERE status = 'completed') "
            "AND attachments IS NOT NULL AND attachments != '[]' AND attachments != '{}'"
        )
        att_rows = await cursor.fetchall()
        sat_files = set()
        for (att_raw,) in att_rows:
            try:
                val = json.loads(att_raw) if att_raw else None
                if isinstance(val, dict):
                    for urls in val.values():
                        if isinstance(urls, list):
                            for u in urls:
                                if u:
                                    sat_files.add(str(u).split('/')[-1])
                elif isinstance(val, list):
                    for u in val:
                        if u:
                            sat_files.add(str(u).split('/')[-1])
            except Exception:
                pass

        # Build the ZIP: Excel sheet + images
        with zipfile.ZipFile(str(backup_path), 'w', zipfile.ZIP_DEFLATED) as zf:
            zf.write(str(excel_tmp), arcname=excel_name)
            img_count = 0
            for fname in sat_files:
                img_path = uploads_dir / fname
                if img_path.exists():
                    zf.write(str(img_path), arcname=f"uploads/{fname}")
                    img_count += 1

        # Remove the temp Excel file (it's now inside the ZIP)
        try:
            excel_tmp.unlink()
        except Exception:
            pass

        backup_type = "Auto" if is_auto else "Manual"
        relative_path = f"{now.year}/{now.month:02d}_{now.strftime('%B')}/{backup_name}"
        logging.info(f"{backup_type} backup created: {relative_path} (Excel + {img_count} SAT images)")
        return backup_name, relative_path
        
    finally:
        await db.close()


def create_backup_sync(is_auto=False):
    """Synchronous wrapper for creating Excel backup.
    
    Args:
        is_auto: If True, creates auto backup, else manual backup
        
    Returns:
        tuple: (backup_name, relative_path)
    """
    loop = asyncio.new_event_loop()
    asyncio.set_event_loop(loop)
    try:
        return loop.run_until_complete(create_backup_excel(is_auto=is_auto))
    finally:
        loop.close()


@api_router.post("/backups")
async def create_backup(current_user: User = Depends(require_permission('manage_backups'))):
    """Create a manual database backup as Excel file, then delete completed requests."""
    db = await get_db()
    try:
        # Get counts before backup
        cursor = await db.execute("SELECT COUNT(*) FROM requests")
        total_before = (await cursor.fetchone())[0]
        
        cursor = await db.execute("SELECT COUNT(*) FROM requests WHERE status = 'completed'")
        completed_count = (await cursor.fetchone())[0]
        
        # Create manual backup
        backup_name, relative_path = await create_backup_excel(is_auto=False)
        if not backup_name:
            raise HTTPException(status_code=500, detail="Database export failed")
        
        # Delete completed requests and their associated process steps (CASCADE will handle steps)
        await db.execute("DELETE FROM requests WHERE status = 'completed'")
        await db.commit()
        
        # Get count after deletion
        cursor = await db.execute("SELECT COUNT(*) FROM requests")
        total_after = (await cursor.fetchone())[0]
        
        # Update backup tracking
        await db.execute(
            """UPDATE backup_tracking 
            SET last_critical_backup_at = ?, 
                last_backup_request_count = ?,
                last_backup_downloaded = 0
            WHERE id = 1""",
            (datetime.now(timezone.utc).isoformat(), total_after)
        )
        await db.commit()
        
        logging.info(f"Backup created: {relative_path}. Deleted {completed_count} completed requests. Remaining: {total_after}")
        
        return {
            "message": "Backup created successfully",
            "filename": backup_name,
            "relative_path": relative_path,
            "deleted_completed_requests": completed_count,
            "remaining_requests": total_after
        }
    except Exception as e:
        await db.rollback()
        raise HTTPException(status_code=500, detail=f"Backup failed: {str(e)}")
    finally:
        await db.close()


@api_router.get("/backups")
async def list_backups(current_user: User = Depends(require_permission('manage_backups'))):
    """List all database backups from both manual and auto folders with year/month organization."""
    backups = []
    
    # Helper function to scan directory recursively
    def scan_backup_folder(base_path, backup_type):
        result = []
        if not base_path.exists():
            return result
            
        # Scan for backup files in year/month subfolders
        for year_folder in base_path.iterdir():
            if year_folder.is_dir() and year_folder.name.isdigit():
                for month_folder in year_folder.iterdir():
                    if month_folder.is_dir():
                        for file in list(month_folder.glob("rel_database_backup_*.zip")) + list(month_folder.glob("rel_database_backup_*.xlsx")):
                            stat = file.stat()
                            relative_path = f"{year_folder.name}/{month_folder.name}/{file.name}"
                            result.append({
                                "filename": file.name,
                                "relative_path": relative_path,
                                "type": backup_type,
                                "size_mb": round(stat.st_size / (1024 * 1024), 2),
                                "created_at": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                            })
        return result
    
    # Collect backups from Manual_Backups folder
    backups.extend(scan_backup_folder(MANUAL_BACKUP_DIR, "Manual"))
    
    # Collect backups from Auto_Backups folder
    backups.extend(scan_backup_folder(AUTO_BACKUP_DIR, "Auto"))
    
    # Sort by creation date, newest first
    backups.sort(key=lambda x: x['created_at'], reverse=True)
    return backups


@api_router.get("/backups/{filename}")
async def download_backup(filename: str, current_user: User = Depends(require_permission('manage_backups'))):
    """Download a specific backup file from organized year/month folders."""
    if not filename.startswith("rel_database_backup_"):
        raise HTTPException(status_code=404, detail="Backup not found")
    
    # Helper function to find file in nested folders
    def find_backup_file(base_path):
        if not base_path.exists():
            return None
        for year_folder in base_path.iterdir():
            if year_folder.is_dir():
                for month_folder in year_folder.iterdir():
                    if month_folder.is_dir():
                        file_path = month_folder / filename
                        if file_path.exists():
                            return file_path
        return None
    
    # Search in both Manual and Auto backup folders
    path = find_backup_file(MANUAL_BACKUP_DIR) or find_backup_file(AUTO_BACKUP_DIR)
    
    if not path:
        raise HTTPException(status_code=404, detail="Backup not found")

    def iterfile():
        with open(path, "rb") as f:
            while chunk := f.read(8192):
                yield chunk

    if filename.endswith(".zip"):
        media_type = "application/zip"
    elif filename.endswith(".xlsx"):
        media_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    else:
        media_type = "application/octet-stream"
    return StreamingResponse(iterfile(), media_type=media_type,
                             headers={"Content-Disposition": f"attachment; filename={filename}"})


@api_router.post("/backups/{filename}/confirm-download")
async def confirm_backup_download(filename: str, current_user: User = Depends(require_permission('manage_backups'))):
    """Confirm that a backup has been downloaded, clearing the critical backup requirement."""
    db = await get_db()
    try:
        if not filename.startswith("rel_database_backup_"):
            raise HTTPException(status_code=404, detail="Backup not found")
        
        # Helper function to check file existence in nested folders
        def check_backup_exists(base_path):
            if not base_path.exists():
                return False
            for year_folder in base_path.iterdir():
                if year_folder.is_dir():
                    for month_folder in year_folder.iterdir():
                        if month_folder.is_dir():
                            if (month_folder / filename).exists():
                                return True
            return False
        
        # Check if file exists in either folder
        if not (check_backup_exists(MANUAL_BACKUP_DIR) or check_backup_exists(AUTO_BACKUP_DIR)):
            raise HTTPException(status_code=404, detail="Backup not found")
        
        # Mark backup as downloaded and clear critical requirement
        await db.execute(
            """UPDATE backup_tracking 
            SET last_backup_downloaded = 1,
                critical_backup_required = 0
            WHERE id = 1"""
        )
        await db.commit()
        
        return {"message": "Backup download confirmed"}
    finally:
        await db.close()


@api_router.get("/backups/{filename}/preview")
async def preview_backup(filename: str, current_user: User = Depends(require_permission('manage_backups'))):
    """Parse a backup ZIP (or legacy XLSX) and return structured request + step data for in-browser review."""
    import zipfile, io

    if not filename.startswith("rel_database_backup_"):
        raise HTTPException(status_code=404, detail="Backup not found")

    # Locate the file in nested year/month folders
    def find_file(base_path):
        if not base_path.exists():
            return None
        for yf in base_path.iterdir():
            if yf.is_dir():
                for mf in yf.iterdir():
                    if mf.is_dir():
                        p = mf / filename
                        if p.exists():
                            return p
        return None

    path = find_file(MANUAL_BACKUP_DIR) or find_file(AUTO_BACKUP_DIR)
    if not path:
        raise HTTPException(status_code=404, detail="Backup not found")

    try:
        # Extract Excel bytes
        if filename.endswith(".zip"):
            with zipfile.ZipFile(str(path), 'r') as zf:
                xlsx_names = [n for n in zf.namelist() if n.endswith('.xlsx')]
                if not xlsx_names:
                    raise HTTPException(status_code=422, detail="No Excel found inside ZIP")
                excel_bytes = zf.read(xlsx_names[0])
            wb = load_workbook(filename=io.BytesIO(excel_bytes), read_only=True, data_only=True)
        else:
            wb = load_workbook(filename=str(path), read_only=True, data_only=True)

        def ws_rows(sheet_name):
            if sheet_name not in wb.sheetnames:
                return [], []
            ws = wb[sheet_name]
            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                return [], []
            return list(rows[0]), list(rows[1:])

        # Parse Requests sheet
        req_headers, req_rows = ws_rows("Requests")
        requests_list = []
        for row in req_rows:
            if not row or all(v is None for v in row):
                continue
            d = {req_headers[i]: row[i] for i in range(min(len(req_headers), len(row)))}
            d['automotive'] = bool(d.get('automotive'))
            requests_list.append(d)

        # Parse Process Steps sheet
        step_headers, step_rows = ws_rows("Process Steps")
        steps_by_req = {}
        for row in step_rows:
            if not row or all(v is None for v in row):
                continue
            d = {step_headers[i]: row[i] for i in range(min(len(step_headers), len(row)))}
            rid = d.get('request_id')
            if not rid:
                continue
            # Parse JSON fields
            try:
                d['attachments'] = json.loads(d['attachments']) if d.get('attachments') else []
            except Exception:
                d['attachments'] = []
            try:
                d['custom_fields'] = json.loads(d['custom_fields']) if d.get('custom_fields') else {}
            except Exception:
                d['custom_fields'] = {}
            steps_by_req.setdefault(rid, []).append(d)

        # Attach steps to requests and sort
        for req in requests_list:
            rid = req.get('id')
            steps = steps_by_req.get(rid, [])
            steps.sort(key=lambda s: (s.get('leg') or 1, s.get('step_number') or 0))
            req['steps'] = steps

        wb.close()
        return {"requests": requests_list, "count": len(requests_list)}

    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Backup preview error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to parse backup: {str(e)}")


@api_router.post("/backups/upload-preview")
async def upload_preview_backup(
    file: UploadFile = File(...),
    current_user: User = Depends(require_permission('manage_backups'))
):
    """Upload a local backup file (.zip or .xlsx) and parse it for in-browser review without saving."""
    import zipfile, io

    filename = file.filename or ""
    if not (filename.endswith('.zip') or filename.endswith('.xlsx')):
        raise HTTPException(status_code=422, detail="Only .zip or .xlsx backup files are supported")

    try:
        contents = await file.read()

        if filename.endswith('.zip'):
            with zipfile.ZipFile(io.BytesIO(contents), 'r') as zf:
                xlsx_names = [n for n in zf.namelist() if n.endswith('.xlsx')]
                if not xlsx_names:
                    raise HTTPException(status_code=422, detail="No Excel found inside ZIP")
                excel_bytes = zf.read(xlsx_names[0])
            wb = load_workbook(filename=io.BytesIO(excel_bytes), read_only=True, data_only=True)
        else:
            wb = load_workbook(filename=io.BytesIO(contents), read_only=True, data_only=True)

        def ws_rows(sheet_name):
            if sheet_name not in wb.sheetnames:
                return [], []
            ws = wb[sheet_name]
            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                return [], []
            return list(rows[0]), list(rows[1:])

        req_headers, req_rows = ws_rows("Requests")
        requests_list = []
        for row in req_rows:
            if not row or all(v is None for v in row):
                continue
            d = {req_headers[i]: row[i] for i in range(min(len(req_headers), len(row)))}
            d['automotive'] = bool(d.get('automotive'))
            requests_list.append(d)

        step_headers, step_rows = ws_rows("Process Steps")
        steps_by_req = {}
        for row in step_rows:
            if not row or all(v is None for v in row):
                continue
            d = {step_headers[i]: row[i] for i in range(min(len(step_headers), len(row)))}
            rid = d.get('request_id')
            if not rid:
                continue
            try:
                d['attachments'] = json.loads(d['attachments']) if d.get('attachments') else []
            except Exception:
                d['attachments'] = []
            try:
                d['custom_fields'] = json.loads(d['custom_fields']) if d.get('custom_fields') else {}
            except Exception:
                d['custom_fields'] = {}
            steps_by_req.setdefault(rid, []).append(d)

        for req in requests_list:
            rid = req.get('id')
            steps = steps_by_req.get(rid, [])
            steps.sort(key=lambda s: (s.get('leg') or 1, s.get('step_number') or 0))
            req['steps'] = steps

        wb.close()
        return {"requests": requests_list, "count": len(requests_list), "filename": filename}

    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Backup upload-preview error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to parse backup file: {str(e)}")


@api_router.get("/backups/status/check")
async def check_backup_status(current_user: User = Depends(get_current_user)):
    """Check if critical backup is required and get completed request count."""
    db = await get_db()
    try:
        # Get request count
        cursor = await db.execute("SELECT COUNT(*) FROM requests")
        request_count = (await cursor.fetchone())[0]
        
        # Get completed request count
        cursor = await db.execute("SELECT COUNT(*) FROM requests WHERE status = 'completed'")
        completed_count = (await cursor.fetchone())[0]
        
        # Get backup tracking
        cursor = await db.execute(
            "SELECT last_critical_backup_at, last_backup_request_count, critical_backup_required, last_backup_downloaded FROM backup_tracking WHERE id = 1"
        )
        backup_track = await cursor.fetchone()
        
        requires_backup = False
        if request_count >= 1000 and backup_track:
            last_count = backup_track[1] or 0
            was_downloaded = backup_track[3] or 0
            if last_count < 1000 or not was_downloaded:
                requires_backup = True
        elif request_count >= 1000:
            requires_backup = True
            
        return {
            "requires_critical_backup": requires_backup,
            "request_count": request_count,
            "completed_count": completed_count,
            "last_backup_count": backup_track[1] if backup_track else 0,
            "last_backup_downloaded": backup_track[3] if backup_track else 0
        }
    finally:
        await db.close()


@api_router.delete("/backups/{filename}")
async def delete_backup(filename: str, current_user: User = Depends(require_permission('manage_backups'))):
    """Delete a specific backup file from organized year/month folders."""
    if not filename.startswith("rel_database_backup_"):
        raise HTTPException(status_code=404, detail="Backup not found")
    
    # Helper function to find and delete file in nested folders
    def find_and_delete_file(base_path):
        if not base_path.exists():
            return False
        for year_folder in base_path.iterdir():
            if year_folder.is_dir():
                for month_folder in year_folder.iterdir():
                    if month_folder.is_dir():
                        file_path = month_folder / filename
                        if file_path.exists():
                            file_path.unlink()
                            # Clean up empty month folders
                            if not any(month_folder.iterdir()):
                                month_folder.rmdir()
                            # Clean up empty year folders
                            if not any(year_folder.iterdir()):
                                year_folder.rmdir()
                            return True
        return False
    
    # Try to delete from either Manual or Auto backup folders
    if not (find_and_delete_file(MANUAL_BACKUP_DIR) or find_and_delete_file(AUTO_BACKUP_DIR)):
        raise HTTPException(status_code=404, detail="Backup not found")
    
    return {"message": "Backup deleted successfully"}


# â”€â”€ Filter-Backups endpoints (Admin + Reliability Engineer, no manage_backups perm needed) â”€â”€

@api_router.get("/filter-backups")
async def list_filter_backups(
    current_user: User = Depends(require_role([UserRole.ADMIN, UserRole.RELIABILITY_ENGINEER]))
):
    """List all backups accessible to Admin and Reliability Engineer for the filter page."""
    backups = []

    def scan(base_path, backup_type):
        result = []
        if not base_path.exists():
            return result
        for year_folder in base_path.iterdir():
            if year_folder.is_dir() and year_folder.name.isdigit():
                for month_folder in year_folder.iterdir():
                    if month_folder.is_dir():
                        for file in (
                            list(month_folder.glob("rel_database_backup_*.zip")) +
                            list(month_folder.glob("rel_database_backup_*.xlsx"))
                        ):
                            stat = file.stat()
                            result.append({
                                "filename": file.name,
                                "relative_path": f"{year_folder.name}/{month_folder.name}/{file.name}",
                                "type": backup_type,
                                "size_mb": round(stat.st_size / (1024 * 1024), 2),
                                "created_at": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                            })
        return result

    backups.extend(scan(MANUAL_BACKUP_DIR, "Manual"))
    backups.extend(scan(AUTO_BACKUP_DIR, "Auto"))
    backups.sort(key=lambda x: x["created_at"], reverse=True)
    return backups


@api_router.get("/filter-backups/{filename}")
async def preview_filter_backup(
    filename: str,
    current_user: User = Depends(require_role([UserRole.ADMIN, UserRole.RELIABILITY_ENGINEER]))
):
    """Parse a backup file and return all requests + steps for client-side filtering."""
    import zipfile, io

    if not filename.startswith("rel_database_backup_"):
        raise HTTPException(status_code=404, detail="Backup not found")

    def find_file(base_path):
        if not base_path.exists():
            return None
        for yf in base_path.iterdir():
            if yf.is_dir():
                for mf in yf.iterdir():
                    if mf.is_dir():
                        p = mf / filename
                        if p.exists():
                            return p
        return None

    path = find_file(MANUAL_BACKUP_DIR) or find_file(AUTO_BACKUP_DIR)
    if not path:
        raise HTTPException(status_code=404, detail="Backup not found")

    try:
        if filename.endswith(".zip"):
            with zipfile.ZipFile(str(path), "r") as zf:
                xlsx_names = [n for n in zf.namelist() if n.endswith(".xlsx")]
                if not xlsx_names:
                    raise HTTPException(status_code=422, detail="No Excel found inside ZIP")
                excel_bytes = zf.read(xlsx_names[0])
            wb = load_workbook(filename=io.BytesIO(excel_bytes), read_only=True, data_only=True)
        else:
            wb = load_workbook(filename=str(path), read_only=True, data_only=True)

        def ws_rows(sheet_name):
            if sheet_name not in wb.sheetnames:
                return [], []
            ws = wb[sheet_name]
            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                return [], []
            return list(rows[0]), list(rows[1:])

        req_headers, req_rows = ws_rows("Requests")
        requests_list = []
        for row in req_rows:
            if not row or all(v is None for v in row):
                continue
            d = {req_headers[i]: row[i] for i in range(min(len(req_headers), len(row)))}
            d["automotive"] = bool(d.get("automotive"))
            requests_list.append(d)

        step_headers, step_rows = ws_rows("Process Steps")
        steps_by_req = {}
        for row in step_rows:
            if not row or all(v is None for v in row):
                continue
            d = {step_headers[i]: row[i] for i in range(min(len(step_headers), len(row)))}
            rid = d.get("request_id")
            if not rid:
                continue
            try:
                d["attachments"] = json.loads(d["attachments"]) if d.get("attachments") else []
            except Exception:
                d["attachments"] = []
            try:
                d["custom_fields"] = json.loads(d["custom_fields"]) if d.get("custom_fields") else {}
            except Exception:
                d["custom_fields"] = {}
            steps_by_req.setdefault(rid, []).append(d)

        for req in requests_list:
            rid = req.get("id")
            steps = steps_by_req.get(rid, [])
            steps.sort(key=lambda s: (s.get("leg") or 1, s.get("step_number") or 0))
            req["steps"] = steps

        wb.close()
        return {"requests": requests_list, "count": len(requests_list)}

    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Filter-backup preview error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to parse backup: {str(e)}")


@api_router.post("/filter-backups/upload-preview")
async def upload_preview_filter_backup(
    file: UploadFile = File(...),
    current_user: User = Depends(require_role([UserRole.ADMIN, UserRole.RELIABILITY_ENGINEER]))
):
    """Upload a local backup file (.zip or .xlsx) and parse it for the filter page."""
    import zipfile, io

    filename = file.filename or ""
    if not (filename.endswith('.zip') or filename.endswith('.xlsx')):
        raise HTTPException(status_code=422, detail="Only .zip or .xlsx backup files are supported")

    try:
        contents = await file.read()

        if filename.endswith('.zip'):
            with zipfile.ZipFile(io.BytesIO(contents), 'r') as zf:
                xlsx_names = [n for n in zf.namelist() if n.endswith('.xlsx')]
                if not xlsx_names:
                    raise HTTPException(status_code=422, detail="No Excel found inside ZIP")
                excel_bytes = zf.read(xlsx_names[0])
            wb = load_workbook(filename=io.BytesIO(excel_bytes), read_only=True, data_only=True)
        else:
            wb = load_workbook(filename=io.BytesIO(contents), read_only=True, data_only=True)

        def ws_rows(sheet_name):
            if sheet_name not in wb.sheetnames:
                return [], []
            ws = wb[sheet_name]
            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                return [], []
            return list(rows[0]), list(rows[1:])

        req_headers, req_rows = ws_rows("Requests")
        requests_list = []
        for row in req_rows:
            if not row or all(v is None for v in row):
                continue
            d = {req_headers[i]: row[i] for i in range(min(len(req_headers), len(row)))}
            d['automotive'] = bool(d.get('automotive'))
            requests_list.append(d)

        step_headers, step_rows = ws_rows("Process Steps")
        steps_by_req = {}
        for row in step_rows:
            if not row or all(v is None for v in row):
                continue
            d = {step_headers[i]: row[i] for i in range(min(len(step_headers), len(row)))}
            rid = d.get('request_id')
            if not rid:
                continue
            try:
                d['attachments'] = json.loads(d['attachments']) if d.get('attachments') else []
            except Exception:
                d['attachments'] = []
            try:
                d['custom_fields'] = json.loads(d['custom_fields']) if d.get('custom_fields') else {}
            except Exception:
                d['custom_fields'] = {}
            steps_by_req.setdefault(rid, []).append(d)

        for req in requests_list:
            rid = req.get('id')
            steps = steps_by_req.get(rid, [])
            steps.sort(key=lambda s: (s.get('leg') or 1, s.get('step_number') or 0))
            req['steps'] = steps

        wb.close()
        return {"requests": requests_list, "count": len(requests_list), "filename": filename}

    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Filter-backup upload-preview error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to parse backup file: {str(e)}")


# ========================
# Training Masterlist Routes
# ========================

@api_router.post("/masterlist/upload")
async def upload_masterlist(
    file: UploadFile = File(...),
    current_user: User = Depends(require_role([UserRole.ADMIN, UserRole.PLANNER])),
):
    """Upload and parse the training masterlist Excel file (Copy of Masterlist 2026 - for dbase)."""
    fname = file.filename or ""
    if not (fname.lower().endswith(".xlsx") or fname.lower().endswith(".xls")):
        raise HTTPException(status_code=422, detail="Only Excel files (.xlsx / .xls) are supported")

    contents = await file.read()
    try:
        if fname.lower().endswith(".xls") and _XLRD_AVAILABLE:
            import xlrd as _xlrd_local
            book = _xlrd_local.open_workbook(file_contents=contents)
            wb_obj = _XlrdWb(book)
            ws = wb_obj.active
            raw_rows = list(ws.iter_rows(values_only=True))
        else:
            wb_obj = load_workbook(filename=io.BytesIO(contents), read_only=True, data_only=True)
            ws = wb_obj.active
            raw_rows = list(ws.iter_rows(values_only=True))
            wb_obj.close()

        if not raw_rows:
            raise HTTPException(status_code=422, detail="Excel file appears empty")

        # Locate header row: scan first 10 rows for planning masterlist keywords
        header_row_idx = 0
        for i, row in enumerate(raw_rows[:10]):
            row_lower = [str(v).strip().lower() if v is not None else "" for v in row]
            if any(("ww" == v or "rrs" in v or "date" in v or "purpose" in v or "qual" in v) for v in row_lower if v):
                header_row_idx = i
                break

        header = raw_rows[header_row_idx]

        # Map header names to column indices (case-insensitive, partial match)
        COL_ALIASES = {
            "ww":             ["ww", "work week"],
            "date_received":  ["date received", "date received at rellab", "date received at rel"],
            "rrs_no":         ["rrs no", "rrs no.", "rrs#", "rrs number"],
            "purpose":        ["purpose"],
            "qual_type":      ["qual type", "qual. type", "qualification type", "qual"],
            "customer":       ["customer", "cust"],
            "pkg_type":       ["pkg. type", "pkg type", "package type", "pkg"],
            "lc_bc":          ["l/c b/c", "lc bc", "l/c", "b/c", "lc/bc"],
            "rr_agile_no":    ["rr/agile no", "rr/agile no.", "agile no", "agile no.", "rr no", "lot no"],
            "test_level":     ["test level", "test lvl"],
            "qty":            ["qty", "quantity"],
            "num_days":       ["# of days", "no. of days", "num days", "days"],
            "num_legs":       ["# of legs", "no. of legs", "num legs", "legs"],
            "est_start":      ["estimated date/time of start", "est. date/time of start", "est start", "estimated start"],
            "est_completion": ["estimated date of completion", "est. date of completion", "est completion", "estimated completion"],
            "recommit":       ["re-commit", "recommit", "re commit"],
            "planner_remarks":["planner remarks", "remarks", "planner note", "notes"],
        }
        header_map: Dict[str, int] = {}
        for j, hdr in enumerate(header):
            if hdr is None:
                continue
            h = str(hdr).strip().lower()
            for field, aliases in COL_ALIASES.items():
                if field not in header_map and any(a in h for a in aliases):
                    header_map[field] = j

        # Positional fallback if mapping incomplete
        default_positions = ["ww", "date_received", "rrs_no", "purpose", "qual_type", "customer",
                             "pkg_type", "lc_bc", "rr_agile_no", "test_level", "qty",
                             "num_days", "num_legs", "est_start", "est_completion", "recommit", "planner_remarks"]
        for i, field in enumerate(default_positions):
            if field not in header_map and i < len(header):
                header_map[field] = i

        def get_col(row, field):
            idx = header_map.get(field)
            if idx is None or idx >= len(row):
                return ""
            v = row[idx]
            return str(v).strip() if v is not None else ""

        records = []
        uploaded_at = datetime.now(timezone.utc).isoformat()
        for row in raw_rows[header_row_idx + 1:]:
            if not row or all(v is None or str(v).strip() == "" for v in row):
                continue
            rrs_no = get_col(row, "rrs_no")
            ww = get_col(row, "ww")
            if not rrs_no and not ww:
                continue
            records.append((
                ww,
                get_col(row, "date_received"),
                rrs_no,
                get_col(row, "purpose"),
                get_col(row, "qual_type"),
                get_col(row, "customer"),
                get_col(row, "pkg_type"),
                get_col(row, "lc_bc"),
                get_col(row, "rr_agile_no"),
                get_col(row, "test_level"),
                get_col(row, "qty"),
                get_col(row, "num_days"),
                get_col(row, "num_legs"),
                get_col(row, "est_start"),
                get_col(row, "est_completion"),
                get_col(row, "recommit"),
                get_col(row, "planner_remarks"),
                uploaded_at,
            ))

        if not records:
            raise HTTPException(status_code=422, detail="No data rows found in the Excel file")

        db = await get_db()
        try:
            await db.execute("DELETE FROM masterlist_2026")
            await db.executemany(
                "INSERT INTO masterlist_2026 (ww, date_received, rrs_no, purpose, qual_type, customer, "
                "pkg_type, lc_bc, rr_agile_no, test_level, qty, num_days, num_legs, "
                "est_start, est_completion, recommit, planner_remarks, uploaded_at) "
                "VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
                records,
            )
            await db.commit()
        finally:
            await db.close()

        return {"imported": len(records), "message": f"Successfully imported {len(records)} records"}

    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Masterlist upload error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to parse Excel file: {str(e)}")


@api_router.get("/masterlist")
async def get_masterlist(current_user: User = Depends(get_current_user)):
    """Get all Masterlist 2026 records."""
    db = await get_db()
    try:
        cursor = await db.execute(
            "SELECT id, ww, date_received, rrs_no, purpose, qual_type, customer, "
            "pkg_type, lc_bc, rr_agile_no, test_level, qty, num_days, num_legs, "
            "est_start, est_completion, recommit, planner_remarks, uploaded_at "
            "FROM masterlist_2026 ORDER BY id"
        )
        rows = await cursor.fetchall()
        return [
            {
                "id": r[0], "ww": r[1], "date_received": r[2], "rrs_no": r[3],
                "purpose": r[4], "qual_type": r[5], "customer": r[6],
                "pkg_type": r[7], "lc_bc": r[8], "rr_agile_no": r[9],
                "test_level": r[10], "qty": r[11], "num_days": r[12],
                "num_legs": r[13], "est_start": r[14], "est_completion": r[15],
                "recommit": r[16], "planner_remarks": r[17], "uploaded_at": r[18],
            }
            for r in rows
        ]
    finally:
        await db.close()


@api_router.delete("/masterlist")
async def clear_masterlist(current_user: User = Depends(require_role([UserRole.ADMIN]))):
    """Clear all Masterlist 2026 records."""
    db = await get_db()
    try:
        await db.execute("DELETE FROM masterlist_2026")
        await db.commit()
        return {"message": "Masterlist cleared successfully"}
    finally:
        await db.close()


class MasterlistRecord(BaseModel):
    ww: Optional[str] = None
    date_received: Optional[str] = None
    rrs_no: Optional[str] = None
    purpose: Optional[str] = None
    qual_type: Optional[str] = None
    customer: Optional[str] = None
    pkg_type: Optional[str] = None
    lc_bc: Optional[str] = None
    rr_agile_no: Optional[str] = None
    test_level: Optional[str] = None
    qty: Optional[str] = None
    num_days: Optional[str] = None
    num_legs: Optional[str] = None
    est_start: Optional[str] = None
    est_completion: Optional[str] = None
    recommit: Optional[str] = None
    planner_remarks: Optional[str] = None


@api_router.post("/masterlist", status_code=201)
async def add_masterlist_record(
    data: MasterlistRecord,
    current_user: User = Depends(require_role([UserRole.ADMIN, UserRole.PLANNER])),
):
    """Add a single record to Masterlist 2026."""
    db = await get_db()
    try:
        uploaded_at = datetime.now(timezone.utc).isoformat()
        cursor = await db.execute(
            "INSERT INTO masterlist_2026 (ww, date_received, rrs_no, purpose, qual_type, customer, "
            "pkg_type, lc_bc, rr_agile_no, test_level, qty, num_days, num_legs, "
            "est_start, est_completion, recommit, planner_remarks, uploaded_at) "
            "VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
            (data.ww, data.date_received, data.rrs_no, data.purpose, data.qual_type,
             data.customer, data.pkg_type, data.lc_bc, data.rr_agile_no, data.test_level,
             data.qty, data.num_days, data.num_legs, data.est_start, data.est_completion,
             data.recommit, data.planner_remarks, uploaded_at),
        )
        await db.commit()
        return {"id": cursor.lastrowid, "message": "Record added successfully"}
    finally:
        await db.close()


@api_router.put("/masterlist/{record_id}")
async def update_masterlist_record(
    record_id: int,
    data: MasterlistRecord,
    current_user: User = Depends(require_role([UserRole.ADMIN, UserRole.PLANNER])),
):
    """Update a single Masterlist 2026 record."""
    db = await get_db()
    try:
        await db.execute(
            "UPDATE masterlist_2026 SET ww=?, date_received=?, rrs_no=?, purpose=?, qual_type=?, "
            "customer=?, pkg_type=?, lc_bc=?, rr_agile_no=?, test_level=?, qty=?, num_days=?, "
            "num_legs=?, est_start=?, est_completion=?, recommit=?, planner_remarks=? WHERE id=?",
            (data.ww, data.date_received, data.rrs_no, data.purpose, data.qual_type,
             data.customer, data.pkg_type, data.lc_bc, data.rr_agile_no, data.test_level,
             data.qty, data.num_days, data.num_legs, data.est_start, data.est_completion,
             data.recommit, data.planner_remarks, record_id),
        )
        await db.commit()
        return {"message": "Record updated successfully"}
    finally:
        await db.close()


@api_router.delete("/masterlist/{record_id}")
async def delete_masterlist_record(
    record_id: int,
    current_user: User = Depends(require_role([UserRole.ADMIN, UserRole.PLANNER])),
):
    """Delete a single Masterlist 2026 record."""
    db = await get_db()
    try:
        await db.execute("DELETE FROM masterlist_2026 WHERE id=?", (record_id,))
        await db.commit()
        return {"message": "Record deleted successfully"}
    finally:
        await db.close()


class MasterlistPlannerFields(BaseModel):
    test_level: Optional[str] = None
    ml_qty: Optional[str] = None
    num_days: Optional[str] = None
    num_legs: Optional[str] = None
    recommit: Optional[str] = None
    planner_est_start: Optional[str] = None
    planner_est_end: Optional[str] = None
    planner_note: Optional[str] = None


def _ml_compute_ww(date_str: str) -> str:
    """Return work-week string like 'ww01' from a date string.
    WW01 = Jan 1 through the first Sunday of the year.
    WW02+ = each subsequent Monday-Sunday block.
    Example: 2026 Jan 1 is Thursday, so WW01=Jan 1-4, WW02=Jan 5-11, etc.
    """
    if not date_str:
        return ""
    try:
        dt = datetime.fromisoformat(date_str.replace('Z', '+00:00'))
        d = dt.date() if hasattr(dt, 'date') else dt
        jan1 = d.replace(month=1, day=1)
        # Days from Jan 1 to the first Sunday (Mon=0 ... Sun=6)
        days_to_sunday = (6 - jan1.weekday()) % 7
        first_sunday = jan1 + timedelta(days=days_to_sunday)
        if d <= first_sunday:
            week_num = 1
        else:
            first_monday = first_sunday + timedelta(days=1)
            week_num = 2 + (d - first_monday).days // 7
        return f"ww{week_num:02d}"
    except Exception:
        return ""


def _ml_compute_lc_bc(ball_pitch, ball_count, lead_pitch, lead_count) -> str:
    """Auto-build L/C B/C string.
    Format: B/C:{ball_pitch}/{ball_count} | L/C:{lead_pitch}/{lead_count}
    """
    parts = []
    if ball_pitch is not None or ball_count is not None:
        bp = f"{float(ball_pitch):.2f}" if ball_pitch is not None else "—"
        bc = str(int(ball_count)) if ball_count is not None else "—"
        parts.append(f"B/C:{bp}/{bc}")
    if lead_pitch is not None or lead_count is not None:
        lp = f"{float(lead_pitch):.2f}" if lead_pitch is not None else "—"
        lc = str(int(lead_count)) if lead_count is not None else "—"
        parts.append(f"L/C:{lp}/{lc}")
    return " | ".join(parts)


@api_router.get("/masterlist/requests")
async def get_masterlist_requests(current_user: User = Depends(get_current_user)):
    """Return all requests formatted as masterlist rows for the planning monitor."""
    db = await get_db()
    try:
        cursor = await db.execute(
            "SELECT r.id, r.request_number, r.classification, r.customer, r.pkg_info, r.lot_no, "
            "r.date_ltc, r.created_at, r.purpose, r.total_ss, "
            "r.planner_est_start, r.planner_est_end, r.planner_note, "
            "r.ww, r.lc_bc, r.ml_qty, r.num_days, r.num_legs, r.recommit, r.status, "
            "(SELECT ps.step_name FROM process_steps ps "
            " WHERE ps.request_id = r.id AND ps.leg = 1 AND ps.step_number = r.current_step "
            " LIMIT 1) as current_step_name, "
            "r.approved_at, r.ball_pitch, r.ball_count, r.lead_pitch, r.lead_count, r.test_level "
            "FROM requests r ORDER BY r.created_at DESC"
        )
        rows = await cursor.fetchall()
        result = []
        for r in rows:
            # r[20]=current_step_name  r[21]=approved_at  r[22]=ball_pitch
            # r[23]=ball_count  r[24]=lead_pitch  r[25]=lead_count  r[26]=test_level
            approved_at = r[21] or ""
            # Resolved date used for display (date_received column)
            date_received = approved_at or r[6] or r[7] or ""

            # WW: use stored value; if empty, auto-derive from available dates
            # Try approved_at → date_ltc → created_at, using first that parses
            stored_ww = r[13] or ""
            if stored_ww:
                ww = stored_ww
            else:
                ww = ""
                for _d in filter(None, [r[21], r[6], r[7]]):
                    _w = _ml_compute_ww(str(_d))
                    if _w:
                        ww = _w
                        break

            # L/C B/C: use stored value; if empty, auto-derive from ball/lead fields
            stored_lc_bc = r[14] or ""
            lc_bc = stored_lc_bc if stored_lc_bc else _ml_compute_lc_bc(r[22], r[23], r[24], r[25])

            # Test Level: use stored value; if empty, fall back to current step name
            test_level = r[26] or r[20] or ""

            result.append({
                "id": r[0],
                "request_number": r[1],
                "ww": ww,
                "date_received": date_received,
                "rrs_no": r[1] or "",
                "purpose": r[8] or "",
                "qual_type": r[2] or "",
                "customer": r[3] or "",
                "pkg_type": r[4] or "",
                "lc_bc": lc_bc,
                "rr_agile_no": r[5] or "",
                "test_level": test_level,
                "qty": r[15] or r[9] or "",
                "num_days": r[16] or "",
                "num_legs": r[17] or "",
                "est_start": r[10] or "",
                "est_completion": r[11] or "",
                "recommit": r[18] or "",
                "planner_remarks": r[12] or "",
                "status": r[19] or "",
            })
        return result
    finally:
        await db.close()


@api_router.patch("/masterlist/requests/{request_id}")
async def update_request_masterlist_fields(
    request_id: str,
    data: MasterlistPlannerFields,
    current_user: User = Depends(require_role([UserRole.ADMIN, UserRole.PLANNER])),
):
    """Update planner-specific masterlist fields on a request."""
    db = await get_db()
    try:
        db_field_map = {
            'test_level': data.test_level,
            'ml_qty': data.ml_qty, 'num_days': data.num_days, 'num_legs': data.num_legs,
            'recommit': data.recommit, 'planner_est_start': data.planner_est_start,
            'planner_est_end': data.planner_est_end, 'planner_note': data.planner_note,
        }
        to_update = {col: val for col, val in db_field_map.items() if val is not None}
        if not to_update:
            return {"message": "Nothing to update"}
        set_clause = ', '.join(f"{col}=?" for col in to_update)
        values = list(to_update.values()) + [request_id]
        await db.execute(f"UPDATE requests SET {set_clause} WHERE id=?", values)
        await db.commit()
        return {"message": "Updated successfully"}
    finally:
        await db.close()


async def declined_user_cleanup():
    """Background task: permanently delete users whose status has been 'declined'
    for 5 or more minutes."""
    while True:
        await asyncio.sleep(60)  # check every minute
        try:
            db = await get_db()
            try:
                cursor = await db.execute(
                    """SELECT id, username FROM users
                       WHERE user_status = 'declined'
                         AND declined_at IS NOT NULL
                         AND datetime(declined_at, '+5 minutes') <= datetime('now')"""
                )
                rows = await cursor.fetchall()
                for row in rows:
                    uid = row[0]
                    uname = row[1]
                    await db.execute("DELETE FROM users WHERE id = ?", (uid,))
                    logger.info(f"Auto-deleted declined user id={uid} username={uname} (5-min grace expired)")
                if rows:
                    await db.commit()
            finally:
                await db.close()
        except Exception as e:
            logger.error(f"declined_user_cleanup error: {e}")


async def monthly_backup_scheduler():
    """Background task that creates an automatic backup on the 1st of each month."""
    while True:
        now = datetime.now()
        # Calculate next 1st of month at 2:00 AM
        if now.day == 1 and now.hour < 2:
            next_run = now.replace(hour=2, minute=0, second=0, microsecond=0)
        else:
            if now.month == 12:
                next_run = now.replace(year=now.year + 1, month=1, day=1, hour=2, minute=0, second=0, microsecond=0)
            else:
                next_run = now.replace(month=now.month + 1, day=1, hour=2, minute=0, second=0, microsecond=0)

        wait_seconds = (next_run - now).total_seconds()
        logging.info(f"Next auto-backup scheduled for {next_run.isoformat()} (in {wait_seconds/3600:.1f} hours)")
        await asyncio.sleep(wait_seconds)

        try:
            # Create automatic backup
            create_backup_sync(is_auto=True)
            
            # Keep only last 12 monthly backups in Auto_Backups folder
            # Collect all auto backups with their creation dates
            auto_backups = []
            for year_folder in AUTO_BACKUP_DIR.iterdir():
                if year_folder.is_dir() and year_folder.name.isdigit():
                    for month_folder in year_folder.iterdir():
                        if month_folder.is_dir():
                            for file in month_folder.glob("rel_database_backup_*.xlsx"):
                                auto_backups.append((file.stat().st_mtime, file, month_folder, year_folder))
            
            # Sort by modification time (oldest first)
            auto_backups.sort(key=lambda x: x[0])
            
            # Remove oldest backups if more than 12
            while len(auto_backups) > 12:
                _, oldest_file, month_folder, year_folder = auto_backups.pop(0)
                oldest_file.unlink()
                logging.info(f"Old auto-backup removed: {oldest_file.name}")
                # Clean up empty folders
                if not any(month_folder.iterdir()):
                    month_folder.rmdir()
                if not any(year_folder.iterdir()):
                    year_folder.rmdir()
        except Exception as e:
            logging.error(f"Auto-backup failed: {e}")


# ========================
# Loading / Unloading Monitor
# ========================
MONITORED_STEP_NAMES = [
    'reliability test', 't&h soak', 'forced convection reflow (fcr)',
    'preconditioning (precon)', 'temperature cycle',
    'moisture resistance test', 'bake', 'dry bake',
]

@api_router.get("/loading-unloading")
async def get_loading_unloading(current_user: User = Depends(get_current_user)):
    """Return active + recently completed monitored steps with their request info."""
    placeholders = ','.join('?' * len(MONITORED_STEP_NAMES))
    query = f"""
        SELECT
            ps.id, ps.request_id, ps.step_number, ps.step_name, ps.leg,
            ps.status, ps.started_at, ps.completed_at, ps.machine_no, ps.rack_no,
            ps.custom_fields,
            r.request_number, r.device_name, r.customer, r.lot_no,
            COALESCE(e.name, '') AS employee_name,
            COALESCE(m.description, '') AS machine_desc
        FROM process_steps ps
        JOIN requests r ON ps.request_id = r.id
        LEFT JOIN employees e ON ps.operator_id = e.id
        LEFT JOIN machines m ON LOWER(ps.machine_no) = LOWER(m.machine_no)
        WHERE LOWER(ps.step_name) IN ({placeholders})
          AND r.status NOT IN ('discontinued')
          AND ps.status IN ('in_progress', 'completed')
          AND ps.started_at IS NOT NULL
        ORDER BY
            CASE ps.status WHEN 'in_progress' THEN 0 ELSE 1 END,
            ps.started_at DESC
    """
    db = await get_db()
    try:
        cursor = await db.execute(query, MONITORED_STEP_NAMES)
        rows = await cursor.fetchall()
        result = []
        for row in rows:
            cf = {}
            try:
                cf = json.loads(row[10]) if row[10] else {}
            except Exception:
                pass
            result.append({
                'id': row[0],
                'request_id': row[1],
                'step_number': row[2],
                'step_name': row[3],
                'leg': row[4],
                'status': row[5],
                'started_at': row[6],
                'completed_at': row[7],
                'machine_no': row[8] or '',
                'rack_no': row[9] or '',
                'test_item': cf.get('test_item', ''),
                'test_condition': cf.get('test_condition', ''),
                'request_number': row[11],
                'device_name': row[12] or '',
                'customer': row[13] or '',
                'lot_no': row[14] or '',
                'employee_name': row[15] or '',
                'machine_desc': row[16] or '',
            })
        return result
    finally:
        await db.close()


@api_router.get("/loading-unloading/export")
async def export_loading_unloading(
    step_name: Optional[str] = None,
    current_user: User = Depends(get_current_user)
):
    """Generate an Excel report for loading/unloading operations."""
    placeholders = ','.join('?' * len(MONITORED_STEP_NAMES))
    conditions = [f"LOWER(ps.step_name) IN ({placeholders})"]
    params: list = list(MONITORED_STEP_NAMES)
    if step_name:
        conditions.append("LOWER(ps.step_name) = LOWER(?)")
        params.append(step_name)
    query = f"""
        SELECT ps.step_name, ps.leg, ps.step_number, ps.status,
               ps.started_at, ps.completed_at, ps.machine_no, ps.rack_no,
               r.request_number, r.device_name, r.customer, r.lot_no,
               COALESCE(e.name, '') AS employee_name,
               COALESCE(m.description, '') AS machine_desc
        FROM process_steps ps
        JOIN requests r ON ps.request_id = r.id
        LEFT JOIN employees e ON ps.operator_id = e.id
        LEFT JOIN machines m ON LOWER(ps.machine_no) = LOWER(m.machine_no)
        WHERE {' AND '.join(conditions)}
          AND r.status NOT IN ('discontinued')
          AND ps.status IN ('in_progress', 'completed')
          AND ps.started_at IS NOT NULL
        ORDER BY ps.started_at DESC
    """
    db = await get_db()
    try:
        cursor = await db.execute(query, params)
        rows = await cursor.fetchall()
    finally:
        await db.close()

    wb = Workbook()
    ws = wb.active
    ws.title = "L/U Monitor Report"
    header_fill = PatternFill(start_color="1E3A5F", end_color="1E3A5F", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF")
    thin_border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )
    headers = [
        "Request #", "Device", "Customer", "Lot #",
        "Step Name", "Leg", "Step #", "Status",
        "Machine No", "Machine Name", "Rack No", "Operator",
        "Started At", "Completed At"
    ]
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = thin_border
    ws.row_dimensions[1].height = 18

    alt_fill = PatternFill(start_color="EFF3FB", end_color="EFF3FB", fill_type="solid")
    for i, row in enumerate(rows, 2):
        data = [
            row[8], row[9], row[10], row[11],  # request_number, device_name, customer, lot_no
            row[0], row[1], row[2], row[3],    # step_name, leg, step_number, status
            row[6], row[13], row[7],           # machine_no, machine_desc, rack_no
            row[12], row[4], row[5],           # employee_name, started_at, completed_at
        ]
        fill = alt_fill if i % 2 == 0 else None
        for col, val in enumerate(data, 1):
            cell = ws.cell(row=i, column=col, value=val)
            cell.border = thin_border
            if fill:
                cell.fill = fill
    for col_cells in ws.columns:
        max_len = max((len(str(c.value or '')) for c in col_cells), default=10)
        ws.column_dimensions[col_cells[0].column_letter].width = min(max_len + 3, 42)

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    ts = datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S')
    fname = f"LU_Report_{ts}.xlsx"
    if step_name:
        safe = step_name.replace('/', '_').replace(' ', '_')
        fname = f"LU_{safe}_{ts[:8]}.xlsx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={fname}"}
    )


@api_router.get("/loading-unloading/history")
async def get_loading_unloading_history(
    machine: Optional[str] = None,
    current_user: User = Depends(get_current_user)
):
    """Return all historical loading/unloading records (completed + in_progress), searchable by machine."""
    placeholders = ','.join('?' * len(MONITORED_STEP_NAMES))
    conditions = [f"LOWER(ps.step_name) IN ({placeholders})"]
    params: list = list(MONITORED_STEP_NAMES)
    conditions.append("r.status NOT IN ('discontinued')")
    conditions.append("ps.status IN ('in_progress', 'completed')")
    if machine:
        conditions.append("(LOWER(ps.machine_no) LIKE LOWER(?) OR LOWER(COALESCE(m.description, '')) LIKE LOWER(?))")
        params.extend([f"%{machine}%", f"%{machine}%"])
    query = f"""
        SELECT ps.id, ps.step_name, ps.leg, ps.step_number, ps.status,
               ps.started_at, ps.completed_at, ps.machine_no, ps.rack_no,
               r.id AS request_id, r.request_number, r.device_name, r.customer, r.lot_no,
               COALESCE(e.name, '') AS employee_name,
               COALESCE(m.description, '') AS machine_desc
        FROM process_steps ps
        JOIN requests r ON ps.request_id = r.id
        LEFT JOIN employees e ON ps.operator_id = e.id
        LEFT JOIN machines m ON LOWER(ps.machine_no) = LOWER(m.machine_no)
        WHERE {' AND '.join(conditions)}
        ORDER BY ps.started_at DESC
        LIMIT 500
    """
    db = await get_db()
    try:
        cursor = await db.execute(query, params)
        rows = await cursor.fetchall()
        result = []
        for row in rows:
            result.append({
                "id": row[0],
                "step_name": row[1],
                "leg": row[2],
                "step_number": row[3],
                "status": row[4],
                "started_at": row[5],
                "completed_at": row[6],
                "machine_no": row[7],
                "rack_no": row[8],
                "request_id": row[9],
                "request_number": row[10],
                "device_name": row[11],
                "customer": row[12],
                "lot_no": row[13],
                "employee_name": row[14],
                "machine_desc": row[15],
            })
        return result
    finally:
        await db.close()


@api_router.get("/performance/employees")
async def get_employee_performance(
    days: int = 30,
    current_user: User = Depends(get_current_user)
):
    """Return performance stats per employee for the given period."""
    db = await get_db()
    try:
        cutoff = (datetime.now(timezone.utc) - timedelta(days=days)).isoformat()
        cursor = await db.execute("""
            SELECT ps.operator_id,
                   COALESCE(e.name, ps.operator_id) AS employee_name,
                   COUNT(*) AS steps_completed,
                   COUNT(DISTINCT DATE(ps.completed_at)) AS active_days,
                   COUNT(DISTINCT ps.request_id) AS requests_touched,
                   GROUP_CONCAT(DISTINCT ps.step_name) AS step_types,
                   GROUP_CONCAT(ps.completed_at) AS completed_dates
            FROM process_steps ps
            LEFT JOIN employees e ON ps.operator_id = e.id
            WHERE ps.status = 'completed'
              AND ps.completed_at >= ?
              AND ps.operator_id IS NOT NULL
              AND ps.operator_id != ''
            GROUP BY ps.operator_id
            ORDER BY steps_completed DESC
        """, (cutoff,))
        rows = await cursor.fetchall()
        result = []
        for row in rows:
            active_days = row[3] or 1
            # Compute work weeks from completed dates
            ww_set = set()
            if row[6]:
                for dt_str in row[6].split(','):
                    ww = _ml_compute_ww(dt_str.strip())
                    if ww:
                        ww_set.add(ww.upper())
            ww_list = sorted(ww_set)
            result.append({
                "operator_id": row[0],
                "employee_name": row[1],
                "steps_completed": row[2],
                "active_days": active_days,
                "steps_per_day": round(row[2] / active_days, 1),
                "requests_touched": row[4],
                "step_types": row[5].split(',') if row[5] else [],
                "work_weeks": ww_list,
            })
        return result
    finally:
        await db.close()


@api_router.get("/performance/daily")
async def get_daily_performance(
    days: int = 30,
    current_user: User = Depends(get_current_user)
):
    """Return per-day performance breakdown."""
    db = await get_db()
    try:
        cutoff = (datetime.now(timezone.utc) - timedelta(days=days)).isoformat()
        cursor = await db.execute("""
            SELECT DATE(ps.completed_at) AS day,
                   ps.operator_id,
                   COALESCE(e.name, ps.operator_id) AS employee_name,
                   COUNT(*) AS steps_completed,
                   GROUP_CONCAT(DISTINCT ps.step_name) AS step_types,
                   COUNT(DISTINCT ps.request_id) AS requests_touched
            FROM process_steps ps
            LEFT JOIN employees e ON ps.operator_id = e.id
            WHERE ps.status = 'completed'
              AND ps.completed_at >= ?
              AND ps.operator_id IS NOT NULL
              AND ps.operator_id != ''
            GROUP BY DATE(ps.completed_at), ps.operator_id
            ORDER BY day DESC, steps_completed DESC
        """, (cutoff,))
        rows = await cursor.fetchall()
        result = []
        for row in rows:
            ww = _ml_compute_ww(row[0]) if row[0] else ''
            result.append({
                "date": row[0],
                "work_week": ww.upper() if ww else '',
                "operator_id": row[1],
                "employee_name": row[2],
                "steps_completed": row[3],
                "step_types": row[4].split(',') if row[4] else [],
                "requests_touched": row[5],
            })
        return result
    finally:
        await db.close()


# ========================
# RELMON – Reliability Monitor Summary
# ========================

class RelMonSaveRequest(BaseModel):
    site: str
    sheet: str
    rows: List[List[Optional[Any]]]
    merges: Optional[List[Dict[str, int]]] = None
    form_data: Dict[str, Any] = Field(default_factory=dict)

# Paths to the quarterly Rel Monitor Excel files (root of project, beside backend/)
_RELMON_FILES: dict[str, Path] = {
    "ATP1": ROOT_DIR.parent / "ATP1 Q4'2025 Rel Monitor Summary 011326.xlsx",
    "ATP3": ROOT_DIR.parent / "ATP3 Q4'2025 Rel Monitor Summary 011326.xlsx",
}

# Some users refer to ATP3 workbook content as ATP2 in the UI/process.
_RELMON_SITE_ALIASES: dict[str, str] = {
    "ATP2": "ATP3",
}

# Simple in-memory cache so repeated requests don't re-read the file
_relmon_cache: dict[tuple[str, str], dict] = {}


def _safe_relmon_json_load(raw: Optional[str], fallback: Any) -> Any:
    if not raw:
        return fallback
    try:
        return json.loads(raw)
    except Exception:
        return fallback


def _coerce_relmon_cell(value: Any) -> Any:
    if value is None:
        return None
    if isinstance(value, (int, float)):
        return value

    s = str(value).strip()
    if s == "":
        return None

    try:
        if "." not in s and "e" not in s.lower():
            return int(s)
    except Exception:
        pass

    try:
        return float(s)
    except Exception:
        return s


def _normalize_relmon_rows(rows: Any) -> tuple[list[list[Any]], int]:
    normalized: list[list[Any]] = []
    max_cols = 0

    if not isinstance(rows, list):
        return normalized, 0

    for row in rows:
        if isinstance(row, list):
            values = row
        elif isinstance(row, tuple):
            values = list(row)
        else:
            values = [row]

        cleaned = [_coerce_relmon_cell(v) for v in values]
        normalized.append(cleaned)
        if len(cleaned) > max_cols:
            max_cols = len(cleaned)

    if max_cols > 0:
        for row in normalized:
            if len(row) < max_cols:
                row.extend([None] * (max_cols - len(row)))

    return normalized, max_cols


def _parse_relmon_sheet(site: str, sheet_name: str) -> dict:
    """Parse a single sheet from the RELMON Excel file and return structured data."""
    cache_key = (site, sheet_name)
    if cache_key in _relmon_cache:
        return _relmon_cache[cache_key]

    path = _RELMON_FILES.get(site)
    if path is None or not path.exists():
        raise FileNotFoundError(f"RELMON file for site '{site}' not found")
    wb = None
    try:
        wb = load_workbook(str(path), data_only=True)
        if sheet_name not in wb.sheetnames:
            raise ValueError(f"Sheet '{sheet_name}' not found in {site}")

        ws = wb[sheet_name]

        # Collect raw cell values (None stays None, everything else becomes str or number)
        rows: list[list] = []
        for row in ws.iter_rows(values_only=True):
            rows.append([
                (v if isinstance(v, (int, float)) else (str(v) if v is not None else None))
                for v in row
            ])

        # Collect merged-cell regions so the frontend can render colspan/rowspan
        merges: list[dict] = []
        for rng in ws.merged_cells.ranges:
            merges.append({
                "min_row": rng.min_row - 1,  # convert to 0-based
                "max_row": rng.max_row - 1,
                "min_col": rng.min_col - 1,
                "max_col": rng.max_col - 1,
            })

        result = {
            "site": site,
            "sheet": sheet_name,
            "num_rows": len(rows),
            "num_cols": ws.max_column,
            "rows": rows,
            "merges": merges,
        }
        _relmon_cache[cache_key] = result
        return result
    except PermissionError as e:
        raise PermissionError(
            f"RELMON workbook is locked by another process for site '{site}': {path.name}"
        ) from e
    finally:
        if wb is not None:
            wb.close()


async def _get_saved_relmon_record(site: str, sheet: str) -> Optional[aiosqlite.Row]:
    db = await aiosqlite.connect(DB_PATH)
    db.row_factory = aiosqlite.Row
    try:
        cursor = await db.execute(
            "SELECT rows_json, merges_json, form_json, updated_at, updated_by "
            "FROM relmon_sheet_data WHERE site = ? AND sheet = ?",
            (site, sheet),
        )
        return await cursor.fetchone()
    finally:
        await db.close()


def _resolve_relmon_site(site: str) -> str:
    return _RELMON_SITE_ALIASES.get(site, site)


def _relmon_sheet_device_type(sheet_name: str) -> str:
    s = (sheet_name or "").strip()
    if not s:
        return ""
    m = _re.match(r"^(.+?)\s*\((.+)\)$", s)
    if m:
        return m.group(1).strip()
    return s


async def _get_saved_relmon_sheet_names(site: str) -> list[str]:
    db = await aiosqlite.connect(DB_PATH)
    try:
        cursor = await db.execute(
            "SELECT DISTINCT sheet FROM relmon_sheet_data WHERE site = ? ORDER BY sheet",
            (site,),
        )
        rows = await cursor.fetchall()
        return [r[0] for r in rows if r and r[0]]
    finally:
        await db.close()


async def _load_relmon_sheet_names_for_site(site: str) -> tuple[list[str], str]:
    """Load sheet names from workbook; if locked/unavailable, fallback to saved DB names."""
    resolved = _resolve_relmon_site(site)
    path = _RELMON_FILES.get(resolved)
    if path is None or not path.exists():
        saved = await _get_saved_relmon_sheet_names(resolved)
        return saved, "saved_db"

    wb = None
    try:
        wb = load_workbook(str(path), data_only=True, read_only=True)
        return list(wb.sheetnames), "workbook"
    except PermissionError:
        saved = await _get_saved_relmon_sheet_names(resolved)
        return saved, "saved_db"
    except Exception:
        saved = await _get_saved_relmon_sheet_names(resolved)
        return saved, "saved_db"
    finally:
        if wb is not None:
            wb.close()


@api_router.get("/relmon/sheets")
async def get_relmon_sheets():
    """Return the available sites and their sheet names."""
    out: dict[str, list[str]] = {}
    for site, path in _RELMON_FILES.items():
        if path.exists():
            try:
                wb = load_workbook(str(path), data_only=True, read_only=True)
                out[site] = list(wb.sheetnames)
                wb.close()
            except Exception:
                out[site] = []
        else:
            out[site] = []
    return out


@api_router.get("/relmon/device-types")
async def get_relmon_device_types(site: Optional[str] = None, order: str = "asc"):
    """Return sortable RELMON device-type lists grouped like the workbook sheets.

    Query params:
    - site: ATP1 | ATP2 | ATP3 (optional; default returns ATP1 and ATP2)
    - order: asc | desc (sort order by device type)
    """
    order_norm = (order or "asc").strip().lower()
    if order_norm not in ("asc", "desc"):
        raise HTTPException(status_code=400, detail="order must be 'asc' or 'desc'")

    requested_sites = [site] if site else ["ATP1", "ATP2"]
    response_sites: dict[str, Any] = {}

    for requested in requested_sites:
        resolved = _resolve_relmon_site(requested)
        if resolved not in _RELMON_FILES:
            raise HTTPException(status_code=404, detail=f"Site '{requested}' not found")

        sheet_names, source = await _load_relmon_sheet_names_for_site(requested)

        grouped: dict[str, list[str]] = {}
        for s in sheet_names:
            dev_type = _relmon_sheet_device_type(s)
            if not dev_type:
                continue
            grouped.setdefault(dev_type, []).append(s)

        for k in list(grouped.keys()):
            grouped[k] = sorted(grouped[k])

        reverse = order_norm == "desc"
        sorted_types = sorted(grouped.keys(), reverse=reverse)
        grouped_sheets = [{"device_type": t, "sheets": grouped[t]} for t in sorted_types]

        response_sites[requested] = {
            "resolved_site": resolved,
            "source": source,
            "count": len(sorted_types),
            "device_types": sorted_types,
            "grouped_sheets": grouped_sheets,
        }

    return {
        "order": order_norm,
        "sites": response_sites,
    }


@api_router.get("/relmon/data")
async def get_relmon_data(site: str, sheet: str):
    """Return header rows, merge info, and data rows for a specific RELMON sheet."""
    if site not in _RELMON_FILES:
        raise HTTPException(status_code=404, detail=f"Site '{site}' not found")

    parse_error: Optional[Exception] = None
    payload: Optional[dict] = None
    try:
        payload = _parse_relmon_sheet(site, sheet)
    except Exception as e:
        parse_error = e

    try:
        saved = await _get_saved_relmon_record(site, sheet)

        if payload is None and saved:
            saved_rows = _safe_relmon_json_load(saved["rows_json"], [])
            saved_merges = _safe_relmon_json_load(saved["merges_json"], [])
            saved_form = _safe_relmon_json_load(saved["form_json"], {})

            rows, max_cols = _normalize_relmon_rows(saved_rows)
            payload = {
                "site": site,
                "sheet": sheet,
                "num_rows": len(rows),
                "num_cols": max_cols,
                "rows": rows,
                "merges": saved_merges if isinstance(saved_merges, list) else [],
                "form_data": saved_form if isinstance(saved_form, dict) else {},
                "updated_at": saved["updated_at"],
                "updated_by": saved["updated_by"],
                "source_unavailable": True,
            }
            return payload

        if payload is None:
            if isinstance(parse_error, FileNotFoundError):
                raise HTTPException(status_code=404, detail=str(parse_error))
            if isinstance(parse_error, ValueError):
                raise HTTPException(status_code=404, detail=str(parse_error))
            if isinstance(parse_error, PermissionError):
                raise HTTPException(
                    status_code=423,
                    detail=(
                        f"{parse_error}. Close the workbook in Excel and retry, "
                        "or save once after source access is restored."
                    ),
                )
            raise HTTPException(status_code=500, detail=f"Failed to read sheet: {parse_error}")

        if saved:
            saved_rows = _safe_relmon_json_load(saved["rows_json"], payload["rows"])
            saved_merges = _safe_relmon_json_load(saved["merges_json"], payload["merges"])
            saved_form = _safe_relmon_json_load(saved["form_json"], {})

            rows, max_cols = _normalize_relmon_rows(saved_rows)
            if rows:
                payload["rows"] = rows
                payload["num_rows"] = len(rows)
                payload["num_cols"] = max_cols

            if isinstance(saved_merges, list):
                payload["merges"] = saved_merges

            payload["form_data"] = saved_form if isinstance(saved_form, dict) else {}
            payload["updated_at"] = saved["updated_at"]
            payload["updated_by"] = saved["updated_by"]
        else:
            payload["form_data"] = {}
            payload["updated_at"] = None
            payload["updated_by"] = None

        return payload
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to read sheet: {e}")


@api_router.put("/relmon/data")
async def save_relmon_data(payload: RelMonSaveRequest, current_user: User = Depends(get_current_user)):
    """Save editable RELMON rows and input-form data for a specific site/sheet."""
    if payload.site not in _RELMON_FILES:
        raise HTTPException(status_code=404, detail=f"Site '{payload.site}' not found")
    if not payload.sheet or not payload.sheet.strip():
        raise HTTPException(status_code=400, detail="Sheet cannot be empty")

    # Best-effort source validation while allowing saves when the workbook is temporarily locked.
    try:
        _parse_relmon_sheet(payload.site, payload.sheet)
    except PermissionError:
        existing = await _get_saved_relmon_record(payload.site, payload.sheet)
        if not existing:
            raise HTTPException(
                status_code=423,
                detail="Source workbook is locked and this sheet has no saved baseline yet",
            )
    except FileNotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except ValueError:
        existing = await _get_saved_relmon_record(payload.site, payload.sheet)
        if not existing:
            raise HTTPException(status_code=404, detail=f"Sheet '{payload.sheet}' not found in source workbook")

    rows, max_cols = _normalize_relmon_rows(payload.rows)
    if not rows or max_cols == 0:
        raise HTTPException(status_code=400, detail="Rows cannot be empty")

    merges = payload.merges if isinstance(payload.merges, list) else []
    form_data = payload.form_data if isinstance(payload.form_data, dict) else {}
    now = datetime.now(timezone.utc).isoformat()

    db = await aiosqlite.connect(DB_PATH)
    try:
        await db.execute(
            """
            INSERT INTO relmon_sheet_data (
                site, sheet, rows_json, merges_json, form_json, created_at, updated_at, updated_by
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            ON CONFLICT(site, sheet) DO UPDATE SET
                rows_json = excluded.rows_json,
                merges_json = excluded.merges_json,
                form_json = excluded.form_json,
                updated_at = excluded.updated_at,
                updated_by = excluded.updated_by
            """,
            (
                payload.site,
                payload.sheet,
                json.dumps(rows, ensure_ascii=True),
                json.dumps(merges, ensure_ascii=True),
                json.dumps(form_data, ensure_ascii=True),
                now,
                now,
                current_user.username,
            ),
        )
        await db.commit()
    finally:
        await db.close()

    return {
        "ok": True,
        "site": payload.site,
        "sheet": payload.sheet,
        "num_rows": len(rows),
        "num_cols": max_cols,
        "updated_at": now,
        "updated_by": current_user.username,
    }


# ========================
# App Setup
# ========================
app.include_router(api_router)

upload_dir = ROOT_DIR / "uploads"
upload_dir.mkdir(exist_ok=True)
app.mount("/uploads", StaticFiles(directory=str(upload_dir)), name="uploads")

app.add_middleware(
    CORSMiddleware, allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"], allow_headers=["*"],
)

# Setup Jinja2 template engine
TEMPLATES_DIR = ROOT_DIR / "templates"
if TEMPLATES_DIR.exists():
    templates = Jinja2Templates(directory=str(TEMPLATES_DIR))
    # Provide Flask-style get_flashed_messages for templates
    def _get_flashed_messages(**kwargs):
        return []
    templates.env.globals["get_flashed_messages"] = _get_flashed_messages

    # Custom filter: safely format dates that may be ISO strings or datetime objects
    def _datefmt(value, fmt='%Y-%m-%d'):
        if not value:
            return '-'
        if isinstance(value, str):
            try:
                value = datetime.fromisoformat(value)
            except (ValueError, TypeError):
                return value[:10] if len(value) >= 10 else value
        try:
            return value.strftime(fmt)
        except (AttributeError, ValueError):
            return str(value)
    templates.env.filters["datefmt"] = _datefmt
else:
    templates = None

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# ── Serve React SPA Frontend ─────────────────────────────────────────────────
# The React app is built into frontend/dist/. We serve it as a static SPA:
#   - /assets/* → static files (JS, CSS, images)
#   - All other non-/api/* routes → index.html (client-side routing)
FRONTEND_DIST = Path(__file__).resolve().parent.parent / "frontend" / "dist"
if FRONTEND_DIST.exists():
    app.mount("/assets", StaticFiles(directory=str(FRONTEND_DIST / "assets")), name="spa_assets")

    @app.get("/{full_path:path}", response_class=HTMLResponse)
    async def serve_spa(request: Request, full_path: str):
        """Serve the React SPA index.html for all non-API routes."""
        # Don't serve SPA for API routes or uploads (they're already handled above)
        index_file = FRONTEND_DIST / "index.html"
        return HTMLResponse(content=index_file.read_text(encoding="utf-8"))
else:
    logger.warning("React frontend build not found at %s — run 'npm run build' in frontend/", FRONTEND_DIST)

# ── (Legacy Jinja2 template routes removed — React SPA handles all pages) ────
if False:  # kept for reference; never executes
    from starlette.responses import RedirectResponse as _RR

    @app.exception_handler(_HTMLAuthRequired)
    async def _html_auth_redirect(request: Request, exc: _HTMLAuthRequired):
        """Redirect unauthenticated HTML page requests to the login screen."""
        response = _RR(url="/login", status_code=302)
        response.delete_cookie("access_token")
        return response

    @app.get("/login", response_class=HTMLResponse)
    async def login_page(request: Request):
        """Render login page."""
        return templates.TemplateResponse("pages/login.html", {"request": request})

    @app.post("/login", response_class=HTMLResponse)
    async def login_post(request: Request, db: aiosqlite.Connection = Depends(get_db)):
        """Handle login form submission."""
        try:
            form = await request.form()
            email = form.get("email", "").strip()
            password = form.get("password", "").strip()

            if not email or not password:
                return templates.TemplateResponse("pages/login.html", {
                    "request": request,
                    "error": "Email and password are required"
                }, status_code=400)

            cursor = await db.execute("SELECT * FROM users WHERE email = ?", (email,))
            user = await cursor.fetchone()

            if not user or not verify_password(password, user["password"]):
                return templates.TemplateResponse("pages/login.html", {
                    "request": request,
                    "error": "Invalid email or password"
                }, status_code=401)

            if not user["approved"]:
                return templates.TemplateResponse("pages/login.html", {
                    "request": request,
                    "error": "Account not approved yet"
                }, status_code=403)

            expiration = datetime.now(timezone.utc) + timedelta(hours=24)
            token_payload = {
                "sub": user["id"],
                "user_id": user["id"],
                "email": user["email"],
                "username": user["username"],
                "role": user["role"],
                "exp": expiration
            }
            token = jwt.encode(token_payload, SECRET_KEY, ALGORITHM)

            await db.execute(
                "INSERT INTO login_logs (user_id, email, username, role, login_at) VALUES (?, ?, ?, ?, ?)",
                (user["id"], user["email"], user["username"], user["role"],
                 datetime.now(timezone.utc).isoformat())
            )
            await db.commit()

            response = _RR(url="/", status_code=302)
            response.set_cookie("access_token", token, max_age=86400, httponly=True)
            return response

        except _HTMLAuthRequired:
            raise
        except Exception as e:
            logger.error(f"Login error: {e}")
            return templates.TemplateResponse("pages/login.html", {
                "request": request,
                "error": "An error occurred during login"
            }, status_code=500)

    @app.get("/logout")
    async def logout(request: Request):
        """Logout user."""
        response = _RR(url="/login", status_code=302)
        response.delete_cookie("access_token")
        return response

    # ── Dashboard ────────────────────────────────────────────────────────────
    @app.get("/", response_class=HTMLResponse)
    async def dashboard(request: Request, current_user: User = Depends(get_current_user_html)):
        """Render dashboard page. Data is fetched client-side from /api/dashboard/stats."""
        return templates.TemplateResponse("pages/dashboard.html", {
            "request": request,
            "user": current_user,
        })

    # ── All Requests ──────────────────────────────────────────────────────────
    @app.get("/requests", response_class=HTMLResponse)
    async def requests_list(
        request: Request,
        db: aiosqlite.Connection = Depends(get_db),
        current_user: User = Depends(get_current_user_html)
    ):
        """Render all requests list page."""
        try:
            cursor = await db.execute(
                "SELECT id, request_number, status, classification, created_at, created_by_username "
                "FROM requests ORDER BY created_at DESC"
            )
            reqs = await cursor.fetchall()
            return templates.TemplateResponse("pages/requests.html", {
                "request": request,
                "user": current_user,
                "requests": reqs
            })
        finally:
            await db.close()

    # ── My Requests ───────────────────────────────────────────────────────────
    @app.get("/my-requests", response_class=HTMLResponse)
    async def my_requests(
        request: Request,
        db: aiosqlite.Connection = Depends(get_db),
        current_user: User = Depends(get_current_user_html)
    ):
        """Render current user's own requests."""
        try:
            cursor = await db.execute(
                "SELECT id, request_number, status, created_at, updated_at "
                "FROM requests WHERE created_by = ? ORDER BY created_at DESC",
                (current_user.id,)
            )
            reqs = await cursor.fetchall()
            return templates.TemplateResponse("pages/my_requests.html", {
                "request": request,
                "user": current_user,
                "requests": reqs
            })
        finally:
            await db.close()

    # ── Request Detail / New Request ──────────────────────────────────────────
    @app.get("/requests/new", response_class=HTMLResponse)
    async def new_request(
        request: Request,
        current_user: User = Depends(get_current_user_html)
    ):
        """Render new request form."""
        return templates.TemplateResponse("pages/request_new.html", {
            "request": request,
            "user": current_user
        })

    @app.post("/requests/new", response_class=HTMLResponse)
    async def create_new_request(
        request: Request,
        db: aiosqlite.Connection = Depends(get_db),
        current_user: User = Depends(get_current_user_html)
    ):
        """Handle new request creation."""
        try:
            form = await request.form()
            req_id = str(uuid.uuid4())
            # Determine request type and generate sequential number
            request_type = (form.get("request_type") or "REL").strip().upper()
            if request_type not in ("REL", "RMS"):
                request_type = "REL"
            year = datetime.now(timezone.utc).year
            cursor = await db.execute(
                "SELECT MAX(CAST(SUBSTR(request_number, ?) AS INTEGER)) FROM requests WHERE request_number LIKE ?",
                (len(f"{request_type}{year}") + 1, f"{request_type}{year}%")
            )
            max_row = await cursor.fetchone()
            request_number = f"{request_type}{year}{(max_row[0] or 0) + 1:05d}"
            await db.execute(
                """INSERT INTO requests (id, request_number, request_type, classification, plant, device_name, lot_no,
                   customer, purpose, engineer_special_instruction, deadline, body_size_x, body_size_y,
                   ball_pitch, ball_count, created_by, created_by_username, created_at, updated_at, status)
                   VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
                (req_id, request_number, request_type, form.get("classification"),
                 form.get("plant"), form.get("device_name"), form.get("lot_no"),
                 form.get("customer"), form.get("purpose"), form.get("engineer_special_instruction"),
                 form.get("deadline"), form.get("body_size_x"), form.get("body_size_y"),
                 form.get("ball_pitch"), form.get("ball_count"), current_user.id,
                 current_user.username, datetime.now(timezone.utc).isoformat(),
                 datetime.now(timezone.utc).isoformat(), "pending")
            )
            await db.commit()
            return _RR(url=f"/requests/{req_id}", status_code=302)
        finally:
            await db.close()

    @app.get("/requests/{req_id}", response_class=HTMLResponse)
    async def request_detail(
        request: Request,
        req_id: str,
        db: aiosqlite.Connection = Depends(get_db),
        current_user: User = Depends(get_current_user_html)
    ):
        """Render request detail page."""
        try:
            cursor = await db.execute("SELECT * FROM requests WHERE id = ?", (req_id,))
            req_data = await cursor.fetchone()
            if not req_data:
                raise HTTPException(status_code=404, detail="Request not found")
            cursor = await db.execute(
                "SELECT id, leg, step_number, step_name, status, started_at, completed_at, "
                "machine_no, rack_no, operator_id, tray_no, qty_in, qty_out, notes, updated_by, custom_fields "
                "FROM process_steps WHERE request_id = ? ORDER BY leg, step_number",
                (req_id,)
            )
            raw_steps = await cursor.fetchall()
            steps = []
            for s in raw_steps:
                sd = {k: s[k] for k in s.keys()}
                sd['custom_fields'] = json.loads(s['custom_fields']) if s['custom_fields'] else {}
                steps.append(sd)
            e_cursor = await db.execute("SELECT id, name, position FROM employees ORDER BY name")
            raw_emps = await e_cursor.fetchall()
            employees = [{k: e[k] for k in e.keys()} for e in raw_emps]
            m_cursor = await db.execute("SELECT machine_no, description FROM machines ORDER BY machine_no")
            raw_machines = await m_cursor.fetchall()
            machines = [{k: m[k] for k in m.keys()} for m in raw_machines]
            # Convert request_data to dict for easier template access
            request_data = {k: req_data[k] for k in req_data.keys()}
            return templates.TemplateResponse("pages/request_detail.html", {
                "request": request,
                "user": current_user,
                "request_data": request_data,
                "steps": steps,
                "employees": employees,
                "machines": machines,
                "available_step_names": AVAILABLE_STEP_NAMES,
            })
        finally:
            await db.close()

    # ── Completed ─────────────────────────────────────────────────────────────
    @app.get("/completed", response_class=HTMLResponse)
    async def completed_requests(
        request: Request,
        db: aiosqlite.Connection = Depends(get_db),
        current_user: User = Depends(get_current_user_html)
    ):
        """Render completed requests page."""
        try:
            cursor = await db.execute(
                "SELECT * FROM requests WHERE status = 'completed' ORDER BY updated_at DESC"
            )
            reqs = await cursor.fetchall()
            return templates.TemplateResponse("pages/completed.html", {
                "request": request,
                "user": current_user,
                "completed_requests": reqs
            })
        finally:
            await db.close()

    # ── Approval Queue ────────────────────────────────────────────────────────
    @app.get("/approval", response_class=HTMLResponse)
    async def approval_queue(
        request: Request,
        db: aiosqlite.Connection = Depends(get_db),
        current_user: User = Depends(get_current_user_html)
    ):
        """Render approval queue page."""
        if current_user.role not in ["Admin", "ADMIN", "Planner", "PLANNER"]:
            raise HTTPException(status_code=403, detail="Insufficient permissions")
        try:
            cursor = await db.execute(
                "SELECT * FROM requests WHERE status = 'pending' ORDER BY created_at"
            )
            pending = await cursor.fetchall()
            return templates.TemplateResponse("pages/approval.html", {
                "request": request,
                "user": current_user,
                "pending_requests": pending,
                "stats": {
                    "pending": len(pending) if pending else 0,
                    "this_week": 0,
                    "overdue": 0,
                    "approved_this_month": 0
                },
                "now": datetime.now(timezone.utc)
            })
        finally:
            await db.close()

    # ── Retention Monitor ─────────────────────────────────────────────────────
    @app.get("/retention-monitor", response_class=HTMLResponse)
    async def retention_monitor(
        request: Request,
        db: aiosqlite.Connection = Depends(get_db),
        current_user: User = Depends(get_current_user_html)
    ):
        """Render retention monitor page (API-driven)."""
        try:
            return templates.TemplateResponse("pages/retention_monitor.html", {
                "request": request,
                "user": current_user,
            })
        finally:
            await db.close()

    # ── Filter & Analytics ────────────────────────────────────────────────────
    @app.get("/request-filter", response_class=HTMLResponse)
    async def request_filter(
        request: Request,
        db: aiosqlite.Connection = Depends(get_db),
        current_user: User = Depends(get_current_user_html)
    ):
        """Render RELDMS filter page (API-driven, no server-side data needed)."""
        try:
            return templates.TemplateResponse("pages/request_filter.html", {
                "request": request,
                "user": current_user,
            })
        finally:
            await db.close()

    # ── Loading / Unloading ───────────────────────────────────────────────────
    @app.get("/loading-unloading", response_class=HTMLResponse)
    async def loading_unloading(
        request: Request,
        search: Optional[str] = None,
        step_filter: Optional[str] = None,
        db: aiosqlite.Connection = Depends(get_db),
        current_user: User = Depends(get_current_user_html)
    ):
        """Render loading/unloading monitor page."""
        try:
            placeholders = ','.join('?' * len(MONITORED_STEP_NAMES))
            query = f"""
                SELECT ps.id, ps.request_id, ps.leg, ps.step_number, ps.step_name,
                       ps.status, ps.started_at, ps.completed_at,
                       ps.machine_no, ps.rack_no,
                       r.request_number, r.device_name, r.customer, r.lot_no,
                       COALESCE(e.name, '') AS employee_name,
                       COALESCE(m.description, '') AS machine_desc
                FROM process_steps ps
                JOIN requests r ON ps.request_id = r.id
                LEFT JOIN employees e ON ps.operator_id = e.id
                LEFT JOIN machines m ON LOWER(ps.machine_no) = LOWER(m.machine_no)
                WHERE LOWER(ps.step_name) IN ({placeholders})
                  AND r.status NOT IN ('discontinued')
                  AND ps.status IN ('in_progress', 'completed')
                  AND ps.started_at IS NOT NULL
                ORDER BY
                    CASE ps.status WHEN 'in_progress' THEN 0 ELSE 1 END,
                    ps.started_at DESC
                LIMIT 300
            """
            cursor = await db.execute(query, MONITORED_STEP_NAMES)
            rows = await cursor.fetchall()
            ops = []
            for row in rows:
                machine_display = row['machine_no'] or ''
                if row['machine_desc']:
                    machine_display = f"{row['machine_no']} — {row['machine_desc']}"
                ops.append({
                    'id': row['id'],
                    'request_id': row['request_id'],
                    'step_name': row['step_name'],
                    'leg': row['leg'],
                    'step_number': row['step_number'],
                    'status': row['status'],
                    'started_at': _safe_isoparse(row['started_at']),
                    'completed_at': _safe_isoparse(row['completed_at']),
                    'machine_no': row['machine_no'] or '',
                    'rack_no': row['rack_no'] or '',
                    'machine_display': machine_display,
                    'request_number': row['request_number'],
                    'device_name': row['device_name'] or '',
                    'customer': row['customer'] or '',
                    'lot_no': row['lot_no'] or '',
                    'employee_name': row['employee_name'] or '',
                })
            # Apply client-side filters
            if search:
                s = search.lower()
                ops = [op for op in ops if
                    s in op['machine_no'].lower() or
                    s in op['machine_display'].lower() or
                    s in op['request_number'].lower() or
                    s in op['step_name'].lower() or
                    s in op['employee_name'].lower() or
                    s in op['device_name'].lower()]
            if step_filter:
                sf = step_filter.lower()
                ops = [op for op in ops if op['step_name'].lower() == sf]
            active_ops = [op for op in ops if op['status'] == 'in_progress']
            history_ops = [op for op in ops if op['status'] == 'completed']
            today_str = datetime.now(timezone.utc).strftime('%Y-%m-%d')
            completed_today = sum(
                1 for op in history_ops
                if op['completed_at'] and op['completed_at'].strftime('%Y-%m-%d') == today_str
            )
            durations = []
            for op in history_ops:
                if op['started_at'] and op['completed_at']:
                    try:
                        durations.append((op['completed_at'] - op['started_at']).total_seconds())
                    except Exception:
                        pass
            avg_secs = sum(durations) / len(durations) if durations else 0
            avg_h = int(avg_secs // 3600)
            avg_m = int((avg_secs % 3600) // 60)
            return templates.TemplateResponse("pages/loading_unloading.html", {
                "request": request,
                "user": current_user,
                "stats": {
                    "active_jobs": len(active_ops),
                    "completed_today": completed_today,
                    "scheduled": 0,
                    "avg_duration": f"{avg_h}h {avg_m}m"
                },
                "active_operations": active_ops,
                "history_operations": history_ops,
                "search": search or '',
                "step_filter": step_filter or '',
                "monitored_step_names": MONITORED_STEP_NAMES,
            })
        finally:
            await db.close()

    # ── Performance Monitor ───────────────────────────────────────────────────
    @app.get("/performance", response_class=HTMLResponse)
    async def performance_monitor(
        request: Request,
        period: Optional[str] = 'week',
        db: aiosqlite.Connection = Depends(get_db),
        current_user: User = Depends(get_current_user_html)
    ):
        """Render performance monitor page."""
        try:
            cursor = await db.execute("SELECT id, name, position FROM employees ORDER BY name")
            employees = await cursor.fetchall()
            return templates.TemplateResponse("pages/performance_monitor.html", {
                "request": request,
                "user": current_user,
                "employees": employees,
                "default_period": period,
            })
        finally:
            await db.close()

    # ── Task Manager (ADMIN) ──────────────────────────────────────────────────
    @app.get("/task-manager", response_class=HTMLResponse)
    async def task_manager(
        request: Request,
        current_user: User = Depends(get_current_user_html)
    ):
        """Render task manager page (API-driven)."""
        if current_user.role not in ["Admin", "ADMIN"]:
            raise HTTPException(status_code=403, detail="Insufficient permissions")
        return templates.TemplateResponse("pages/task_manager.html", {
            "request": request,
            "user": current_user,
        })

    # ── Backup Viewer (ADMIN) ─────────────────────────────────────────────────
    @app.get("/backup-viewer", response_class=HTMLResponse)
    async def backup_viewer(
        request: Request,
        current_user: User = Depends(get_current_user_html)
    ):
        """Render backup viewer page."""
        if current_user.role not in ("Admin", "ADMIN"):
            raise HTTPException(status_code=403, detail="Insufficient permissions")
        return templates.TemplateResponse("pages/backup_viewer.html", {
            "request": request,
            "user": current_user,
        })

    # ── User Management (ADMIN) ───────────────────────────────────────────────
    @app.get("/users", response_class=HTMLResponse)
    async def users_page(
        request: Request,
        current_user: User = Depends(get_current_user_html)
    ):
        """Render user management page (admin only)."""
        if current_user.role not in ("Admin", "ADMIN"):
            raise HTTPException(status_code=403, detail="Insufficient permissions")
        return templates.TemplateResponse("pages/users.html", {
            "request": request,
            "user": current_user,
        })

    # ── Settings (ADMIN) ──────────────────────────────────────────────────────
    @app.get("/settings", response_class=HTMLResponse)
    async def settings_page(
        request: Request,
        current_user: User = Depends(get_current_user_html)
    ):
        """Render settings page."""
        return templates.TemplateResponse("pages/settings.html", {
            "request": request,
            "user": current_user
        })


@app.on_event("startup")
async def startup():
    await init_db()
    # Start monthly auto-backup scheduler
    asyncio.create_task(monthly_backup_scheduler())
    # Auto-delete declined users after 5-minute grace period
    asyncio.create_task(declined_user_cleanup())
    logger.info("Application started - SQLite database ready, auto-backup scheduler active")

@app.on_event("shutdown")
async def shutdown():
    logger.info("Application shutting down")


if __name__ == "__main__":
    import uvicorn
    # ── Load configuration from environment
    port = int(os.getenv("PORT", "8000"))
    host = os.getenv("HOST", "0.0.0.0")  # Listen on all network interfaces for LAN sharing

    logger.info(f"Starting server on {host}:{port}...")
    uvicorn.run(
        app,
        host=host,
        port=port,
        log_level="info"
    )
